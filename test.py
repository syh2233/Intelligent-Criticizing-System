from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()
doc.add_heading('答题卡', 0)

# Add student information section
doc.add_paragraph('姓名：_________________________')
doc.add_paragraph('学号：_________________________')
doc.add_paragraph('科目：_________________________')
doc.add_paragraph('日期：_________________________')

# Add instructions
doc.add_paragraph('请选择每个问题的正确答案，涂黑相应的圆圈。')

# Create a table for the answer choices (5 rows, 4 columns)
table = doc.add_table(rows=51, cols=5)

# Fill the header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '题号'
hdr_cells[1].text = 'A'
hdr_cells[2].text = 'B'
hdr_cells[3].text = 'C'
hdr_cells[4].text = 'D'

# Fill the table with question numbers and empty answer spaces
for i in range(1, 51):
    row_cells = table.rows[i].cells
    row_cells[0].text = str(i)
    row_cells[1].text = '◯'
    row_cells[2].text = '◯'
    row_cells[3].text = '◯'
    row_cells[4].text = '◯'

# Save the document
file_path = "答题卡.docx"
doc.save(file_path)

file_path
