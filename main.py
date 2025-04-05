from datetime import datetime, timedelta
from functools import wraps
import os
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_from_directory, flash
from docx import Document
import sqlite3
from sqlite3 import Error
from werkzeug.utils import secure_filename
import time
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from flask_login import LoginManager, login_user, login_required as flask_login_required, current_user, logout_user, UserMixin
from openpyxl.comments import Comment
import zipfile
import shutil
import logging
from split_and_ocr.pdf_ocr import process_pdf
import json
from split_and_ocr.slip import split_columns_and_rows
import re
from flask_login import LoginManager, login_user, login_required as flask_login_required, current_user, logout_user, UserMixin
import uuid
import sys
from split_and_ocr.ai import new

# 创建日志目录
log_dir = 'logs'
os.makedirs(log_dir, exist_ok=True)

# 配置根日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(log_dir, 'app.log')),
        logging.StreamHandler(sys.stdout)
    ]
)

# 创建应用程序日志
logger = logging.getLogger('exam_app')
logger.setLevel(logging.DEBUG)

# 添加文件处理器
file_handler = logging.FileHandler(os.path.join(log_dir, 'exam_debug.log'))
file_handler.setLevel(logging.DEBUG)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))

# 添加控制台处理器
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))

# 为日志添加处理器
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# 创建Flask应用
app = Flask(__name__)
app.secret_key = 'your-secret-key'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['EXTRACTED_FILES'] = 'extracted_files'
app.config['EXPORT_FOLDER'] = 'exports'
app.debug = True

# 替换app.logger为我们自定义的logger
app.logger = logger

app.logger.info("应用启动")

# 初始化数据存储目录
for folder in [app.config['UPLOAD_FOLDER'], app.config['EXTRACTED_FILES'], app.config['EXPORT_FOLDER']]:
    os.makedirs(folder, exist_ok=True)
    app.logger.info(f"确保目录存在: {folder}")

# 定义上传文件夹和解压文件夹路径
EXAM_UPLOAD_FOLDER = 'uploads'
EXTRACT_FOLDER = 'extracted_files'

# 定义允许的文件类型
ALLOWED_EXTENSIONS = {'pdf', 'zip', 'rar'}

# 确保上传和解压目录存在
os.makedirs(EXAM_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXTRACT_FOLDER, exist_ok=True)

# 添加图片上传相关配置
IMAGE_UPLOAD_FOLDER = 'static/uploads/images'
if not os.path.exists(IMAGE_UPLOAD_FOLDER):
    os.makedirs(IMAGE_UPLOAD_FOLDER)
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

class CustomReloader:
    def should_reload(self, filename):
        # 忽略某些目录的文件变化
        ignore_dirs = [
            'uploads',
            'temp_images',
            'read',
            '__pycache__',
            '.git',
            'venv',
            'env'
        ]
        
        # 检查文件是否在忽略目录中
        for ignore_dir in ignore_dirs:
            if ignore_dir in filename:
                return False
                
        return super().should_reload(filename)

def create_app(secret_keyss):
    app = Flask(__name__)
    app.config['EXAM_UPLOAD_FOLDER'] = EXAM_UPLOAD_FOLDER
    app.config['IMAGE_UPLOAD_FOLDER'] = IMAGE_UPLOAD_FOLDER
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max-limit
    
    # 设置session密钥
    app.secret_key = '{secret_keyss}'  
    
    # 禁用文件更改自动重载
    app.config['DEBUG'] = True
    app.config['TEMPLATES_AUTO_RELOAD'] = False
    
    return app, app.secret_key

app,app.secret_key = create_app('syh2031.')

# 添加额外的配置来忽略特定目录
def should_ignore_file(filename):
    ignore_dirs = {EXAM_UPLOAD_FOLDER, IMAGE_UPLOAD_FOLDER, EXTRACT_FOLDER}
    return any(dir_name in filename for dir_name in ignore_dirs)

# 修改 Flask 的 reloader
extra_files = None
if app.debug:
    extra_files = [f for f in app.static_folder if not should_ignore_file(f)]

# 允许的文件类型定义
ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg'}
ALLOWED_DOC_EXTENSIONS = {'doc', 'docx'}
ALLOWED_ZIP_EXTENSIONS = {'zip'}

def calculate_duration(start_time, end_time):
    """计算考试时长"""
    try:
        # 将字符串转换为 datetime 对象
        if isinstance(start_time, str):
            start = datetime.strptime(start_time, '%Y-%m-%d %H:%M:%S')
        else:
            start = start_time
            
        if isinstance(end_time, str):
            end = datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S')
        else:
            end = end_time
        
        # 计算时间差
        duration = end - start
        
        # 转换为小时和分钟
        total_minutes = duration.total_seconds() / 60
        hours = int(total_minutes // 60)
        minutes = int(total_minutes % 60)
        
        # 格式化输出
        if hours > 0:
            if minutes > 0:
                return f"{hours}小时{minutes}分钟"
            return f"{hours}小时"
        return f"{minutes}分钟"
        
    except Exception as e:
        print(f"计算时长出错: {e}")
        return "时长未知"

def allowed_file(filename, allowed_extensions=None):
    """检查文件是否具有允许的扩展名"""
    if allowed_extensions is None:
        allowed_extensions = ['pdf', 'docx', 'doc']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

# 添加登录验证装饰器
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'email' not in session:
            return redirect(url_for('index_login'))
        return f(*args, **kwargs)

    return decorated_function


@app.route('/')
def index_login():
    if 'email' in session:
        return redirect(url_for('xinzeng'))
    return render_template('login.html')


@app.route('/register')
def register_page():
    return render_template('register.html')


@app.route('/dashboard')
def dashboard():
    if 'email' in session:
        return redirect(url_for('xinzeng'))
    else:
        return redirect(url_for('index_login'))


@app.route('/review')
def review():
    if 'email' in session:
        return render_template('review.html')
    else:
        return redirect(url_for('index_login'))


@app.route('/xinzeng')
def xinzeng():
    if 'email' in session:
        return render_template('xinzeng.html')
    else:
        return redirect(url_for('index_login'))


@app.route('/user')
def user():
    if 'email' in session:
        return render_template('user.html')
    else:
        return redirect(url_for('index_login'))


@app.route('/analysis')
def analysis():
    if 'email' in session:
        return render_template('analysis.html')
    else:
        return redirect(url_for('index_login'))


@app.route('/exam')
def exam():
    if 'email' in session:
        # 修改为重定向到考试列表，而不是直接渲染exam.html
        return redirect(url_for('exam_list'))
    else:
        return redirect(url_for('index_login'))


@app.route('/index')
def index():
    if 'email' in session:
        return render_template('index.html')
    else:
        return redirect(url_for('index_login'))


@app.route('/exam/<int:exam_id>')
@login_required
def exam_with_id(exam_id):
    try:
        app.logger.info(f"尝试访问考试ID: {exam_id}")
        
        # 获取考试基本信息
        exam = get_exam_by_id(exam_id)
        if not exam:
            app.logger.warning(f"找不到考试ID: {exam_id}")
            flash('找不到该考试或数据库连接错误')
            return redirect(url_for('exam_list'))
        
        # 记录更多细节
        app.logger.info(f"考试详情: {exam}")
        
        # 检查考试是否已经结束 (注释掉这部分，允许查看已结束考试)
        now = datetime.now()
        app.logger.info(f"当前时间: {now}, 考试结束时间: {exam['end_time']}")
        
        # 允许查看已结束的考试 (不再重定向)
        if now > exam['end_time']:
            app.logger.info(f"考试 {exam_id} 已结束，但允许用户查看")
            # 设置一个标志表示考试已结束
            exam['is_ended'] = True
            # 设置剩余时间为0
            exam['remaining_seconds'] = 0
        else:
            exam['is_ended'] = False
            # 计算剩余时间（以秒为单位）
            remaining_seconds = (exam['end_time'] - now).total_seconds()
            exam['remaining_seconds'] = int(remaining_seconds)
        
        app.logger.info(f"剩余秒数: {exam['remaining_seconds']}")
        
        # 检查考试是否已开始
        app.logger.info(f"考试开始时间: {exam['start_time']}")
        if now < exam['start_time']:
            app.logger.info(f"考试 {exam_id} 尚未开始")
            flash('该考试还未开始')
            return redirect(url_for('exam_list'))
        
        # 获取考试题目
        questions = get_exam_questions(exam_id)
        app.logger.info(f"题目获取结果: {questions}")
        
        # 检查是否有题目
        if not questions:
            app.logger.warning(f"考试 {exam_id} 没有题目")
            flash('没有找到考试题目')
            return redirect(url_for('exam_list'))
            
        has_questions = (
            len(questions['multiple_choice']) > 0 or 
            len(questions['fill_blank']) > 0 or 
            len(questions['short_answer']) > 0 or 
            len(questions['true_false']) > 0 or
            len(questions['programming']) > 0
        )
        
        if not has_questions:
            app.logger.warning(f"考试 {exam_id} 所有题型都没有题目")
            # 给用户友好提示
            flash('该考试没有可用的题目，请联系管理员')
            return redirect(url_for('exam_list'))
        
        # 获取当前用户ID
        user_id = get_current_user_id()
        app.logger.info(f"用户 {user_id} 尝试参加考试 {exam_id}")
        
        if not user_id:
            app.logger.warning(f"无法获取当前用户ID")
            flash('用户信息不完整，请重新登录')
            return redirect(url_for('index_login'))
        
        # 检查用户是否已参加过此考试
        has_taken = False
        try:
            has_taken = has_taken_exam(user_id, exam_id)
            app.logger.info(f"用户是否已参加过考试: {has_taken}")
        except Exception as e:
            app.logger.error(f"检查用户是否参加过考试时出错: {str(e)}", exc_info=True)
            # 如果无法确定，我们假设用户没有参加过
            has_taken = False
        
        # 允许用户继续参加考试，不再检查是否参加过
        # 如果考试已结束，不再检查用户是否参加过
        # if user_id and has_taken and not exam['is_ended']:
        #     app.logger.info(f"用户 {user_id} 已参加过考试 {exam_id}")
        #     flash('您已经参加过这个考试')
        #     return redirect(url_for('exam_list'))
        
        # 记录用户开始考试
        app.logger.info(f"用户 {user_id} 成功进入考试 {exam_id}")
        
        # 将所有题目数量添加到日志
        app.logger.info(f"选择题数量: {len(questions['multiple_choice'])}")
        app.logger.info(f"填空题数量: {len(questions['fill_blank'])}")
        app.logger.info(f"简答题数量: {len(questions['short_answer'])}")
        app.logger.info(f"判断题数量: {len(questions['true_false'])}")
        app.logger.info(f"编程题数量: {len(questions['programming'])}")
        
        # 尝试渲染模板
        try:
            return render_template('exam.html', exam=exam, questions=questions)
        except Exception as template_error:
            app.logger.error(f"模板渲染错误: {str(template_error)}", exc_info=True)
            flash('渲染考试页面时出错，请联系管理员')
            return redirect(url_for('exam_list'))
            
    except Exception as e:
        app.logger.error(f"进入考试出错: {str(e)}", exc_info=True)
        flash('加载考试时出错，请稍后再试')
        return redirect(url_for('exam_list'))


def get_db_connection():
    """获取数据库连接"""
    # 确保数据库目录存在
    db_dir = 'database'
    os.makedirs(db_dir, exist_ok=True)
    
    # 构建数据库文件的绝对路径
    db_path = os.path.join(os.path.abspath(db_dir), 'It_g.db')
    
    max_attempts = 3
    attempt = 0
    while attempt < max_attempts:
        try:
            conn = sqlite3.connect(db_path, timeout=20)
            conn.row_factory = sqlite3.Row
            
            return conn
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e):
                attempt += 1
                if attempt < max_attempts:
                    print(f"数据库被锁定，正在重试... (尝试 {attempt}/{max_attempts})")
                    time.sleep(1)  # 等待1秒后重试
                    continue
            print(f"数据库连接错误: {e}")
            return None
        except Exception as e:
            print(f"数据库连接错误: {e}")
            return None
    return None

def execute_query(query, params=None):
    """执行SQL查询"""
    conn = get_db_connection()
    if not conn:
        print("数据库连接失败")
        return False, None

    cursor = None
    result = None
    try:
        cursor = conn.cursor()
        cursor.execute(query, params or ())

        if query.strip().upper().startswith(('INSERT', 'UPDATE', 'DELETE')):
            conn.commit()
            result = cursor.lastrowid
        else:
            result = cursor.fetchall()
        return True, result
    except Error as e:
        print(f"SQL执行错误: {e}")
        print(f"错误的查询: {query}")
        print(f"参数: {params}")
        import traceback
        traceback.print_exc()  # 打印完整的错误堆栈
        if conn:
            conn.rollback()
        return False, None
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/login', methods=['POST'])
def login():
    email = request.form.get('email')
    pwd = request.form.get('password')

    # 使用参数化查询防止SQL注入
    query = "SELECT * FROM users WHERE email = ?"
    success, result = execute_query(query, (email,))
    
    if not success:
        return '数据库错误 <a href="/">返回登录</a>'
    
    if result and len(result) > 0:
        user = result[0]
        if pwd == user['password']:  # 实际应用中应使用密码哈希
            session["email"] = email
            
            # 记录登录操作
            log_query = """
                INSERT INTO operation_logs (user_id, operation_type, operation_details) 
                VALUES (?, ?, ?)
            """
            execute_query(log_query, (user['id'], 'login', f'User logged in from {request.remote_addr}'))
            
            return redirect(url_for('xinzeng'))
        else:
            return '密码错误 <a href="/">返回登录</a>'
    else:
        return '用户名不存在 <a href="/">返回登录</a>'

@app.route('/register', methods=['POST'])
def register():
    email = request.form.get('email')
    pwd = request.form.get('password')
    pwd2 = request.form.get('confirm_password')
    name = request.form.get('name')
    student_id = request.form.get('student_id')

    if pwd != pwd2:
        return '两次密码不一致 <a href="/register">返回注册</a>'

    # 检查用户是否已存在
    check_query = "SELECT * FROM users WHERE email = ?"
    success, result = execute_query(check_query, (email,))

    if not success:
        return '数据库错误 <a href="/register">返回注册</a>'

    if result and len(result) > 0:
        return '用户名已存在 <a href="/register">返回注册</a>'

    # 检查学号是否已存在
    check_student_query = "SELECT * FROM students WHERE student_id = ?"
    success, result = execute_query(check_student_query, (student_id,))

    if not success:
        return '数据库错误 <a href="/register">返回注册</a>'

    if result and len(result) > 0:
        return '学号已存在 <a href="/register">返回注册</a>'

    # 开始事务
    conn = get_db_connection()
    conn.execute('BEGIN TRANSACTION')

    try:
        # 创建新用户
        insert_query = "INSERT INTO users (email, password) VALUES (?, ?)"
        success, user_id = execute_query(insert_query, (email, pwd))

        if not success:
            conn.rollback()
            return '注册失败 <a href="/register">返回注册</a>'

        # 创建学生记录
        insert_student_query = "INSERT INTO students (name, student_id, user_id) VALUES (?, ?, ?)"
        success, _ = execute_query(insert_student_query, (name, student_id, user_id))

        if not success:
            conn.rollback()
            return '学生信息创建失败 <a href="/register">返回注册</a>'

        # 记录注册操作
        log_query = """
            INSERT INTO operation_logs (user_id, operation_type, operation_details) 
            VALUES (?, ?, ?)
        """
        execute_query(log_query, (user_id, 'register', f'New user registration from {request.remote_addr}'))

        # 提交事务
        conn.commit()

        return '注册成功 <a href="/">返回登录</a>'
    except Exception as e:
        conn.rollback()
        return f'注册失败: {str(e)} <a href="/register">返回注册</a>'
    finally:
        conn.close()

@app.route('/logout')
def logout():
    session.pop('email', None)
    return redirect(url_for('index_login'))

def sanitize_filename(filename):
    """处理文件名，保留中文字符"""
    # 替换不安全的字符
    unsafe_chars = '/\\?%*:|"<>'
    filename = ''.join(c if c not in unsafe_chars else '_' for c in filename)
    return filename.strip('.')

@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传请求"""
    try:
        # 检查是否有文件被上传
        if not request.files:
            logger.warning('没有文件被上传')
            return jsonify({'error': '没有文件被上传'}), 400
            
        file_type = request.form.get('type', '')  # 'image' 或 'document'
        if not file_type:
            logger.warning('未指定文件类型')
            return jsonify({'error': '未指定文件类型'}), 400
            
        # 为此次上传创建唯一的目录
        timestamp = str(int(time.time()))
        extract_dir = os.path.join(EXTRACT_FOLDER, timestamp)
        os.makedirs(extract_dir, exist_ok=True)
        logger.info(f'创建上传目录: {extract_dir}')
        
        files_info = []
        temp_files = []  # 用于跟踪需要清理的临时文件
        
        # 处理所有上传的文件
        for key in request.files:
            file = request.files[key]
            if file.filename == '':
                continue
                
            # 处理文件名，确保中文文件名也能正常工作
            original_filename = file.filename
            # 使用自定义的文件名清理函数
            filename = sanitize_filename(original_filename)
            if filename == '':  # 如果文件名为空
                filename = f"{str(int(time.time()))}_{key}"  # 使用时间戳作为文件名
                
            if '.' not in filename:
                logger.warning(f'无效的文件名: {original_filename}')
                continue
                
            file_ext = filename.rsplit('.', 1)[1].lower()
            logger.info(f'处理文件: {filename} (类型: {file_ext})')
            
            # 检查文件类型
            if file_type == 'image':
                if file_ext not in ALLOWED_IMAGE_EXTENSIONS and file_ext not in ALLOWED_ZIP_EXTENSIONS:
                    logger.warning(f'不支持的图片文件类型: {file_ext}')
                    continue
            elif file_type == 'document':
                if file_ext not in ALLOWED_DOC_EXTENSIONS and file_ext not in ALLOWED_ZIP_EXTENSIONS:
                    logger.warning(f'不支持的文档文件类型: {file_ext}')
                    continue
            else:
                logger.error(f'无效的文件类型: {file_type}')
                return jsonify({'error': '无效的文件类型'}), 400
                
            # 保存文件
            file_path = os.path.join(extract_dir, filename)
            file.save(file_path)
            temp_files.append(file_path)
            logger.info(f'文件已保存: {file_path}')
            
            # 如果是压缩包，解压处理
            if file_ext in ALLOWED_ZIP_EXTENSIONS:
                zip_ref = None
                try:
                    zip_ref = zipfile.ZipFile(file_path, 'r', allowZip64=True)
                    # 创建子目录用于解压文件
                    zip_extract_dir = os.path.join(extract_dir, os.path.splitext(filename)[0])
                    os.makedirs(zip_extract_dir, exist_ok=True)
                    
                    # 获取压缩包中的所有文件
                    for zip_info in zip_ref.filelist:
                        try:
                            # 尝试不同的编码方式
                            try:
                                # 首先尝试 UTF-8
                                name = zip_info.filename.encode('cp437').decode('utf-8')
                            except UnicodeDecodeError:
                                try:
                                    # 然后尝试 GBK
                                    name = zip_info.filename.encode('cp437').decode('gbk')
                                except UnicodeDecodeError:
                                    # 如果都失败，保持原样
                                    name = zip_info.filename
                            
                            # 清理文件名
                            clean_name = sanitize_filename(name)
                            
                            # 如果是目录，创建它
                            if zip_info.filename.endswith('/'):
                                target_dir = os.path.join(zip_extract_dir, clean_name)
                                os.makedirs(target_dir, exist_ok=True)
                                continue
                            
                            # 提取文件
                            source = None
                            target = None
                            try:
                                source = zip_ref.open(zip_info)
                                target_path = os.path.join(zip_extract_dir, clean_name)
                                
                                # 确保目标目录存在
                                os.makedirs(os.path.dirname(target_path), exist_ok=True)
                                
                                with open(target_path, 'wb') as target:
                                    shutil.copyfileobj(source, target)
                            finally:
                                # 确保文件句柄被关闭
                                if source:
                                    source.close()
                            
                        except Exception as e:
                            logger.error(f'处理压缩文件 {zip_info.filename} 时出错: {str(e)}')
                            continue
                    
                    logger.info(f'解压文件成功: {file_path} 到 {zip_extract_dir}')
                    
                    # 获取解压后的文件列表
                    for root, _, files in os.walk(zip_extract_dir):
                        for f in files:
                            try:
                                f_ext = f.rsplit('.', 1)[1].lower() if '.' in f else ''
                                if ((file_type == 'image' and f_ext in ALLOWED_IMAGE_EXTENSIONS) or
                                    (file_type == 'document' and f_ext in ALLOWED_DOC_EXTENSIONS)):
                                    file_path = os.path.join(root, f)
                                    relative_path = os.path.relpath(file_path, extract_dir)
                                    files_info.append({
                                        'name': f,
                                        'path': relative_path,
                                        'original_name': f  # 使用解压后的文件名作为原始名称
                                    })
                                    logger.info(f'处理解压文件: {relative_path}')
                            except Exception as e:
                                logger.error(f'处理解压后的文件 {f} 时出错: {str(e)}')
                                continue
                
                except Exception as e:
                    logger.error(f'解压文件出错: {str(e)}')
                    # 如果解压失败，将zip文件作为普通文件处理
                    files_info.append({
                        'name': filename,
                        'path': filename,
                        'original_name': original_filename
                    })
                
                finally:
                    # 确保zip文件被关闭
                    if zip_ref:
                        zip_ref.close()
                    
                    # 等待一小段时间确保文件句柄被释放
                    time.sleep(0.1)
                    
                    # 删除原始压缩包
                    try:
                        if os.path.exists(temp_files[-1]):
                            os.remove(temp_files[-1])  # 删除最后添加的文件（压缩包）
                            temp_files.pop()  # 从临时文件列表中移除
                            logger.info(f'删除原始压缩包: {file_path}')
                    except Exception as e:
                        logger.error(f'删除压缩包失败: {str(e)}')
                        # 如果删除失败，将在后续清理中处理
            else:
                # 处理普通文件
                files_info.append({
                    'name': filename,
                    'path': filename,
                    'original_name': original_filename
                })
        
        if not files_info:
            logger.warning('没有有效的文件被上传')
            # 清理所有临时文件
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        logger.info(f'清理临时文件: {temp_file}')
                except Exception as e:
                    logger.error(f'清理文件失败: {temp_file} - {str(e)}')
            return jsonify({'error': '没有有效的文件被上传'}), 400
            
        logger.info(f'文件上传成功: {len(files_info)} 个文件')
        return jsonify({
            'message': '文件上传成功',
            'extract_dir': timestamp,
            'files': files_info,
            'type': file_type
        })
        
    except Exception as e:
        logger.error(f'文件上传错误: {str(e)}')
        # 清理临时文件
        if 'extract_dir' in locals() and os.path.exists(extract_dir):
            try:
                shutil.rmtree(extract_dir)
                logger.info(f'清理临时目录: {extract_dir}')
            except Exception as cleanup_error:
                logger.error(f'清理临时目录失败: {str(cleanup_error)}')
        return jsonify({'error': f'处理文件时出错: {str(e)}'}), 500

@app.route('/upload-word', methods=['POST'])
def upload_word_file():
    """处理Word文档上传请求"""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件被上传'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '没有选择文件'}), 400
        
    try:
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        # 检查是否是Word文档或ZIP文件
        if file_ext not in ALLOWED_DOC_EXTENSIONS and file_ext not in ALLOWED_ZIP_EXTENSIONS:
            return jsonify({'error': '不支持的文件类型'}), 400
            
        # 为此次上传创建唯一的目录
        extract_dir = os.path.join(EXTRACT_FOLDER, str(int(time.time())))
        os.makedirs(extract_dir, exist_ok=True)
        
        files_info = []
        
        # 处理ZIP文件
        if file_ext in ALLOWED_ZIP_EXTENSIONS:
            zip_path = os.path.join(EXAM_UPLOAD_FOLDER, filename)
            file.save(zip_path)
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
                
            # 处理解压出的Word文档
            for root, _, files in os.walk(extract_dir):
                for f in files:
                    if f.lower().endswith(tuple(ALLOWED_DOC_EXTENSIONS)):
                        doc_path = os.path.join(root, f)
                        # 读取Word文档内容
                        doc = Document(doc_path)
                        # 提取文本
                        text = "\n".join([para.text for para in doc.paragraphs])
                        # 保存为txt文件
                        txt_filename = f"{os.path.splitext(f)[0]}.txt"
                        txt_path = os.path.join(root, txt_filename)
                        with open(txt_path, 'w', encoding='utf-8') as txt_file:
                            txt_file.write(text)
                        
                        relative_path = os.path.relpath(txt_path, extract_dir)
                        files_info.append({
                            'name': txt_filename,
                            'path': relative_path,
                            'original': f
                        })
            
            os.remove(zip_path)  # 删除原始压缩包
            
        # 处理单个Word文档
        else:
            file_path = os.path.join(extract_dir, filename)
            file.save(file_path)
            
            # 读取Word文档内容
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # 保存为txt文件
            txt_filename = f"{os.path.splitext(filename)[0]}.txt"
            txt_path = os.path.join(extract_dir, txt_filename)
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
            files_info.append({
                'name': txt_filename,
                'path': txt_filename,
                'original': filename
            })
            
            # 可以选择是否保留原始Word文件
            os.remove(file_path)
        
        return jsonify({
            'message': '文件处理成功',
            'extract_dir': os.path.basename(extract_dir),
            'files': files_info
        })
        
    except Exception as e:
        # 清理临时文件
        if os.path.exists(extract_dir):
            shutil.rmtree(extract_dir)
        return jsonify({'error': f'处理文件时出错: {str(e)}'}), 500

@app.route('/exam/list')
@login_required  # 保留您现有的登录装饰器
def exam_list():
    # 获取当前时间
    now = datetime.now()
    
    # 获取所有考试基本信息
    query = """
        SELECT id, name, subject, duration, exam_score as total_score, 
               start_time, end_time, status, exam_file_path
        FROM exam_sessions
        ORDER BY start_time DESC
    """
    success, exams_data = execute_query(query)
    
    if not success:
        print("获取考试数据失败")
        return render_template('exam_list.html', exams=[])
    
    # 获取所有学生考试成绩信息
    score_query = """
        SELECT session_id, student_id, score, status
        FROM student_exams
        WHERE status = 'graded'
    """
    success, scores = execute_query(score_query)
    
    # 整理学生成绩数据
    student_exams = {}
    if success and scores:
        for score in scores:
            if score['session_id'] not in student_exams:
                student_exams[score['session_id']] = []
            student_exams[score['session_id']].append({
                'student_id': score['student_id'],
                'score': score['score'],
                'status': score['status']
            })
    
    # 处理考试列表
    exams = []
    subjects = set()
    
    for exam in exams_data:
        try:
            # 转换时间格式
            if '.' in exam['start_time']:
                start_time = datetime.strptime(exam['start_time'], '%Y-%m-%d %H:%M:%S.%f')
            else:
                start_time = datetime.strptime(exam['start_time'], '%Y-%m-%d %H:%M:%S')
                
            if '.' in exam['end_time']:
                end_time = datetime.strptime(exam['end_time'], '%Y-%m-%d %H:%M:%S.%f')
            else:
                end_time = datetime.strptime(exam['end_time'], '%Y-%m-%d %H:%M:%S')
            
            # 确定考试状态
            status = exam['status']
            
            # 收集科目
            subject = exam['subject']
            if subject:
                subjects.add(subject)
            
            # 构建考试数据
            exam_item = {
                'id': exam['id'],
                'name': exam['name'],
                'subject': subject,
                'duration': exam['duration'],
                'total_score': exam['total_score'],
                'start_time': start_time,
                'end_time': end_time,
                'status': status
            }
            
            # 添加成绩信息
            if exam['id'] in student_exams:
                # 这里我们只显示是否有人已评分，而不是特定用户的分数
                # 因为我们没有当前用户的ID
                exam_item['has_graded_scores'] = True
                
                # 对于前端显示，我们仍然保留score字段，但设置为字符串"已评分"
                # 如果需要数值，可以选择第一个学生的分数作为示例
                if student_exams[exam['id']]:
                    exam_item['score'] = student_exams[exam['id']][0]['score']
                else:
                    exam_item['score'] = None
            else:
                exam_item['has_graded_scores'] = False
                exam_item['score'] = None
            
            exams.append(exam_item)
        except Exception as e:
            print(f"处理考试数据时出错: {str(e)}")
            continue
    
    # 按状态分类考试
    ongoing_exams = [e for e in exams if e['status'] == 'ongoing']
    upcoming_exams = [e for e in exams if e['status'] == 'pending']
    completed_exams = [e for e in exams if e['status'] in ['completed', 'graded']]
    
    # 返回模板
    return render_template('exam_list.html', 
                          exams=exams,
                          subjects=sorted(list(subjects)))


@app.route('/exam/<int:exam_id>/detail')
@login_required
def exam_detail(exam_id):
    # 获取当前用户信息
    user_email = session.get('email')
    
    SID = session.get('student_id')
    # 添加额外的调试信息
    debug_query = """
        SELECT 
            sa.student_id,
            s.name,
            sa.question_id,
            q.question_type,
            sa.ai_score,
            sa.final_score,
            sa.review_status
        FROM student_answers sa
        JOIN students s ON sa.student_id = s.id
        JOIN questions q ON sa.question_id = q.id
        WHERE sa.session_id = ? AND sa.student_id = ?
        ORDER BY sa.student_id, sa.question_id
    """
    success_debug, debug_data = execute_query(debug_query, (exam_id, SID))
    
    if success_debug:
        app.logger.info(f"考试 {exam_id} 的答案记录总数: {len(debug_data)}")
        
        # 计算每个学生的总分
        student_scores = {}
        for row in debug_data:
            student_id = row['student_id']
            score = row['final_score'] if row['final_score'] is not None else row['ai_score']
            
            if student_id not in student_scores:
                student_scores[student_id] = 0
                
            if score is not None:
                student_scores[student_id] += score
                
        for student_id, score in student_scores.items():
            student_name = next((row['name'] for row in debug_data if row['student_id'] == student_id), "未知")
            app.logger.info(f"学生 {student_name} (ID: {student_id}) 的总分: {score}")
    
    # 获取考试基本信息
    exam_query = """
        SELECT id, name, subject, duration, exam_score as total_score, start_time, end_time
        FROM exam_sessions
        WHERE id = ?
    """
    success, exam_data = execute_query(exam_query, (exam_id,))
    
    if not success or not exam_data:
        return redirect(url_for('exam_list'))
        
    app.logger.info(f"获取到考试基本信息: {exam_data}")

    # 获取学生信息 - 使用用户名查询，而不是email
    student_query = """
        SELECT id FROM students WHERE student_id = ? OR name = ?
    """
    success, student_data = execute_query(student_query, (user_email, user_email))
    
    if not success or not student_data:
        # 如果找不到学生信息，创建临时会话使用的学生记录
        temp_student_id = f"temp_{user_email}_{int(time.time())}"
        insert_query = """
            INSERT INTO students (name, student_id)
            VALUES (?, ?)
        """
        success, student_id = execute_query(insert_query, (user_email, temp_student_id))
        
        if not success:
            return redirect(url_for('exam_list'))
    else:
        student_id = student_data[0]['id']
    
    app.logger.info(f"学生ID: {student_id}")
        
    # 获取考试的总分 - 首先尝试获取当前学生的所有答案总分
    score_query = """
        SELECT SUM(final_score) as total_score 
        FROM student_answers 
        WHERE session_id = ? AND student_id = ?
    """
    success, score_data = execute_query(score_query, (exam_id, student_id))
    
    app.logger.info(f"总分查询结果: {score_data}")
    
    # 如果没有找到分数，尝试不使用review_status条件，第二次查询
    if not success or not score_data or score_data[0]['total_score'] is None:
        app.logger.info("未找到评分数据，尝试获取所有答案...")
        score_query_alt = """
            SELECT 
                question_id,
                MAX(COALESCE(final_score, ai_score, 0)) as best_score
            FROM student_answers 
            WHERE session_id = ? AND student_id = ?
            GROUP BY question_id
        """
        success, question_scores = execute_query(score_query_alt, (exam_id, student_id))
        
        # 如果找到了单个题目的最佳分数，手动计算总和
        if success and question_scores:
            total_score = sum(q['best_score'] for q in question_scores)
            score_data = [{'total_score': total_score}]
            app.logger.info(f"通过单题分数计算得到总分: {total_score}")
        else:
            app.logger.warning("仍然未找到有效的分数数据")
    
    # 获取考试的统计数据：平均分、排名等
    user_id = get_current_user_id()
    stats_query = """
        SELECT 
            AVG(total_score) as avg_score,
            COUNT(*) as student_count
        FROM (
            SELECT 
                student_id, 
                SUM(COALESCE(final_score, ai_score, 0)) as total_score
            FROM student_answers
            WHERE session_id = ?
            GROUP BY student_id
        ) as student_scores
    """
    success, stats_data = execute_query(stats_query, (exam_id,))
    
    # 获取学生排名
    rank_query = """
        WITH student_scores AS (
            SELECT 
                student_id,
                SUM(COALESCE(final_score, ai_score, 0)) as total_score
            FROM student_answers 
            WHERE session_id = ?
            GROUP BY student_id
        ),
        class_scores AS (
            SELECT 
                student_id,
                total_score
            FROM student_scores
        )
        SELECT 
            cs1.student_id,
            COUNT(DISTINCT cs2.student_id) + 1 as rank
        FROM class_scores cs1
        LEFT JOIN class_scores cs2 ON cs2.total_score > cs1.total_score
        WHERE cs1.student_id = ?
        GROUP BY cs1.student_id
    """
    success_rank, rank_data = execute_query(rank_query, (exam_id,user_id))
    
    # 默认值
    avg_score = 0
    student_count = 0
    student_rank = 0
    
    if success and stats_data:
        avg_score = stats_data[0]['avg_score'] or 0
        print(f"平均分: {avg_score}")
        student_count = stats_data[0]['student_count'] or 0
        print(f"学生人数: {student_count}")
    
    if success_rank and rank_data:
        student_rank = rank_data[0]['rank']
        print(f"学生排名: {student_rank}")
    # 设置固定分数为70分
    user_id = get_current_user_id()
    rescore = """
        SELECT SUM(final_score) as total_score
        FROM student_answers
        WHERE session_id = ? AND student_id = ?;
    """
    success_rescore, rescore_data = execute_query(rescore, (exam_id, user_id))
    print(success_rescore, rescore_data)
    student_score = rescore_data[0]['total_score']
    print(f"学生分数: {student_score}")
    if success_rescore and rescore_data and rescore_data[0]['total_score'] is not None:
        student_score = rescore_data[0]['total_score']
        print(f"学生分数: {student_score}")
    # 计算得分率
    score_rate = (student_score / 100) * 100 if student_score is not None else 0
    # 根据得分率确定等级
    score_level = "优秀" if score_rate >= 90 else "良好" if score_rate >= 80 else "中等" if score_rate >= 70 else "及格" if score_rate >= 60 else "不及格"
    
    app.logger.info(f"报错固定分数为52分，得分率为{score_rate}%，等级为{score_level}")
    # 构建考试信息
    exam = {
        'id': exam_data[0]['id'],
        'name': exam_data[0]['name'],
        'duration': exam_data[0]['duration'],
        'total_score': exam_data[0]['total_score'],
        'start_time': datetime.strptime(exam_data[0]['start_time'], '%Y-%m-%d %H:%M:%S.%f') if '.' in exam_data[0]['start_time'] else datetime.strptime(exam_data[0]['start_time'], '%Y-%m-%d %H:%M:%S'),
        'score': student_score,
        'stats': {
            'rank': f"{student_rank}/{student_count}",
            'rank_percent': round(((student_count - student_rank) / student_count) * 100) if student_count > 0 else 0,
            'avg_score': round(avg_score, 1),
            'score_diff': round(student_score - avg_score, 1) if student_score is not None else 0,
            'score_rate': score_rate,
            'score_level': score_level
        },
        'questions': {
            'multiple_choice': [],
            'fill_blanks': [],
            'short_answer': [],
            'true_false': [],
            'programming': []
        }
    }
    
    # 获取试题和答题情况
    # 获取考试中的所有题目
    questions_query = """
        SELECT id, question_type, question_text, score, question_order
        FROM questions
        WHERE session_id = ?
        ORDER BY question_order
    """
    success, questions = execute_query(questions_query, (exam_id,))
    
    app.logger.info(f"找到题目数量: {len(questions) if success and questions else 0}")
    
    if success and questions:
        for q in questions:
            # 获取学生对该题的作答 - 不限制学生ID，获取所有学生的答案
            answer_query = """
                SELECT answer_text, ai_score, ai_feedback, scoring_details, final_score, manual_feedback, review_status, student_id
                FROM student_answers
                WHERE session_id = ? AND question_id = ?
                ORDER BY 
                    CASE review_status 
                        WHEN 'reviewed' THEN 1 
                        WHEN 'disputed' THEN 2
                        ELSE 3 
                    END,
                    COALESCE(final_score, 0) DESC,
                    COALESCE(ai_score, 0) DESC
                LIMIT 1
            """
            success, answer = execute_query(answer_query, (exam_id, q['id']))
            
            app.logger.info(f"题目ID {q['id']} 的答案数据: {answer}")
            
            question_data = {
                'id': q['id'],
                'question': q['question_text'],
                'score': q['score'],
                'user_answer': answer[0]['answer_text'] if success and answer else None,
                'final_score': answer[0]['final_score'] if success and answer and answer[0]['final_score'] is not None else (answer[0]['ai_score'] if success and answer and answer[0]['ai_score'] is not None else 0),
                'feedback': answer[0]['manual_feedback'] if success and answer and answer[0]['manual_feedback'] else (answer[0]['ai_feedback'] if success and answer else "暂无评分反馈")
            }
            
            app.logger.info(f"构建的题目数据: {question_data}")
            
            # 如果是学生自己的答案，记录下来
            student_answer = False
            if success and answer and answer[0]['student_id'] == student_id:
                student_answer = True
                app.logger.info(f"找到学生自己的答案！")
            
            question_data['is_your_answer'] = student_answer
            
            # 根据题型添加到对应类别
            if q['question_type'] == 'multiple_choice':
                # 获取选择题选项
                options_query = """
                    SELECT option_a, option_b, option_c, option_d, correct_answer
                    FROM multiple_choice_questions
                    WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                """
                success, options = execute_query(options_query, (q['id'],))
                
                if success and options:
                    question_data['options'] = [
                        options[0]['option_a'],
                        options[0]['option_b'],
                        options[0]['option_c'],
                        options[0]['option_d']
                    ]
                    question_data['correct_answer'] = options[0]['correct_answer']
                
                exam['questions']['multiple_choice'].append(question_data)
                
            elif q['question_type'] == 'fill_blank':
                # 获取填空题答案
                answer_query = """
                    SELECT correct_answer
                    FROM fill_blank_questions
                    WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                """
                success, answer_data = execute_query(answer_query, (q['id'],))
                
                if success and answer_data:
                    # 确保正确答案作为字符串存储，以便于模板中的比较
                    question_data['correct_answer'] = str(answer_data[0]['correct_answer'])
                
                # 如果用户答案存在，也将其转换为字符串
                if question_data.get('user_answer') is not None:
                    question_data['user_answer'] = str(question_data['user_answer'])
                
                exam['questions']['fill_blanks'].append(question_data)
                
            elif q['question_type'] == 'short_answer':
                exam['questions']['short_answer'].append(question_data)
            
            elif q['question_type'] == 'true_false':
                # 获取判断题的正确答案
                answer_query = """
                    SELECT correct_answer
                    FROM true_false_questions
                    WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                """
                success, answer_data = execute_query(answer_query, (q['id'],))
                
                if success and answer_data:
                    # 确保正确答案作为字符串存储，以便于模板中的比较
                    question_data['correct_answer'] = str(answer_data[0]['correct_answer'])
                
                # 如果用户答案存在，也将其转换为字符串
                if question_data.get('user_answer') is not None:
                    question_data['user_answer'] = str(question_data['user_answer'])
                
                exam['questions']['true_false'].append(question_data)
            
            elif q['question_type'] == 'programming':
                # 获取编程题的详细信息
                programming_query = """
                    SELECT sample_input, sample_output, time_limit, memory_limit, hints
                    FROM programming_questions
                    WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                """
                success, programming_data = execute_query(programming_query, (q['id'],))
                
                if success and programming_data:
                    question_data['sample_input'] = programming_data[0]['sample_input']
                    question_data['sample_output'] = programming_data[0]['sample_output']
                    question_data['time_limit'] = programming_data[0]['time_limit']
                    question_data['memory_limit'] = programming_data[0]['memory_limit']
                    question_data['hints'] = programming_data[0]['hints']
                
                exam['questions']['programming'].append(question_data)
    
    return render_template('exam_detail.html', exam=exam, chr=chr)


@app.route('/exam/<int:exam_id>/save', methods=['POST'])
@login_required
def save_exam_answers(exam_id):
    data = request.get_json()

    try:
        # 获取当前用户ID
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'status': 'error', 'message': '无法获取用户信息'}), 401
        
        # 保存答案到数据库
        save_answers_to_db(user_id, exam_id, data['answers'])
        return jsonify({'status': 'success', 'message': '保存成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/exam/<int:exam_id>/get_saved_answers', methods=['GET'])
@login_required
def get_saved_answers(exam_id):
    try:
        # 获取当前用户ID
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'status': 'error', 'message': '无法获取用户信息'}), 401
        
        # 获取学生ID
        student_query = """
            SELECT id FROM students WHERE user_id = ? OR student_id = ?
        """
        success, student_data = execute_query(student_query, (user_id, user_id))
        
        if not success or not student_data:
            return jsonify({'status': 'success', 'answers': {}})
            
        student_id = student_data[0]['id']
        
        # 获取保存的答案
        answers_query = """
            SELECT question_id, question_type, answer_text
            FROM student_answers
            WHERE student_id = ? AND session_id = ?
        """
        success, answers_data = execute_query(answers_query, (student_id, exam_id))
        
        if not success or not answers_data:
            return jsonify({'status': 'success', 'answers': {}})
            
        # 整理答案数据
        answers = {
            'multiple_choice': {},
            'fill_blank': {},
            'short_answer': {},
            'true_false': {},
            'programming': {}
        }
        
        for answer in answers_data:
            question_id = answer['question_id']
            question_type = answer['question_type']
            answer_text = answer['answer_text']
            
            if question_type == 'multiple_choice':
                answers['multiple_choice'][str(question_id)] = answer_text
            elif question_type == 'fill_blank':
                answers['fill_blank'][str(question_id)] = answer_text
            elif question_type == 'short_answer':
                answers['short_answer'][str(question_id)] = answer_text
            elif question_type == 'true_false':
                answers['true_false'][str(question_id)] = answer_text
            elif question_type == 'programming':
                answers['programming'][str(question_id)] = answer_text
                
        return jsonify({'status': 'success', 'answers': answers})
    except Exception as e:
        app.logger.error(f"获取保存的答案出错: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/exam/<int:exam_id>/submit', methods=['POST'])
@login_required
def submit_exam(exam_id):
    data = request.get_json()

    try:
        # 检查是否在考试时间内
        exam = get_exam_by_id(exam_id)
        now = datetime.now()
        if now > exam['end_time']:
            return jsonify({'status': 'error', 'message': '考试已结束，无法提交'}), 400

        # 获取当前用户ID
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'status': 'error', 'message': '无法获取用户信息'}), 401

        # 保存最终答案
        save_answers_to_db(user_id, exam_id, data['answers'])

        # 标记考试为已完成
        mark_exam_completed(user_id, exam_id)

        # 返回成功消息
        return jsonify({
            'status': 'success',
            'message': '试卷提交成功',
            'redirect': url_for('exam_list')
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


def get_exam_by_id(exam_id):
    try:
        app.logger.info(f"尝试获取考试ID: {exam_id}的信息")
        
        query = """
            SELECT id, name, subject, duration, exam_score as total_score, start_time, end_time, status
            FROM exam_sessions
            WHERE id = ?
        """
        
        # 获取数据库连接
        conn = get_db_connection()
        if not conn:
            app.logger.error("无法连接到数据库")
            return None
            
        # 尝试执行查询
        success, exam_data = execute_query(query, (exam_id,))
        
        if not success:
            app.logger.error(f"查询失败: {query}")
            return None
        
        if not exam_data:
            app.logger.warning(f"找不到考试ID: {exam_id}")
            return None
        
        app.logger.info(f"成功获取考试ID: {exam_id}的数据")
        
        exam = exam_data[0]
        # 转换日期时间格式
        try:
            if '.' in exam['start_time']:
                start_time = datetime.strptime(exam['start_time'], '%Y-%m-%d %H:%M:%S.%f')
            else:
                start_time = datetime.strptime(exam['start_time'], '%Y-%m-%d %H:%M:%S')
                
            if '.' in exam['end_time']:
                end_time = datetime.strptime(exam['end_time'], '%Y-%m-%d %H:%M:%S.%f')
            else:
                end_time = datetime.strptime(exam['end_time'], '%Y-%m-%d %H:%M:%S')
        except Exception as e:
            app.logger.error(f"日期转换错误: {str(e)}")
            return None
        
        return {
            'id': exam['id'],
            'name': exam['name'],
            'subject': exam['subject'],
            'duration': exam['duration'],
            'total_score': exam['total_score'],
            'start_time': start_time,
            'end_time': end_time,
            'status': exam['status']
        }
    except Exception as e:
        app.logger.error(f"获取考试数据时出错: {str(e)}", exc_info=True)
        return None


def get_exam_questions(exam_id):
    try:
        app.logger.info(f"获取考试ID: {exam_id}的题目")
        
        result = {
            'multiple_choice': [],
            'fill_blank': [],
            'short_answer': [],
            'true_false': [],
            'programming': []
        }
        
        # 获取数据库连接
        conn = get_db_connection()
        if not conn:
            app.logger.error("无法连接到数据库获取题目")
            return result
        
        # 获取考试中的所有题目
        questions_query = """
            SELECT id, question_type, question_text, score, question_order, source_question_id, source_table
            FROM questions
            WHERE session_id = ?
            ORDER BY question_order
        """
        app.logger.info(f"执行查询: {questions_query} 参数: {exam_id}")
        
        success, questions = execute_query(questions_query, (exam_id,))
        
        if not success:
            app.logger.error(f"获取考试题目失败: 数据库查询错误")
            return result
        
        if not questions:
            app.logger.warning(f"考试ID: {exam_id}没有题目")
            return result
        
        app.logger.info(f"成功获取到 {len(questions)} 道题目")
        
        # 手动处理题目，确保即使数据库中缺少数据也能正常显示
        for q in questions:
            try:
                app.logger.info(f"处理题目: {q['id']} 类型: {q['question_type']}")
                question_data = {
                    'id': q['id'],
                    'question': q['question_text'],
                    'score': q['score']
                }
                
                # 根据题型获取更多信息并添加到对应类别
                if q['question_type'] == 'multiple_choice':
                    try:
                        # 获取选择题选项
                        if q['source_question_id']:
                            options_query = """
                                SELECT option_a, option_b, option_c, option_d
                                FROM multiple_choice_questions
                                WHERE id = ?
                            """
                            app.logger.info(f"获取选择题选项: source_id={q['source_question_id']}")
                            success, options = execute_query(options_query, (q['source_question_id'],))
                            
                            if success and options and len(options) > 0:
                                app.logger.info(f"成功获取选择题选项")
                                question_data['options'] = [
                                    f"A. {options[0]['option_a']}",
                                    f"B. {options[0]['option_b']}",
                                    f"C. {options[0]['option_c']}",
                                    f"D. {options[0]['option_d']}"
                                ]
                            else:
                                # 如果找不到选项，使用默认选项
                                app.logger.warning(f"选择题ID: {q['id']}找不到选项数据，使用默认选项")
                                question_data['options'] = [
                                    "A. 选项A(未找到原始选项)",
                                    "B. 选项B(未找到原始选项)",
                                    "C. 选项C(未找到原始选项)",
                                    "D. 选项D(未找到原始选项)"
                                ]
                        else:
                            app.logger.warning(f"选择题ID: {q['id']}没有source_question_id，使用默认选项")
                            question_data['options'] = [
                                "A. 选项A(未找到原始选项)",
                                "B. 选项B(未找到原始选项)",
                                "C. 选项C(未找到原始选项)",
                                "D. 选项D(未找到原始选项)"
                            ]
                        
                        result['multiple_choice'].append(question_data)
                    except Exception as e:
                        app.logger.error(f"处理选择题时出错: {str(e)}", exc_info=True)
                        
                        # 即使出错，也添加一个基本的选择题对象
                        question_data['options'] = [
                            "A. 选项A(处理错误)",
                            "B. 选项B(处理错误)",
                            "C. 选项C(处理错误)",
                            "D. 选项D(处理错误)"
                        ]
                        result['multiple_choice'].append(question_data)
                        continue
                    
                elif q['question_type'] == 'fill_blank':
                    # 获取填空题答案
                    answer_query = """
                        SELECT correct_answer
                        FROM fill_blank_questions
                        WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                    """
                    success, answer_data = execute_query(answer_query, (q['id'],))
                    
                    if success and answer_data:
                        # 确保正确答案作为字符串存储
                        question_data['correct_answer'] = str(answer_data[0]['correct_answer'])
                    
                    # 如果用户答案存在，也将其转换为字符串
                    if question_data.get('user_answer') is not None:
                        question_data['user_answer'] = str(question_data['user_answer'])
                    
                    result['fill_blank'].append(question_data)
                    
                elif q['question_type'] == 'short_answer':
                    result['short_answer'].append(question_data)
                
                elif q['question_type'] == 'true_false':
                    # 获取判断题的正确答案
                    answer_query = """
                        SELECT correct_answer
                        FROM true_false_questions
                        WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                    """
                    success, answer_data = execute_query(answer_query, (q['id'],))
                    
                    if success and answer_data:
                        # 确保正确答案作为字符串存储，以便于模板中的比较
                        question_data['correct_answer'] = str(answer_data[0]['correct_answer'])
                    
                    # 如果用户答案存在，也将其转换为字符串
                    if question_data.get('user_answer') is not None:
                        question_data['user_answer'] = str(question_data['user_answer'])
                    
                    result['true_false'].append(question_data)
                
                elif q['question_type'] == 'programming':
                    # 获取编程题的详细信息
                    programming_query = """
                        SELECT sample_input, sample_output, time_limit, memory_limit, hints
                        FROM programming_questions
                        WHERE id = (SELECT source_question_id FROM questions WHERE id = ?)
                    """
                    success, programming_data = execute_query(programming_query, (q['id'],))
                    
                    if success and programming_data:
                        question_data['sample_input'] = programming_data[0]['sample_input']
                        question_data['sample_output'] = programming_data[0]['sample_output']
                        question_data['time_limit'] = programming_data[0]['time_limit']
                        question_data['memory_limit'] = programming_data[0]['memory_limit']
                        question_data['hints'] = programming_data[0]['hints']
                    
                    result['programming'].append(question_data)
            except Exception as e:
                app.logger.error(f"处理题目时发生错误: {str(e)}", exc_info=True)
                continue
        
        app.logger.info(f"分类结果: 选择题{len(result['multiple_choice'])}道, 填空题{len(result['fill_blank'])}道, 简答题{len(result['short_answer'])}道, 判断题{len(result['true_false'])}道, 编程题{len(result['programming'])}道")

        return result
    except Exception as e:
        app.logger.error(f"获取考试题目时发生错误: {str(e)}", exc_info=True)
        # 返回带有默认题目的结果
        default_result = {
            'multiple_choice': [],
            'fill_blank': [],
            'short_answer': [],
            'true_false': [],
            'programming': []
        }
        return default_result


def has_taken_exam(user_id, exam_id):
    """检查用户是否已参加过指定考试"""
    try:
        app.logger.info(f"检查用户 {user_id} 是否参加过考试 {exam_id}")
        
        # 首先获取学生ID
        student_query = """
            SELECT id FROM students WHERE user_id = ? OR student_id = ?
        """
        success, student_data = execute_query(student_query, (user_id, user_id))
        
        if not success:
            app.logger.error("查询学生信息失败")
            return False
            
        if not student_data:
            app.logger.info(f"找不到与用户 {user_id} 关联的学生记录")
            return False
            
        student_id = student_data[0]['id']
        app.logger.info(f"找到学生ID: {student_id}")
        
        # 检查是否有该考试的答题记录
        answer_query = """
            SELECT COUNT(*) as count
            FROM student_answers
            WHERE student_id = ? AND session_id = ?
        """
        success, answer_data = execute_query(answer_query, (student_id, exam_id))
        
        if not success:
            app.logger.error("查询答题记录失败")
            return False
            
        if not answer_data:
            app.logger.info("没有找到答题记录")
            return False
            
        count = answer_data[0]['count']
        has_taken = count > 0
        app.logger.info(f"答题记录数: {count}, 是否参加过: {has_taken}")
        
        return has_taken
    except Exception as e:
        app.logger.error(f"检查用户是否参加过考试时出错: {str(e)}", exc_info=True)
        # 如果出错，保守地返回False，使用户能参加考试
        return False


def save_answers_to_db(user_id, exam_id, answers):
    try:
        # 获取学生ID
        student_query = """
            SELECT id FROM students WHERE user_id = ? OR student_id = ?
        """
        success, student_data = execute_query(student_query, (user_id, user_id))
        
        app.logger.info(f"保存答案 - 用户ID: {user_id}, 考试ID: {exam_id}")
        app.logger.info(f"答案数据: {answers}")
        
        if not success or not student_data:
            # 如果找不到学生记录，创建一个新的
            new_student_query = """
                INSERT INTO students (user_id, student_id, name)
                VALUES (?, ?, ?)
            """
            email = session.get('email', f"user_{user_id}")
            success, result = execute_query(new_student_query, (user_id, email, email))
            if not success:
                app.logger.error("无法创建学生记录")
                raise Exception("无法创建学生记录")
            
            # 获取新创建的学生ID
            get_id_query = "SELECT last_insert_rowid() as id"
            success, id_result = execute_query(get_id_query)
            if not success or not id_result:
                app.logger.error("无法获取学生ID")
                raise Exception("无法获取学生ID")
            
            student_id = id_result[0]['id']
        else:
            student_id = student_data[0]['id']
        
        app.logger.info(f"学生ID: {student_id}")
        
        # 清除之前的答案（如果有的话）
        clear_query = """
            DELETE FROM student_answers
            WHERE student_id = ? AND session_id = ?
        """
        execute_query(clear_query, (student_id, exam_id))
        
        # 保存选择题答案
        if 'multiple_choice' in answers:
            app.logger.info(f"处理选择题答案: {answers['multiple_choice']}")
            for question_id, answer in answers['multiple_choice'].items():
                query = """
                    INSERT INTO student_answers
                    (student_id, session_id, question_id, answer_text, question_type)
                    VALUES (?, ?, ?, ?, 'multiple_choice')
                """
                app.logger.info(f"插入选择题答案: 题目ID:{question_id}, 答案:{answer}")
                success_insert, _ = execute_query(query, (student_id, exam_id, question_id, answer))
                app.logger.info(f"选择题答案插入结果: {success_insert}")
        
        # 保存填空题答案
        if 'fill_blank' in answers:
            app.logger.info(f"处理填空题答案: {answers['fill_blank']}")
            for question_id, answer in answers['fill_blank'].items():
                query = """
                    INSERT INTO student_answers
                    (student_id, session_id, question_id, answer_text, question_type)
                    VALUES (?, ?, ?, ?, 'fill_blank')
                """
                success_insert, _ = execute_query(query, (student_id, exam_id, question_id, answer))
                app.logger.info(f"填空题答案插入结果: {success_insert}")
        
        # 保存简答题答案
        if 'short_answer' in answers:
            app.logger.info(f"处理简答题答案: {answers['short_answer']}")
            for question_id, answer in answers['short_answer'].items():
                query = """
                    INSERT INTO student_answers
                    (student_id, session_id, question_id, answer_text, question_type)
                    VALUES (?, ?, ?, ?, 'short_answer')
                """
                success_insert, _ = execute_query(query, (student_id, exam_id, question_id, answer))
                app.logger.info(f"简答题答案插入结果: {success_insert}")
        
        # 保存判断题答案
        if 'true_false' in answers:
            app.logger.info(f"处理判断题答案: {answers['true_false']}")
            for question_id, answer in answers['true_false'].items():
                query = """
                    INSERT INTO student_answers
                    (student_id, session_id, question_id, answer_text, question_type)
                    VALUES (?, ?, ?, ?, 'true_false')
                """
                app.logger.info(f"插入判断题答案: 题目ID:{question_id}, 答案:{answer}")
                success_insert, _ = execute_query(query, (student_id, exam_id, question_id, answer))
                app.logger.info(f"判断题答案插入结果: {success_insert}")
        
        # 保存编程题答案
        if 'programming' in answers:
            app.logger.info(f"处理编程题答案: {answers['programming']}")
            for question_id, answer in answers['programming'].items():
                query = """
                    INSERT INTO student_answers
                    (student_id, session_id, question_id, answer_text, question_type)
                    VALUES (?, ?, ?, ?, 'programming')
                """
                success_insert, _ = execute_query(query, (student_id, exam_id, question_id, answer))
                app.logger.info(f"编程题答案插入结果: {success_insert}")
        
        app.logger.info(f"所有答案保存完成")
        return True
    except Exception as e:
        app.logger.error(f"保存答案出错: {str(e)}")
        raise


def mark_exam_completed(user_id, exam_id):
    """将考试标记为已完成状态，记录完成时间"""
    try:
        app.logger.info(f"尝试将用户ID:{user_id}的考试ID:{exam_id}标记为已完成")
        
        # 获取学生ID
        student_query = """
            SELECT id FROM students WHERE user_id = ? OR student_id = ?
        """
        success, student_data = execute_query(student_query, (user_id, user_id))
        
        if not success or not student_data:
            app.logger.warning(f"找不到学生记录，用户ID:{user_id}")
            return False
            
        student_id = student_data[0]['id']
        
        # 先检查是否已存在记录
        check_query = """
            SELECT id FROM student_exams
            WHERE student_id = ? AND session_id = ?
        """
        success, existing = execute_query(check_query, (student_id, exam_id))
        
        if success and existing:
            # 更新现有记录
            update_query = """
                UPDATE student_exams
                SET status = 'completed', completion_time = datetime('now')
                WHERE student_id = ? AND session_id = ?
            """
            success, _ = execute_query(update_query, (student_id, exam_id))
            if not success:
                app.logger.error(f"更新学生考试状态失败")
                # 尝试插入新记录
                insert_query = """
                    INSERT INTO student_exams (student_id, session_id, status, completion_time)
                    VALUES (?, ?, 'completed', datetime('now'))
                """
                success, _ = execute_query(insert_query, (student_id, exam_id))
                if not success:
                    app.logger.error(f"插入学生考试状态失败")
                    return False
        else:
            # 插入新记录
            insert_query = """
                INSERT INTO student_exams (student_id, session_id, status, completion_time)
                VALUES (?, ?, 'completed', datetime('now'))
            """
            success, _ = execute_query(insert_query, (student_id, exam_id))
            if not success:
                app.logger.error(f"插入学生考试状态失败")
                return False
                
        # 更新答案状态为已提交 - 删除了is_draft字段的引用，因为该字段不存在
        update_answers_query = """
            UPDATE student_answers
            SET review_status = 'pending'
            WHERE student_id = ? AND session_id = ?
        """
        execute_query(update_answers_query, (student_id, exam_id))
        
        # 更新考试状态为已完成
        update_exam_query = """
            UPDATE exam_sessions
            SET status = 'completed'
            WHERE id = ?
        """
        success, _ = execute_query(update_exam_query, (exam_id,))
        if not success:
            app.logger.warning(f"更新考试状态失败，考试ID:{exam_id}")
        
        app.logger.info(f"成功将用户ID:{user_id}的考试ID:{exam_id}标记为已完成")
        return True
    except Exception as e:
        app.logger.error(f"标记考试完成状态时出错: {str(e)}")
        return False

@app.route('/api/sessions', methods=['GET'])
@login_required
def get_exam_sessions():
    """获取所有考试场次"""
    search_term = request.args.get('search', '').strip()
    
    if search_term:
        query = """
            SELECT id, name, subject, start_time, end_time, status 
            FROM exam_sessions 
            WHERE name LIKE ?  
            ORDER BY start_time DESC
        """
        search_param = f'%{search_term}%'
        success, sessions = execute_query(query, (search_param,))
    else:
        # 如果没有搜索词，返回所有考试场次
        query = """
            SELECT id, name, subject, start_time, end_time, status 
            FROM exam_sessions 
            ORDER BY start_time DESC
        """
        success, sessions = execute_query(query)
    
    if not success:
        return jsonify({'error': '获取考试场次失败'}), 500
    
    # 确保返回的是列表
    sessions = sessions if sessions else []
    
    # 格式化日期时间
    formatted_sessions = []
    for session in sessions:
        try:
            # 处理日期时间字段
            start_time = datetime.strptime(session['start_time'], '%Y-%m-%d %H:%M:%S.%f') if '.' in session['start_time'] else datetime.strptime(session['start_time'], '%Y-%m-%d %H:%M:%S')
            end_time = datetime.strptime(session['end_time'], '%Y-%m-%d %H:%M:%S.%f') if '.' in session['end_time'] else datetime.strptime(session['end_time'], '%Y-%m-%d %H:%M:%S')
            
            formatted_session = {
                'id': session['id'],
                'name': session['name'],
                'subject': session['subject'],
                'start_time': start_time.strftime('%Y-%m-%d %H:%M'),
                'end_time': end_time.strftime('%Y-%m-%d %H:%M'),
                'status': session['status']
            }
            formatted_sessions.append(formatted_session)
        except Exception as e:
            print(f"格式化会话数据时出错: {str(e)}")
            continue
    
    return jsonify(formatted_sessions)

@app.route('/api/session/<int:session_id>', methods=['DELETE'])
@login_required
def delete_exam_session(session_id):
    """删除指定的考试场次"""
    try:
        # 直接执行删除操作
        query = "DELETE FROM exam_sessions WHERE id = ?"
        success, _ = execute_query(query, (session_id,))
        
        if not success:
            print(f"删除失败: session_id={session_id}")
            return jsonify({'error': '删除考试场次失败'}), 500
            
        return jsonify({'message': '删除成功'})
    except Exception as e:
        print(f"删除考试场次错误: {str(e)}")
        return jsonify({'error': '删除考试场次失败'}), 500

@app.route('/api/session', methods=['POST'])
@login_required
def create_exam_session():
    try:
        # 获取表单数据
        name = request.form.get('name')
        subject = request.form.get('subject')
        exam_time = request.form.get('examTime')
        duration = request.form.get('duration')
        exam_score = request.form.get('totalScore')
        
        # 验证必填字段
        if not name or not subject or not exam_time or not duration:
            return jsonify({'error': '缺少必要的考试信息'}), 400
        
        # 确保分值是有效的数字，如果提供了分值
        if exam_score:
            try:
                exam_score = float(exam_score)
            except ValueError:
                return jsonify({'error': '考试分值必须是有效的数字'}), 400
        else:
            # 如果未提供分值，设置为默认值100
            exam_score = 100.0
        
        # 计算考试结束时间
        start_time = datetime.strptime(exam_time, '%Y-%m-%dT%H:%M')
        end_time = start_time + timedelta(minutes=int(duration))
        
        # 获取当前用户ID
        user_id = get_current_user_id()
        
        # 创建考试记录并获取ID
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': '数据库连接失败'}), 500
        
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO exam_sessions 
                (name, subject, start_time, end_time, duration, status, created_by, created_at, exam_file_path, exam_score) 
                VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'), ?, ?)
            """, (name, subject, start_time, end_time, duration, 'pending', user_id, '', exam_score))
            conn.commit()
            session_id = cursor.lastrowid
        except Exception as e:
            if conn:
                conn.rollback()
            print(f"创建考试记录失败: {str(e)}")
            return jsonify({'error': f'创建考试失败: {str(e)}'}), 500
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
        
        # 检查是否有上传文件
        if 'examFile' in request.files and request.files['examFile'].filename != '':
            file = request.files['examFile']
            
            # 验证文件类型
            if not allowed_file(file.filename, ['pdf']):
                return jsonify({'error': '不支持的文件类型，仅支持PDF格式'}), 400
            
            # 确保上传目录存在
            os.makedirs(EXAM_UPLOAD_FOLDER, exist_ok=True)
            
            # 保存文件
            safe_filename = secure_filename(file.filename)
            filename = f"{session_id}_{int(time.time())}_{safe_filename}"
            file_path = os.path.join(EXAM_UPLOAD_FOLDER, filename)
            
            try:
                file.save(file_path)
                
                # 更新考试记录，添加文件路径
                conn = get_db_connection()
                if not conn:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    return jsonify({'error': '数据库连接失败'}), 500
                
                cursor = conn.cursor()
                try:
                    cursor.execute("UPDATE exam_sessions SET exam_file_path = ? WHERE id = ?", (filename, session_id))
                    conn.commit()
                except Exception as e:
                    if conn:
                        conn.rollback()
                    # 如果更新失败，删除已上传的文件
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    print(f"更新考试文件信息失败: {str(e)}")
                    return jsonify({'error': f'更新考试文件信息失败: {str(e)}'}), 500
                finally:
                    if cursor:
                        cursor.close()
                    if conn:
                        conn.close()
                # 处理试卷PDF
                print("开始处理试卷PDF...")
                exam_score = process_pdf(file_path, out_input="oocr_results.txt")
                from split_and_ocr.read.questionsplit import oreadexit
                oreadexit()
                # ai识别题目，插入数据库
                # 调用aiexam.py中的run_ocr函数处理试卷
                from split_and_ocr.read.aiexam import AIExam
                AIExam.run_ocr(subject, user_id)
                # if exam_score == 0:
                #     if os.path.exists(file_path):
                #         os.remove(file_path)
                #     return jsonify({'error': '试卷PDF处理失败'}), 500
            except Exception as e:
                print(f"保存文件失败: {str(e)}")
                return jsonify({'error': f'保存试卷文件失败: {str(e)}'}), 500
        
        return jsonify({
            'success': True,
            'id': session_id,
            'message': '考试创建成功'
        })
        
    except Exception as e:
        print(f"创建考试失败: {str(e)}")
        return jsonify({'error': f'创建考试失败: {str(e)}'}), 500

@app.route('/api/review/students/<int:session_id>', methods=['GET'])
@login_required
def get_session_students(session_id):
    """获取指定考试场次的学生列表"""
    search = request.args.get('search', '').strip()
    
    query = """
        SELECT DISTINCT
            s.id,
            s.name,
            s.student_id,
            SUM(sa.ai_score) as initial_score,
            MIN(sa.review_status) as review_status
        FROM students s
        JOIN student_answers sa ON s.id = sa.student_id
        WHERE sa.session_id = ?
        GROUP BY s.id, s.name, s.student_id
    """
    
    params = [session_id]
    
    if search:
        query = """
            SELECT DISTINCT
                s.id,
                s.name,
                s.student_id,
                SUM(sa.ai_score) as initial_score,
                MIN(sa.review_status) as review_status
            FROM students s
            JOIN student_answers sa ON s.id = sa.student_id
            WHERE sa.session_id = ?
            AND (s.name LIKE ? OR s.student_id LIKE ?)
            GROUP BY s.id, s.name, s.student_id
        """
        params.extend([f'%{search}%', f'%{search}%'])
    
    print(f"Executing query: {query}")
    print(f"With params: {params}")
    
    success, students = execute_query(query, params)
    
    if not success:
        print("Query failed")
        return jsonify({'error': '获取学生列表失败'}), 500
    
    if success:
        student_score = students[0]['initial_score']
        print(student_score)
        if student_score is None:
            stats_query = """
                SELECT sum(final_score) FROM student_answers WHERE session_id = ? AND student_id = ?;
            """
            success, stats_data = execute_query(stats_query, (session_id, students[0]['id']))
            student_score = stats_data[0]['sum(final_score)']
            if student_score is None:
                student_score = 0
    print(f"Query results: {students}")
    
    return jsonify([{
        'id': student['id'],
        'name': student['name'],
        'student_id': student['student_id'],
        'initial_score': student['initial_score'] or student_score,
        'review_status': student['review_status'] or 'pending'
    } for student in (students or [])])

@app.route('/api/review/answers/<int:session_id>/<int:student_id>', methods=['GET'])
@login_required
def get_student_answers(session_id, student_id):
    """获取学生的答题详情"""
    print(f"Getting answers for session {session_id}, student {student_id}")  # 调试日志
    
    query = """
        SELECT 
            q.id as question_id,
            q.question_text,
            q.score as total_score,
            q.question_order,
            sa.answer_text,
            sa.ai_score,
            sa.ai_feedback,
            sa.review_status
        FROM questions q
        JOIN student_answers sa ON q.id = sa.question_id
        WHERE sa.session_id = ? AND sa.student_id = ?
        ORDER BY q.question_order
    """
    
    success, answers = execute_query(query, (session_id, student_id))
    print(f"Query results: {answers}")  # 调试日志
    
    if not success:
        print("Query failed")  # 调试日志
        return jsonify({'error': '获取答题详情失败'}), 500
        
    result = [{
        'question_id': answer['question_id'],
        'question_text': answer['question_text'],
        'total_score': answer['total_score'],
        'question_order': answer['question_order'],
        'answer_text': answer['answer_text'],
        'ai_score': answer['ai_score'],
        'ai_feedback': answer['ai_feedback'],
        'review_status': answer['review_status'] or 'pending'
    } for answer in (answers or [])]
    
    print(f"Returning: {result}")  # 调试日志
    return jsonify(result)

@app.route('/api/review/submit', methods=['POST'])
@login_required
def submit_review():
    """提交人工复核结果"""
    data = request.get_json()
    session_id = data.get('session_id')
    student_id = data.get('student_id')
    question_id = data.get('question_id')
    final_score = data.get('final_score')
    
    if not all([session_id, student_id, question_id, final_score is not None]):
        return jsonify({'error': '缺少必要参数'}), 400
        
    update_query = """
        UPDATE student_answers
        SET final_score = ?,
            review_status = 'reviewed',
            reviewed_at = CURRENT_TIMESTAMP
        WHERE session_id = ? AND student_id = ? AND question_id = ?
    """
    
    success, _ = execute_query(update_query, (final_score, session_id, student_id, question_id))
    
    if not success:
        return jsonify({'error': '保存复核结果失败'}), 500
        
    return jsonify({'message': '复核结果已保存'})

@app.route('/api/review/batch', methods=['POST'])
@login_required
def batch_review():
    """批量复核"""
    data = request.get_json()
    session_id = data.get('session_id')
    
    if not session_id:
        return jsonify({'error': '缺少考试场次ID'}), 400
        
    # 获取所有未复核的答卷
    query = """
        UPDATE student_answers
        SET final_score = ai_score,
            review_status = 'reviewed',
            reviewed_at = CURRENT_TIMESTAMP
        WHERE session_id = ? AND (review_status = 'pending' or review_status = 'reviewed')
    """
    
    success, _ = execute_query(query, (session_id,))
    
    if not success:
        return jsonify({'error': '批量复核失败'}), 500
        
    return jsonify({'message': '批量复核完成'})


@app.route('/api/analysis/basic-stats/<int:session_id>')
@login_required
def get_basic_stats(session_id):
    """获取考试的基本统计信息"""
    
    # 从score_statistics表获取基本数据
    stats_query = """
        SELECT ss.average_score, ss.pass_rate, ss.highest_score, ss.highest_score_student_id
        FROM score_statistics ss
        WHERE ss.session_id = ?
    """
    success, stats = execute_query(stats_query, (session_id,))
    
    if not success or not stats:
        return jsonify({
            'average_score': 0,
            'pass_rate': 0,
            'highest_score': 0,
            'total_students': 0,
            'ungraded_count': 0,
            'highest_score_student': {
                'name': '',
                'student_id': ''
            }
        })
    
    # 获取总学生数和未评分数据
    student_count_query = """
        SELECT 
            COUNT(DISTINCT student_id) as total_students,
            COUNT(DISTINCT CASE WHEN status != 'graded' THEN student_id END) as ungraded_count
        FROM student_exams
        WHERE session_id = ?
    """
    success_count, count_data = execute_query(student_count_query, (session_id,))
    
    # 获取最高分学生信息
    student_info_query = """
        SELECT name, student_id as student_number
        FROM students
        WHERE id = ?
    """
    highest_student_id = stats[0]['highest_score_student_id']
    success_student, student_info = execute_query(student_info_query, (highest_student_id,))
    
    # 处理数据
    total_students = count_data[0]['total_students'] if success_count and count_data else 0
    ungraded_count = count_data[0]['ungraded_count'] if success_count and count_data else 0
    
    # 转换pass_rate从小数到百分比
    pass_rate = int(stats[0]['pass_rate'] * 100)
    
    # 获取最高分学生信息
    highest_student = {
        'name': student_info[0]['name'] if success_student and student_info else '',
        'student_id': student_info[0]['student_number'] if success_student and student_info else ''
    }
    
    return jsonify({
        'average_score': stats[0]['average_score'],
        'pass_rate': pass_rate,
        'highest_score': stats[0]['highest_score'],
        'total_students': total_students,
        'ungraded_count': ungraded_count,
        'highest_score_student': highest_student
    })

# 获取分数分布数据
@app.route('/api/analysis/score-distribution/<int:session_id>')
@login_required
def get_score_distribution(session_id):
    # 直接从数据库表查询，而不是重新计算
    query = """
    SELECT score_range, student_count as count, 
           IFNULL(ungraded_count, 0) as ungraded
    FROM score_distribution 
    WHERE session_id = ?
    ORDER BY CASE
        WHEN score_range = '90-100' THEN 1
        WHEN score_range = '80-89' THEN 2
        WHEN score_range = '70-79' THEN 3
        WHEN score_range = '60-69' THEN 4
        WHEN score_range = '0-59' THEN 5
    END
    """
    success, results = execute_query(query, (session_id,))
    
    if not success or not results:
        # 如果没有数据，返回默认的空分布
        default_distribution = [
            {'range': '90-100', 'count': 0, 'ungraded': 0},
            {'range': '80-89', 'count': 0, 'ungraded': 0},
            {'range': '70-79', 'count': 0, 'ungraded': 0},
            {'range': '60-69', 'count': 0, 'ungraded': 0},
            {'range': '0-59', 'count': 0, 'ungraded': 0}
        ]
        return jsonify(default_distribution)
    
    # 转换字段名以匹配前端期望
    formatted_results = []
    for row in results:
        formatted_results.append({
            'range': row['score_range'],
            'count': row['count'],
            'ungraded': row['ungraded']
        })
    
    return jsonify(formatted_results)
# 获取题目分析数据
@app.route('/api/analysis/question-analysis/<int:session_id>')
@login_required
def get_question_analysis(session_id):
    """获取题目分析数据"""
    query = """
        SELECT 
            q.question_text,
            qa.average_score_rate,
            qa.difficulty_coefficient,
            qa.discrimination_degree
        FROM question_analysis qa
        JOIN questions q ON qa.question_id = q.id
        WHERE qa.session_id = ?
        ORDER BY q.question_order
    """
    success, analysis = execute_query(query, (session_id,))
    if not success or not analysis:
        return jsonify({'error': '获取题目分析失败'}), 500
    
    return jsonify([{
        'question': a['question_text'],
        'average_score_rate': a['average_score_rate'],
        'difficulty': a['difficulty_coefficient'],
        'discrimination': a['discrimination_degree']
    } for a in (analysis or [])])

# 获取学生历史成绩
@app.route('/api/analysis/student-history/<string:search_term>')
@login_required
def get_student_history(search_term):
    """获取学生的历史成绩记录"""
    try:
        print(f"搜索学生: {search_term}")  # 调试日志
        
        # 检查数据库连接
        conn = get_db_connection()
        if not conn:
            print("数据库连接失败")
            return jsonify({'error': '数据库连接失败'}), 500
            
        # 先查找学生
        student_query = """
            SELECT id, name, student_id
            FROM students
            WHERE name LIKE ? OR student_id LIKE ?
        """
        search_pattern = f"%{search_term}%"
        print(f"执行查询: {student_query} 参数: {search_pattern}")  # 调试日志
        
        success, students = execute_query(student_query, (search_pattern, search_pattern))
        print(f"查询结果: success={success}, students={students}")  # 调试日志
        
        if not success:
            return jsonify({'error': '数据库查询失败'}), 500
            
        if not students:
            return jsonify({
                'student': None,
                'history': []
            })
            
        # 检查学生数据
        student = students[0]
        print(f"找到学生: {student}")  # 调试日志
        
        # 构造返回数据
        result = {
            'student': {
                'id': student['id'],
                'name': student['name'],
                'student_id': student['student_id']
            },
            'history': []
        }
        
        # 如果找到学生，再查询历史成绩
        if result['student']:
            history_query = """
                WITH class_stats AS (
                    SELECT 
                        session_id,
                        COUNT(DISTINCT student_id) as total_students,
                        AVG(score) as avg_score
                    FROM student_exams
                    WHERE status = 'graded'
                    GROUP BY session_id
                ),
                rankings AS (
                    SELECT 
                        se1.session_id,
                        se1.student_id,
                        (SELECT COUNT(*) + 1 FROM student_exams se2 
                        WHERE se2.session_id = se1.session_id 
                        AND se2.score > se1.score
                        AND se2.status = 'graded') as rank
                    FROM student_exams se1
                    WHERE se1.student_id = ? AND se1.status = 'graded'
                )
                SELECT 
                    se.session_id,
                    es.name as exam_name,
                    es.start_time as exam_time,
                    se.score,
                    es.exam_score as total_possible_score,
                    r.rank as class_rank,
                    cs.total_students,
                    cs.avg_score
                FROM student_exams se
                JOIN exam_sessions es ON se.session_id = es.id
                LEFT JOIN rankings r ON se.session_id = r.session_id AND se.student_id = r.student_id
                JOIN class_stats cs ON se.session_id = cs.session_id
                WHERE se.student_id = ? AND se.status = 'graded'
                ORDER BY es.start_time DESC
            """
            print(f"执行历史查询: student_id={result['student']['id']}")  # 调试日志
            
            success, history = execute_query(history_query, (result['student']['id'], result['student']['id']))
            print(f"历史查询结果: success={success}, history={history}")  # 调试日志
            
            if success and history:
                result['history'] = [{
                    'exam_name': h['exam_name'],
                    'exam_time': h['exam_time'],
                    'score': h['score'],
                    'total_score': h['total_possible_score'],
                    'score_percentage': round(h['score'] / h['total_possible_score'] * 100, 1) if h['total_possible_score'] else 0,
                    'class_rank': f"{h['class_rank']}/{h['total_students']}",
                    'vs_average': round(float(h['score']) - float(h['avg_score']), 1) if h['avg_score'] else 0
                } for h in history]
        
        print("返回数据:", result)  # 调试日志
        return jsonify(result)
        
    except Exception as e:
        print(f"获取学生历史成绩失败: {e}")  # 调试日志
        import traceback
        traceback.print_exc()  # 打印完整的错误堆栈
        return jsonify({'error': str(e)}), 500

# 添加中文字体支持
pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))

# 添加导出相关的路由
@app.route('/api/analysis/export', methods=['POST'])
@login_required
def export_analysis():
    """导出分析报告"""
    try:
        data = request.get_json()
        print("收到导出请求:", data)  # 调试日志
        
        export_type = data.get('type')
        report_type = data.get('report_type')
        format_type = data.get('format')
        session_id = data.get('session_id')
        student_id = data.get('student_id')
        
        if not all([export_type, report_type, format_type]):
            return jsonify({'error': '缺少必要参数'}), 400
            
        if export_type == 'exam' and not session_id:
            return jsonify({'error': '缺少考试场次ID'}), 400
            
        if export_type == 'student' and not student_id:
            return jsonify({'error': '缺少学生ID'}), 400
            
        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{export_type}_{report_type}_{timestamp}"
        
        try:
            # 根据不同的导出类型和格式调用相应的处理函数
            if export_type == 'exam':
                if report_type == 'analysis':
                    file_path = generate_exam_analysis_report(session_id, format_type, filename)
                elif report_type == 'questions':
                    file_path = generate_question_analysis_report(session_id, format_type, filename)
                else:  # scores
                    file_path = generate_score_summary_report(session_id, format_type, filename)
            else:  # student
                if report_type == 'personal':
                    file_path = generate_personal_report(student_id, format_type, filename)
                elif report_type == 'ability':
                    file_path = generate_ability_report(student_id, format_type, filename)
                else:  # improvement
                    file_path = generate_improvement_report(student_id, format_type, filename)
                    
            print(f"文件已生成: {file_path}")  # 调试日志
            
            # 记录导出操作
            user_id = get_current_user_id()
            if user_id:  # 只在获取到用户ID时记录日志
                log_export(user_id, export_type, report_type, format_type, session_id, student_id)
            
            # 返回文件下载链接
            return jsonify({
                'success': True,
                'file_url': url_for('download_report', filename=os.path.basename(file_path))
            })
            
        except Exception as e:
            print(f"生成报告失败: {e}")  # 调试日志
            return jsonify({'error': f'生成报告失败: {str(e)}'}), 500
            
    except Exception as e:
        print(f"导出请求处理失败: {e}")  # 调试日志
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

def log_export(user_id, export_type, report_type, format_type, session_id=None, student_id=None):
    """记录导出操作"""
    query = """
        INSERT INTO export_logs 
        (user_id, export_type, report_type, format_type, session_id, student_id)
        VALUES (?, ?, ?, ?, ?, ?)
    """
    execute_query(query, (user_id, export_type, report_type, format_type, session_id, student_id))

def get_current_user_id():
    """获取当前登录用户的ID"""
    try:
        if 'email' not in session:
            app.logger.warning("会话中没有email信息")
            return None
            
        email = session['email']
        app.logger.info(f"尝试获取用户 {email} 的ID")
        
        query = "SELECT id FROM users WHERE email = ?"
        success, result = execute_query(query, (email,))
        
        if not success:
            app.logger.error(f"查询用户ID失败: {email}")
            return None
            
        if not result or len(result) == 0:
            app.logger.warning(f"找不到用户: {email}")
            return None
            
        user_id = result[0]['id']
        app.logger.info(f"找到用户ID: {user_id}")
        return user_id
    except Exception as e:
        app.logger.error(f"获取当前用户ID时出错: {str(e)}", exc_info=True)
        return None

@app.route('/api/analysis/download/<filename>')
@login_required
def download_report(filename):
    """下载生成的报告"""
    try:
        return send_from_directory('exports', filename, as_attachment=True)
    except Exception as e:
        return jsonify({'error': '文件下载失败'}), 404

# 报告生成函数
def generate_pdf_exam_analysis(session_info, stats, filename):
    """生成 PDF 格式的考试分析报告"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph(f"考试分析报告 - {session_info['name']}", title_style))
    
    # 添加考试基本信息
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    info_text = f"科目: {session_info['subject']}<br/>"
    info_text += f"时间: {session_info['start_time']} 至 {session_info['end_time']}<br/>"
    info_text += f"时长: {calculate_duration(session_info['start_time'], session_info['end_time'])}<br/>"
    story.append(Paragraph(info_text, info_style))
    
    # 添加统计数据
    stats_style = ParagraphStyle(
        'CustomStats',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    stats_text = f"平均分: {stats['average_score']}<br/>"
    stats_text += f"及格率: {stats['pass_rate']}%<br/>"
    stats_text += f"最高分: {stats['highest_score']}<br/>"
    story.append(Paragraph(stats_text, stats_style))
    
    # 添加分数分布统计表
    data = [['分数段', '人数']]
    for score_range in ['90-100', '80-89', '70-79', '60-69', '0-59']:
        count = stats.get(score_range, 0)
        data.append([score_range, str(count)])
    
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    table = Table(data)
    table.setStyle(table_style)
    story.append(table)
    
    # 添加分析结论
    conclusion_style = ParagraphStyle(
        'CustomConclusion',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    conclusion_text = generate_analysis_conclusion(stats['average_score'], stats['pass_rate'])
    story.append(Paragraph(conclusion_text, conclusion_style))
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_excel_exam_analysis(session_info, stats, filename):
    """生成 Excel 格式的考试分析报告"""
    file_path = os.path.join('exports', f'{filename}.xlsx')
    wb = Workbook()
    ws = wb.active
    
    # 添加标题
    ws.append(['考试分析报告'])
    ws.merge_cells('A1:D1')
    ws['A1'].font = Font(name='SimSun', size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # 添加考试基本信息
    ws.append(['考试名称', session_info['name']])
    ws.append(['科目', session_info['subject']])
    ws.append(['时间', f"{session_info['start_time']} 至 {session_info['end_time']}"])
    ws.append(['时长', calculate_duration(session_info['start_time'], session_info['end_time'])])
    
    # 添加统计数据
    ws.append(['平均分', stats['average_score']])
    ws.append(['及格率', f"{stats['pass_rate']}%"])
    ws.append(['最高分', stats['highest_score']])
    
    # 添加分数分布统计表
    ws.append(['分数分布统计'])
    ws.append(['分数段', '人数'])
    for score_range in ['90-100', '80-89', '70-79', '60-69', '0-59']:
        count = stats.get(score_range, 0)
        ws.append([score_range, count])
    
    # 添加分析结论
    ws.append(['分析结论'])
    ws.append([generate_analysis_conclusion(stats['average_score'], stats['pass_rate'])])
    
    # 设置单元格格式
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
        for cell in row:
            cell.font = Font(name='SimSun')
            cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # 调整列宽
    for column_letter in ['A', 'B', 'C', 'D']:
        ws.column_dimensions[column_letter].width = 15
    
    # 保存 Excel 文件
    wb.save(file_path)
    return file_path

def generate_word_exam_analysis(session_info, stats, filename):
    """生成 Word 格式的考试分析报告"""
    file_path = os.path.join('exports', f'{filename}.docx')
    doc = Document()
    
    # 添加标题
    doc.add_heading('考试分析报告', level=1)
    
    # 添加考试基本信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = '考试名称'
    info_table.cell(0, 1).text = session_info['name']
    info_table.cell(1, 0).text = '科目'
    info_table.cell(1, 1).text = session_info['subject']
    info_table.cell(2, 0).text = '时间'
    info_table.cell(2, 1).text = f"{session_info['start_time']} 至 {session_info['end_time']}"
    info_table.cell(3, 0).text = '时长'
    info_table.cell(3, 1).text = calculate_duration(session_info['start_time'], session_info['end_time'])
    
    # 添加统计数据
    stats_table = doc.add_table(rows=3, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.cell(0, 0).text = '平均分'
    stats_table.cell(0, 1).text = str(stats['average_score'])
    stats_table.cell(1, 0).text = '及格率'
    stats_table.cell(1, 1).text = f"{stats['pass_rate']}%"
    stats_table.cell(2, 0).text = '最高分'
    stats_table.cell(2, 1).text = str(stats['highest_score'])
    
    # 添加分数分布统计表
    doc.add_heading('分数分布统计', level=2)
    score_dist_table = doc.add_table(rows=6, cols=2)
    score_dist_table.style = 'Table Grid'
    score_dist_table.cell(0, 0).text = '分数段'
    score_dist_table.cell(0, 1).text = '人数'
    for i, score_range in enumerate(['90-100', '80-89', '70-79', '60-69', '0-59'], start=1):
        count = stats.get(score_range, 0)
        score_dist_table.cell(i, 0).text = score_range
        score_dist_table.cell(i, 1).text = str(count)
    
    # 添加分析结论
    doc.add_heading('分析结论', level=2)
    conclusion = doc.add_paragraph()
    conclusion.add_run(generate_analysis_conclusion(stats['average_score'], stats['pass_rate']))
    
    # 保存 Word 文件
    doc.save(file_path)
    return file_path

def generate_analysis_conclusion(average_score, pass_rate):
    """生成分析结论"""
    conclusion = []
    
    # 分析平均分
    if average_score >= 85:
        conclusion.append("本次考试整体成绩优秀，学生掌握程度很好。")
    elif average_score >= 75:
        conclusion.append("本次考试整体成绩良好，大部分学生掌握了主要知识点。")
    elif average_score >= 60:
        conclusion.append("本次考试整体成绩一般，学生的知识掌握还有提升空间。")
    else:
        conclusion.append("本次考试整体成绩不理想，需要加强基础知识的巩固。")
    
    # 分析及格率
    if pass_rate >= 90:
        conclusion.append(f"及格率达到{pass_rate}%，说明绝大多数学生达到了基本要求。")
    elif pass_rate >= 80:
        conclusion.append(f"及格率为{pass_rate}%，大部分学生达到了基本要求。")
    elif pass_rate >= 60:
        conclusion.append(f"及格率为{pass_rate}%，仍有部分学生未能达到基本要求，需要进行针对性辅导。")
    else:
        conclusion.append(f"及格率仅为{pass_rate}%，建议进行全面的知识点复习和强化训练。")
    
    # 添加建议
    if average_score < 60 or pass_rate < 60:
        conclusion.append("建议：\n1. 组织专题复习，针对性讲解重点难点\n2. 增加练习量，强化基础知识\n3. 开展一对一辅导，帮助学习困难学生")
    elif average_score < 75 or pass_rate < 80:
        conclusion.append("建议：\n1. 查漏补缺，巩固薄弱环节\n2. 适当增加练习的深度和难度\n3. 鼓励学生主动提问，及时解决疑难问题")
    else:
        conclusion.append("建议：\n1. 保持良好的学习状态\n2. 进一步提高解题能力和知识应用能力\n3. 可以尝试更具挑战性的题目")
    
    return "\n".join(conclusion)

def generate_exam_analysis_report(session_id, format_type, filename):
    """生成考试整体分析报告"""
    # 获取考试基本信息
    session_query = """
        SELECT 
            id,
            name, 
            subject, 
            start_time, 
            end_time,
            status
        FROM exam_sessions
        WHERE id = ?
    """
    success, session_info = execute_query(session_query, (session_id,))
    if not success or not session_info:
        raise Exception("获取考试信息失败")
    
    # 将 sqlite3.Row 对象转换为字典
    session_info = dict(session_info[0])
    
    # 获取统计数据
    stats_query = """
        WITH student_total_scores AS (
            SELECT 
                student_id,
                SUM(COALESCE(final_score, ai_score)) as total_score
            FROM student_answers
            WHERE session_id = ?
            GROUP BY student_id
        )
        SELECT 
            ROUND(AVG(total_score), 1) as average_score,
            ROUND((COUNT(CASE WHEN total_score >= 60 THEN 1 END) * 100.0 / COUNT(*)), 1) as pass_rate,
            MAX(total_score) as highest_score
        FROM student_total_scores
    """
    success, stats = execute_query(stats_query, (session_id,))
    if not success:
        raise Exception("获取统计数据失败")
    
    # 如果没有统计数据，提供默认值
    stats = dict(stats[0]) if stats else {
        'average_score': 0,
        'pass_rate': 0,
        'highest_score': 0
    }
    
    # 获取分数分布数据
    distribution_query = """
        WITH student_total_scores AS (
            SELECT 
                student_id,
                SUM(COALESCE(final_score, ai_score)) as total_score
            FROM student_answers
            WHERE session_id = ?
            GROUP BY student_id
        )
        SELECT 
            CASE 
                WHEN total_score >= 90 THEN '90-100'
                WHEN total_score >= 80 THEN '80-89'
                WHEN total_score >= 70 THEN '70-79'
                WHEN total_score >= 60 THEN '60-69'
                ELSE '0-59'
            END as score_range,
            COUNT(*) as student_count
        FROM student_total_scores
        GROUP BY 
            CASE 
                WHEN total_score >= 90 THEN '90-100'
                WHEN total_score >= 80 THEN '80-89'
                WHEN total_score >= 70 THEN '70-79'
                WHEN total_score >= 60 THEN '60-69'
                ELSE '0-59'
            END
        ORDER BY score_range DESC
    """
    success, distribution = execute_query(distribution_query, (session_id,))
    
    # 将分布数据转换为字典格式
    distribution_dict = {}
    if success and distribution:
        for row in distribution:
            distribution_dict[row['score_range']] = row['student_count']
    
    # 将分布数据添加到统计数据中
    stats.update(distribution_dict)
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_exam_analysis(session_info, stats, filename)
    elif format_type == 'excel':
        return generate_excel_exam_analysis(session_info, stats, filename)
    else:  # word
        return generate_word_exam_analysis(session_info, stats, filename)

def generate_question_analysis_report(session_id, format_type, filename):
    """生成题目分析报告"""
    # 获取题目分析数据
    query = """
        SELECT 
            q.question_text,
            qa.average_score_rate,
            qa.difficulty_coefficient,
            qa.discrimination_degree
        FROM question_analysis qa
        JOIN questions q ON qa.question_id = q.id
        WHERE qa.session_id = ?
        ORDER BY q.question_order
    """
    success, analysis = execute_query(query, (session_id,))
    if not success or not analysis:
        raise Exception("获取题目分析数据失败")
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_question_analysis(analysis, filename)
    elif format_type == 'excel':
        return generate_excel_question_analysis(analysis, filename)
    else:  # word
        return generate_word_question_analysis(analysis, filename)

def generate_pdf_question_analysis(analysis, filename):
    """生成 PDF 格式的题目分析报告"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph("题目分析报告", title_style))
    
    # 添加说明文字
    content_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=12
    )
    
    # 创建表格数据
    table_data = [['题目', '平均得分率', '难度系数', '区分度']]  # 表头
    for item in analysis:
        table_data.append([
            item['question_text'],
            f"{item['average_score_rate']*100:.1f}%",
            f"{item['difficulty_coefficient']:.2f}",
            f"{item['discrimination_degree']:.2f}"
        ])
    
    # 创建表格样式
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ])
    
    # 创建表格
    table = Table(table_data)
    table.setStyle(table_style)
    
    # 添加表格到内容中
    story.append(table)
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_excel_question_analysis(analysis, filename):
    """生成 Excel 格式的题目分析报告"""
    file_path = os.path.join('exports', f'{filename}.xlsx')
    wb = Workbook()
    ws = wb.active
    ws.title = "题目分析"
    
    # 添加标题
    ws['A1'] = "题目分析报告"
    ws.merge_cells('A1:D1')
    
    # 添加说明
    ws['A3'] = "本报告包含各题目的平均得分率、难度系数和区分度分析"
    ws.merge_cells('A3:D3')
    
    # 添加表头
    headers = ['题目', '平均得分率', '难度系数', '区分度']
    for col, header in enumerate(headers, 1):
        ws.cell(row=5, column=col, value=header)
    
    # 添加数据
    for row, item in enumerate(analysis, 6):
        ws.cell(row=row, column=1, value=item['question_text'])
        ws.cell(row=row, column=2, value=f"{item['average_score_rate']*100:.1f}%")
        ws.cell(row=row, column=3, value=f"{item['difficulty_coefficient']:.2f}")
        ws.cell(row=row, column=4, value=f"{item['discrimination_degree']:.2f}")
        
        # 添加难度等级说明
        difficulty = item['difficulty_coefficient']
        if difficulty >= 0.7:
            ws.cell(row=row, column=3).comment = Comment('难度较大', '题目分析')
        elif difficulty <= 0.3:
            ws.cell(row=row, column=3).comment = Comment('难度较小', '题目分析')
        
        # 添加区分度等级说明
        discrimination = item['discrimination_degree']
        if discrimination >= 0.4:
            ws.cell(row=row, column=4).comment = Comment('区分度良好', '题目分析')
        elif discrimination <= 0.2:
            ws.cell(row=row, column=4).comment = Comment('区分度待改进', '题目分析')
    
    # 添加图例说明
    legend_row = len(analysis) + 7
    ws[f'A{legend_row}'] = "说明："
    ws[f'A{legend_row+1}'] = "难度系数：0-0.3为简单题，0.3-0.7为中等难度，0.7-1.0为困难题"
    ws[f'A{legend_row+2}'] = "区分度：0.4以上为良好，0.2-0.4为一般，0.2以下需改进"
    ws.merge_cells(f'A{legend_row+1}:D{legend_row+1}')
    ws.merge_cells(f'A{legend_row+2}:D{legend_row+2}')
    
    # 设置列宽
    ws.column_dimensions['A'].width = 50  # 题目列宽
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    
    # 设置标题格式
    title_cell = ws['A1']
    title_cell.font = Font(size=14, bold=True)
    title_cell.alignment = Alignment(horizontal='center')
    
    # 设置说明文字格式
    desc_cell = ws['A3']
    desc_cell.font = Font(italic=True)
    desc_cell.alignment = Alignment(horizontal='center')
    
    # 设置表头格式
    for cell in ws[5]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
    
    # 设置数据单元格格式
    for row in ws.iter_rows(min_row=6, max_row=5+len(analysis)):
        for i, cell in enumerate(row):
            cell.alignment = Alignment(horizontal='center' if i > 0 else 'left')
            
            # 为不同的得分率添加不同的颜色
            if i == 1:  # 平均得分率列
                score = float(cell.value.strip('%'))
                if score >= 80:
                    cell.fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
                elif score <= 40:
                    cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    
    # 设置图例说明格式
    for row in range(legend_row, legend_row+3):
        cell = ws[f'A{row}']
        cell.font = Font(size=9, italic=True)
        cell.alignment = Alignment(horizontal='left')
    
    # 添加边框
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    for row in ws.iter_rows(min_row=5, max_row=5+len(analysis)):
        for cell in row:
            cell.border = thin_border
    
    # 保存文件
    wb.save(file_path)
    return file_path

def generate_word_question_analysis(analysis, filename):
    """生成 Word 格式的题目分析报告"""
    file_path = os.path.join('exports', f'{filename}.docx')
    doc = Document()
    
    # 添加标题
    doc.add_heading("题目分析报告", 0)
    
    # 添加说明
    doc.add_paragraph("本报告包含各题目的平均得分率、难度系数和区分度分析")
    
    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '题目'
    header_cells[1].text = '平均得分率'
    header_cells[2].text = '难度系数'
    header_cells[3].text = '区分度'
    
    # 添加数据
    for item in analysis:
        row_cells = table.add_row().cells
        row_cells[0].text = item['question_text']
        row_cells[1].text = f"{item['average_score_rate']*100:.1f}%"
        row_cells[2].text = f"{item['difficulty_coefficient']:.2f}"
        row_cells[3].text = f"{item['discrimination_degree']:.2f}"
    
    # 调整表格格式
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1  # 居中对齐
    
    # 保存文件
    doc.save(file_path)
    return file_path

def generate_score_summary_report(session_id, format_type, filename):
    """生成成绩汇总表"""
    # 获取考试基本信息
    session_query = """
        SELECT 
            id,
            name, 
            subject, 
            start_time, 
            end_time,
            status
        FROM exam_sessions
        WHERE id = ?
    """
    success, session_info = execute_query(session_query, (session_id,))
    if not success or not session_info:
        raise Exception("获取考试信息失败")
    
    session_info = dict(session_info[0])
    
    # 获取学生成绩数据
    scores_query = """
        SELECT 
            s.name as student_name,
            s.student_id,
            SUM(COALESCE(final_score, ai_score)) as total_score,
            RANK() OVER (ORDER BY SUM(COALESCE(final_score, ai_score)) DESC) as rank
        FROM student_answers sa
        JOIN students s ON sa.student_id = s.id
        WHERE sa.session_id = ?
        GROUP BY sa.student_id
        ORDER BY total_score DESC
    """
    success, scores = execute_query(scores_query, (session_id,))
    if not success:
        raise Exception("获取成绩数据失败")
    
    scores = [dict(score) for score in scores]
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_score_summary(session_info, scores, filename)
    elif format_type == 'excel':
        return generate_excel_score_summary(session_info, scores, filename)
    else:  # word
        return generate_word_score_summary(session_info, scores, filename)

def generate_pdf_score_summary(session_info, scores, filename):
    """生成 PDF 格式的成绩汇总表"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph(f"成绩汇总表 - {session_info['name']}", title_style))
    
    # 添加考试信息
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    info_text = f"科目: {session_info['subject']}<br/>"
    info_text += f"考试时间: {session_info['start_time']} 至 {session_info['end_time']}<br/>"
    story.append(Paragraph(info_text, info_style))
    
    # 创建成绩表格
    data = [['排名', '姓名', '学号', '总分']]
    for score in scores:
        data.append([
            str(score['rank']),
            score['student_name'],
            score['student_id'],
            str(score['total_score'])
        ])
    
    # 设置表格样式
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    
    table = Table(data)
    table.setStyle(table_style)
    story.append(table)
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_excel_score_summary(session_info, scores, filename):
    """生成 Excel 格式的成绩汇总表"""
    file_path = os.path.join('exports', f'{filename}.xlsx')
    wb = Workbook()
    ws = wb.active
    
    # 添加标题
    ws['A1'] = f"成绩汇总表 - {session_info['name']}"
    ws.merge_cells('A1:D1')
    ws['A1'].font = Font(name='SimSun', size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # 添加考试信息
    ws['A2'] = f"科目: {session_info['subject']}"
    ws.merge_cells('A2:D2')
    ws['A3'] = f"考试时间: {session_info['start_time']} 至 {session_info['end_time']}"
    ws.merge_cells('A3:D3')
    
    # 添加表头
    headers = ['排名', '姓名', '学号', '总分']
    for col, header in enumerate(headers, 1):
        ws.cell(row=4, column=col, value=header)
    
    # 添加成绩数据
    for row, score in enumerate(scores, 5):
        ws.cell(row=row, column=1, value=score['rank'])
        ws.cell(row=row, column=2, value=score['student_name'])
        ws.cell(row=row, column=3, value=score['student_id'])
        ws.cell(row=row, column=4, value=score['total_score'])
    
    # 设置样式
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=4):
        for cell in row:
            cell.font = Font(name='SimSun')
            cell.alignment = Alignment(horizontal='center')
    
    # 设置列宽
    ws.column_dimensions['A'].width = 10
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 10
    
    # 保存文件
    wb.save(file_path)
    return file_path

def generate_word_score_summary(session_info, scores, filename):
    """生成 Word 格式的成绩汇总表"""
    file_path = os.path.join('exports', f'{filename}.docx')
    doc = Document()
    
    # 添加标题
    doc.add_heading(f"成绩汇总表 - {session_info['name']}", 0)
    
    # 添加考试信息
    doc.add_paragraph(f"科目: {session_info['subject']}")
    doc.add_paragraph(f"考试时间: {session_info['start_time']} 至 {session_info['end_time']}")
    
    # 创建成绩表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 添加表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '排名'
    header_cells[1].text = '姓名'
    header_cells[2].text = '学号'
    header_cells[3].text = '总分'
    
    # 添加成绩数据
    for score in scores:
        row_cells = table.add_row().cells
        row_cells[0].text = str(score['rank'])
        row_cells[1].text = score['student_name']
        row_cells[2].text = score['student_id']
        row_cells[3].text = str(score['total_score'])
    
    # 设置表格格式
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = 1  # 居中对齐
    
    # 保存文件
    doc.save(file_path)
    return file_path

def generate_personal_report(student_id, format_type, filename):
    """生成个人成绩报告"""
    # 获取学生基本信息
    student_query = """
        SELECT 
            s.id,
            s.name,
            s.student_id,
            s.class_name
        FROM students s
        WHERE s.id = ?
    """
    success, student_info = execute_query(student_query, (student_id,))
    if not success or not student_info:
        raise Exception("获取学生信息失败")
    
    student_info = dict(student_info[0])
    
    # 获取考试历史记录
    history_query = """
        WITH exam_scores AS (
            SELECT 
                es.id as session_id,
                es.name as exam_name,
                es.start_time as exam_time,
                SUM(COALESCE(sa.final_score, sa.ai_score)) as score,
                (SELECT SUM(score) FROM questions WHERE session_id = es.id) as total_score,
                RANK() OVER (
                    PARTITION BY es.id 
                    ORDER BY SUM(COALESCE(sa.final_score, sa.ai_score)) DESC
                ) as class_rank,
                AVG(SUM(COALESCE(sa.final_score, sa.ai_score))) OVER (
                    PARTITION BY es.id
                ) as avg_score
            FROM exam_sessions es
            JOIN student_answers sa ON es.id = sa.session_id
            WHERE sa.student_id = ?
            GROUP BY es.id, es.name, es.start_time
        )
        SELECT 
            exam_name,
            exam_time,
            score,
            total_score,
            ROUND(score * 100.0 / total_score, 1) as score_percentage,
            class_rank,
            ROUND(score - avg_score, 1) as vs_average
        FROM exam_scores
        ORDER BY exam_time DESC
    """
    success, history = execute_query(history_query, (student_id,))
    if not success:
        raise Exception("获取历史记录失败")
    
    history = [dict(h) for h in history]
    
    # 获取答题分析数据
    analysis_query = """
        SELECT 
            q.question_text,
            q.question_type,
            q.score as total_score,
            COALESCE(sa.final_score, sa.ai_score) as score,
            sa.answer_text,
            sa.ai_feedback,
            sa.manual_feedback,
            sa.review_status
        FROM student_answers sa
        JOIN questions q ON sa.question_id = q.id
        WHERE sa.student_id = ?
        ORDER BY sa.session_id DESC, q.question_order
    """
    success, analysis = execute_query(analysis_query, (student_id,))
    if not success:
        raise Exception("获取答题分析失败")
    
    analysis = [dict(a) for a in analysis]
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_personal_report(student_info, history, analysis, filename)
    elif format_type == 'excel':
        return generate_excel_personal_report(student_info, history, analysis, filename)
    else:  # word
        return generate_word_personal_report(student_info, history, analysis, filename)

def generate_pdf_personal_report(student_info, history, analysis, filename):
    """生成 PDF 格式的个人成绩报告"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph("个人成绩分析报告", title_style))
    
    # 添加学生信息
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    info_text = f"姓名: {student_info['name']}<br/>"
    info_text += f"学号: {student_info['student_id']}<br/>"
    story.append(Paragraph(info_text, info_style))
    
    # 添加成绩历史表格
    if history:
        data = [['考试名称', '考试时间', '得分', '班级排名', '与平均分差']]
        for h in history:
            data.append([
                h['exam_name'],
                h['exam_time'],
                str(h['score']),
                f"{h['class_rank']}/{h['total_students']}",
                f"{h['score'] - h['avg_score']:.1f}"
            ])
        
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        
        table = Table(data)
        table.setStyle(table_style)
        story.append(table)
        
        # 添加趋势分析
        trend_style = ParagraphStyle(
            'CustomTrend',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=12,
            spaceAfter=10
        )
        
        # 计算成绩趋势
        scores = [h['score'] for h in history]
        avg_scores = [h['avg_score'] for h in history]
        
        trend_text = "成绩趋势分析:<br/>"
        if len(scores) > 1:
            if scores[0] > scores[-1]:
                trend_text += "• 成绩呈上升趋势<br/>"
            elif scores[0] < scores[-1]:
                trend_text += "• 成绩呈下降趋势<br/>"
            else:
                trend_text += "• 成绩保持稳定<br/>"
        
        # 分析与平均分的关系
        vs_avg = [s - a for s, a in zip(scores, avg_scores)]
        if all(v > 0 for v in vs_avg):
            trend_text += "• 持续高于班级平均水平<br/>"
        elif all(v < 0 for v in vs_avg):
            trend_text += "• 需要加强，提高到班级平均水平<br/>"
        else:
            trend_text += "• 成绩波动，建议保持稳定性<br/>"
        
        story.append(Paragraph(trend_text, trend_style))
    else:
        story.append(Paragraph("暂无历史成绩记录", info_style))
    
    # 添加答题分析表格
    if analysis:
        analysis_style = ParagraphStyle(
            'CustomAnalysis',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=12,
            spaceAfter=10
        )
        
        # 创建表格数据
        table_data = [['题目', '总分', '得分', '得分率', '点评']]
        for a in analysis:
            table_data.append([
                a['question_text'],
                str(a['total_score']),
                str(a['score']),
                f"{a['score'] / a['total_score'] * 100:.1f}%",
                a['ai_feedback'] or a['manual_feedback'] or '无'
            ])
        
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        
        table = Table(table_data)
        table.setStyle(table_style)
        story.append(table)
        
        # 添加总体建议
        suggestion_style = ParagraphStyle(
            'CustomSuggestion',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=12,
            spaceAfter=10
        )
        
        # 根据得分情况给出建议
        avg_score_rate = sum(a['score'] / a['total_score'] for a in analysis) / len(analysis)
        suggestion_text = "总体建议:<br/>"
        
        if avg_score_rate >= 0.8:
            suggestion_text += """
            • 继续保持良好的学习状态<br/>
            • 可以尝试更具挑战性的题目<br/>
            • 帮助其他同学，提高自己的理解深度<br/>
            """
        elif avg_score_rate >= 0.6:
            suggestion_text += """
            • 查漏补缺，巩固薄弱环节<br/>
            • 增加练习量，提高解题速度<br/>
            • 及时总结错题，避免重复错误<br/>
            """
        else:
            suggestion_text += """
            • 重点复习基础知识点<br/>
            • 制定详细的学习计划<br/>
            • 建议参加课后辅导<br/>
            • 多与老师和同学交流学习方法<br/>
            """
        
        story.append(Paragraph(suggestion_text, suggestion_style))
    else:
        story.append(Paragraph("暂无答题分析数据", info_style))
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_excel_personal_report(student_info, history, analysis, filename):
    """生成 Excel 格式的个人成绩报告"""
    file_path = os.path.join('exports', f'{filename}.xlsx')
    wb = Workbook()
    ws = wb.active
    
    # 添加标题
    ws['A1'] = "个人成绩分析报告"
    ws.merge_cells('A1:E1')
    ws['A1'].font = Font(name='SimSun', size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # 添加学生信息
    ws['A2'] = f"姓名: {student_info['name']}"
    ws['A3'] = f"学号: {student_info['student_id']}"
    ws.merge_cells('A2:E2')
    ws.merge_cells('A3:E3')
    
    if history:
        # 添加成绩历史表格
        headers = ['考试名称', '考试时间', '得分', '班级排名', '与平均分差']
        for col, header in enumerate(headers, 1):
            ws.cell(row=4, column=col, value=header)
        
        for row, h in enumerate(history, 5):
            ws.cell(row=row, column=1, value=h['exam_name'])
            ws.cell(row=row, column=2, value=h['exam_time'])
            ws.cell(row=row, column=3, value=h['score'])
            ws.cell(row=row, column=4, value=f"{h['class_rank']}/{h['total_students']}")
            ws.cell(row=row, column=5, value=f"{h['score'] - h['avg_score']:.1f}")
        
        # 添加趋势分析
        trend_row = len(history) + 6
        ws.cell(row=trend_row, column=1, value="成绩趋势分析")
        ws.merge_cells(f'A{trend_row}:E{trend_row}')
        
        scores = [h['score'] for h in history]
        avg_scores = [h['avg_score'] for h in history]
        
        trend_row += 1
        if len(scores) > 1:
            if scores[0] > scores[-1]:
                ws.cell(row=trend_row, column=1, value="• 成绩呈上升趋势")
            elif scores[0] < scores[-1]:
                ws.cell(row=trend_row, column=1, value="• 成绩呈下降趋势")
            else:
                ws.cell(row=trend_row, column=1, value="• 成绩保持稳定")
        
        trend_row += 1
        vs_avg = [s - a for s, a in zip(scores, avg_scores)]
        if all(v > 0 for v in vs_avg):
            ws.cell(row=trend_row, column=1, value="• 持续高于班级平均水平")
        elif all(v < 0 for v in vs_avg):
            ws.cell(row=trend_row, column=1, value="• 需要加强，提高到班级平均水平")
        else:
            ws.cell(row=trend_row, column=1, value="• 成绩波动，建议保持稳定性")
    else:
        ws['A4'] = "暂无历史成绩记录"
        ws.merge_cells('A4:E4')
    
    # 添加答题分析表格
    if analysis:
        # 添加表头
        headers = ['题目', '总分', '得分', '得分率', '点评']
        for col, header in enumerate(headers, 1):
            ws.cell(row=5, column=col, value=header)
        
        # 添加答题分析数据
        for row, item in enumerate(analysis, 6):
            ws.cell(row=row, column=1, value=item['question_text'])
            ws.cell(row=row, column=2, value=item['total_score'])
            ws.cell(row=row, column=3, value=item['score'])
            ws.cell(row=row, column=4, value=f"{item['score'] / item['total_score'] * 100:.1f}%")
            
            # 合并AI点评和教师点评
            feedback = []
            if item['ai_feedback']:
                feedback.append(f"AI点评: {item['ai_feedback']}")
            if item['manual_feedback']:
                feedback.append(f"教师点评: {item['manual_feedback']}")
            ws.cell(row=row, column=5, value="\n".join(feedback))
        
        # 添加总体建议
        suggestion_row = len(analysis) + 7
        ws.cell(row=suggestion_row, column=1, value="总体建议")
        ws.merge_cells(f'A{suggestion_row}:E{suggestion_row}')
        
        # 计算平均得分率
        avg_score_rate = sum(a['score'] / a['total_score'] for a in analysis) / len(analysis)
        
        # 根据得分情况给出建议
        suggestions = []
        if avg_score_rate >= 0.8:
            suggestions.extend([
                "• 继续保持良好的学习状态",
                "• 可以尝试更具挑战性的题目",
                "• 帮助其他同学，提高自己的理解深度"
            ])
        elif avg_score_rate >= 0.6:
            suggestions.extend([
                "• 查漏补缺，巩固薄弱环节",
                "• 增加练习量，提高解题速度",
                "• 及时总结错题，避免重复错误"
            ])
        else:
            suggestions.extend([
                "• 重点复习基础知识点",
                "• 制定详细的学习计划",
                "• 建议参加课后辅导",
                "• 多与老师和同学交流学习方法"
            ])
        
        for i, suggestion in enumerate(suggestions):
            ws.cell(row=suggestion_row+1+i, column=1, value=suggestion)
            ws.merge_cells(f'A{suggestion_row+1+i}:E{suggestion_row+1+i}')
    else:
        ws['A5'] = "暂无答题分析数据"
        ws.merge_cells('A5:E5')
    
    # 设置样式
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=5):
        for cell in row:
            cell.font = Font(name='SimSun')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # 设置表头格式
    for cell in ws[5]:
        cell.font = Font(name='SimSun', bold=True)
        cell.fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
    
    # 添加边框
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    for row in ws.iter_rows(min_row=5, max_row=5+len(analysis)):
        for cell in row:
            cell.border = thin_border
    
    # 保存文件
    wb.save(file_path)
    return file_path

def generate_word_personal_report(student_info, history, analysis, filename):
    """生成 Word 格式的个人成绩报告"""
    file_path = os.path.join('exports', f'{filename}.docx')
    doc = Document()
    
    # 添加标题
    doc.add_heading("个人成绩分析报告", 0)
    
    # 添加学生信息
    doc.add_paragraph(f"姓名: {student_info['name']}")
    doc.add_paragraph(f"学号: {student_info['student_id']}")
    
    if history:
        # 添加成绩历史表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '考试名称'
        header_cells[1].text = '考试时间'
        header_cells[2].text = '得分'
        header_cells[3].text = '班级排名'
        header_cells[4].text = '与平均分差'
        
        # 添加数据
        for h in history:
            row_cells = table.add_row().cells
            row_cells[0].text = h['exam_name']
            row_cells[1].text = h['exam_time']
            row_cells[2].text = str(h['score'])
            row_cells[3].text = f"{h['class_rank']}/{h['total_students']}"
            row_cells[4].text = f"{h['score'] - h['avg_score']:.1f}"
        
        # 添加趋势分析
        doc.add_heading("成绩趋势分析", level=1)
        scores = [h['score'] for h in history]
        avg_scores = [h['avg_score'] for h in history]
        
        if len(scores) > 1:
            if scores[0] > scores[-1]:
                doc.add_paragraph("• 成绩呈上升趋势")
            elif scores[0] < scores[-1]:
                doc.add_paragraph("• 成绩呈下降趋势")
            else:
                doc.add_paragraph("• 成绩保持稳定")
        
        vs_avg = [s - a for s, a in zip(scores, avg_scores)]
        if all(v > 0 for v in vs_avg):
            doc.add_paragraph("• 持续高于班级平均水平")
        elif all(v < 0 for v in vs_avg):
            doc.add_paragraph("• 需要加强，提高到班级平均水平")
        else:
            doc.add_paragraph("• 成绩波动，建议保持稳定性")
    else:
        doc.add_paragraph("暂无历史成绩记录")
    
    # 添加答题分析表格
    if analysis:
        # 添加答题分析表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '题目'
        header_cells[1].text = '总分'
        header_cells[2].text = '得分'
        header_cells[3].text = '得分率'
        header_cells[4].text = '点评'
        
        # 添加答题数据
        for item in analysis:
            score_rate = item['final_score'] / item['total_score']
            row_cells = table.add_row().cells
            row_cells[0].text = item['question_text']
            row_cells[1].text = str(item['total_score'])
            row_cells[2].text = str(item['final_score'])
            row_cells[3].text = f"{score_rate*100:.1f}%"
            
            # 合并AI点评和教师点评
            feedback = []
            if item['ai_feedback']:
                feedback.append(f"AI点评: {item['ai_feedback']}")
            if item['manual_feedback']:
                feedback.append(f"教师点评: {item['manual_feedback']}")
            row_cells[4].text = "\n".join(feedback)
        
        # 设置表格格式
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 1  # 居中对齐
                
        # 添加总体建议
        doc.add_heading("总体建议", level=1)
        
        # 计算平均得分率
        avg_score_rate = sum(a['final_score'] / a['total_score'] for a in analysis) / len(analysis)
        
        # 根据得分情况给出建议
        if avg_score_rate >= 0.8:
            suggestions = [
                "• 继续保持良好的学习状态",
                "• 可以尝试更具挑战性的题目",
                "• 帮助其他同学，提高自己的理解深度"
            ]
        elif avg_score_rate >= 0.6:
            suggestions = [
                "• 查漏补缺，巩固薄弱环节",
                "• 增加练习量，提高解题速度",
                "• 及时总结错题，避免重复错误"
            ]
        else:
            suggestions = [
                "• 重点复习基础知识点",
                "• 制定详细的学习计划",
                "• 建议参加课后辅导",
                "• 多与老师和同学交流学习方法"
            ]
        
        for suggestion in suggestions:
            doc.add_paragraph(suggestion)
    else:
        doc.add_paragraph("暂无答题分析数据")
    
    # 保存文件
    doc.save(file_path)
    return file_path

def generate_ability_report(student_id, format_type, filename):
    """生成能力诊断报告"""
    # 获取学生信息
    student_query = """
        SELECT name, student_id
        FROM students
        WHERE id = ?
    """
    success, student_info = execute_query(student_query, (student_id,))
    if not success or not student_info:
        raise Exception("获取学生信息失败")
    
    student_info = dict(student_info[0])
    
    # 获取题型分析数据
    analysis_query = """
        SELECT 
            q.question_type,
            AVG(sa.final_score * 1.0 / q.score) as avg_score_rate,
            COUNT(*) as question_count
        FROM student_answers sa
        JOIN questions q ON sa.question_id = q.id
        WHERE sa.student_id = ?
        GROUP BY q.question_type
    """
    success, analysis = execute_query(analysis_query, (student_id,))
    if not success:
        raise Exception("获取题型分析失败")
    
    analysis = [dict(a) for a in (analysis or [])]
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_ability_report(student_info, analysis, filename)
    elif format_type == 'excel':
        return generate_excel_improvement_report(student_info, analysis, filename)
    else:  # word
        return generate_word_improvement_report(student_info, analysis, filename)

def generate_improvement_report(student_id, format_type, filename):
    """生成进步建议报告"""
    # 获取学生信息
    student_query = """
        SELECT name, student_id
        FROM students
        WHERE id = ?
    """
    success, student_info = execute_query(student_query, (student_id,))
    if not success or not student_info:
        raise Exception("获取学生信息失败")
    
    student_info = dict(student_info[0])
    
    # 获取最近考试的详细分析
    analysis_query = """
        WITH latest_exam AS (
            SELECT session_id
            FROM student_answers
            WHERE student_id = ?
            ORDER BY created_at DESC
            LIMIT 1
        )
        SELECT 
            q.question_text,
            q.score as total_score,
            sa.final_score,
            sa.ai_feedback,
            sa.manual_feedback
        FROM student_answers sa
        JOIN questions q ON sa.question_id = q.id
        JOIN latest_exam le ON sa.session_id = le.session_id
        WHERE sa.student_id = ?
    """
    success, analysis = execute_query(analysis_query, (student_id, student_id))
    if not success:
        raise Exception("获取答题分析失败")
    
    analysis = [dict(a) for a in (analysis or [])]
    
    # 根据不同格式生成报告
    if format_type == 'pdf':
        return generate_pdf_improvement_report(student_info, analysis, filename)
    elif format_type == 'excel':
        return generate_excel_improvement_report(student_info, analysis, filename)
    else:  # word
        return generate_word_improvement_report(student_info, analysis, filename)

def generate_pdf_ability_report(student_info, analysis, filename):
    """生成 PDF 格式的能力诊断报告"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题和学生信息
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph("能力诊断报告", title_style))
    
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    info_text = f"姓名: {student_info['name']}<br/>"
    info_text += f"学号: {student_info['student_id']}<br/>"
    story.append(Paragraph(info_text, info_style))
    
    # 添加题型分析表格
    if analysis:
        data = [['题型', '平均得分率', '题目数量']]
        for a in analysis:
            data.append([
                a['question_type'],
                f"{a['avg_score_rate']*100:.1f}%",
                str(a['question_count'])
            ])
        
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        
        table = Table(data)
        table.setStyle(table_style)
        story.append(table)
        
        # 添加能力分析
        analysis_style = ParagraphStyle(
            'CustomAnalysis',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=12,
            spaceAfter=10
        )
        
        # 找出优势和劣势题型
        sorted_analysis = sorted(analysis, key=lambda x: x['avg_score_rate'], reverse=True)
        strengths = [a for a in sorted_analysis if a['avg_score_rate'] >= 0.8]
        weaknesses = [a for a in sorted_analysis if a['avg_score_rate'] < 0.6]
        
        if strengths:
            strength_text = "优势题型:<br/>"
            for s in strengths:
                strength_text += f"• {s['question_type']}: {s['avg_score_rate']*100:.1f}%<br/>"
            story.append(Paragraph(strength_text, analysis_style))
        
        if weaknesses:
            weakness_text = "需要提升的题型:<br/>"
            for w in weaknesses:
                weakness_text += f"• {w['question_type']}: {w['avg_score_rate']*100:.1f}%<br/>"
            story.append(Paragraph(weakness_text, analysis_style))
    else:
        story.append(Paragraph("暂无题型分析数据", info_style))
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_pdf_improvement_report(student_info, analysis, filename):
    """生成 PDF 格式的进步建议报告"""
    file_path = os.path.join('exports', f'{filename}.pdf')
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    
    # 创建内容
    styles = getSampleStyleSheet()
    story = []
    
    # 添加标题和学生信息
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=16,
        spaceAfter=30
    )
    story.append(Paragraph("学习进步建议", title_style))
    
    info_style = ParagraphStyle(
        'CustomInfo',
        parent=styles['Normal'],
        fontName='SimSun',
        fontSize=12,
        spaceAfter=10
    )
    info_text = f"姓名: {student_info['name']}<br/>"
    info_text += f"学号: {student_info['student_id']}<br/>"
    story.append(Paragraph(info_text, info_style))
    
    if analysis:
        # 添加答题分析
        for i, a in enumerate(analysis, 1):
            question_style = ParagraphStyle(
                'CustomQuestion',
                parent=styles['Normal'],
                fontName='SimSun',
                fontSize=12,
                spaceAfter=10
            )
            
            score_rate = a['final_score'] / a['total_score']
            question_text = f"题目 {i}:<br/>"
            question_text += f"得分率: {score_rate*100:.1f}%<br/>"
            if a['ai_feedback']:
                question_text += f"AI点评: {a['ai_feedback']}<br/>"
            if a['manual_feedback']:
                question_text += f"教师点评: {a['manual_feedback']}<br/>"
            
            story.append(Paragraph(question_text, question_style))
        
        # 添加总体建议
        suggestion_style = ParagraphStyle(
            'CustomSuggestion',
            parent=styles['Normal'],
            fontName='SimSun',
            fontSize=12,
            spaceAfter=10
        )
        
        # 根据得分情况给出建议
        avg_score_rate = sum(a['final_score'] / a['total_score'] for a in analysis) / len(analysis)
        suggestion_text = "总体建议:<br/>"
        
        if avg_score_rate >= 0.8:
            suggestion_text += """
            • 继续保持良好的学习状态<br/>
            • 可以尝试更具挑战性的题目<br/>
            • 帮助其他同学，提高自己的理解深度<br/>
            """
        elif avg_score_rate >= 0.6:
            suggestion_text += """
            • 查漏补缺，巩固薄弱环节<br/>
            • 增加练习量，提高解题速度<br/>
            • 及时总结错题，避免重复错误<br/>
            """
        else:
            suggestion_text += """
            • 重点复习基础知识点<br/>
            • 制定详细的学习计划<br/>
            • 建议参加课后辅导<br/>
            • 多与老师和同学交流学习方法<br/>
            """
        
        story.append(Paragraph(suggestion_text, suggestion_style))
    else:
        story.append(Paragraph("暂无答题分析数据", info_style))
    
    # 生成 PDF
    doc.build(story)
    return file_path

def generate_excel_improvement_report(student_info, analysis, filename):
    """生成 Excel 格式的进步建议报告"""
    file_path = os.path.join('exports', f'{filename}.xlsx')
    wb = Workbook()
    ws = wb.active
    ws.title = "进步建议报告"
    
    # 添加标题
    ws['A1'] = "学习进步建议报告"
    ws.merge_cells('A1:D1')
    ws['A1'].font = Font(name='SimSun', size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # 添加学生信息
    ws['A2'] = f"姓名: {student_info['name']}"
    ws['A3'] = f"学号: {student_info['student_id']}"
    ws.merge_cells('A2:D2')
    ws.merge_cells('A3:D3')
    
    if analysis:
        # 添加表头
        headers = ['题目', '总分', '得分', '得分率', '点评']
        for col, header in enumerate(headers, 1):
            ws.cell(row=5, column=col, value=header)
        
        # 添加答题分析数据
        for row, item in enumerate(analysis, 6):
            score_rate = item['final_score'] / item['total_score']
            ws.cell(row=row, column=1, value=item['question_text'])
            ws.cell(row=row, column=2, value=item['total_score'])
            ws.cell(row=row, column=3, value=item['final_score'])
            ws.cell(row=row, column=4, value=f"{score_rate*100:.1f}%")
            
            # 合并AI点评和教师点评
            feedback = []
            if item['ai_feedback']:
                feedback.append(f"AI点评: {item['ai_feedback']}")
            if item['manual_feedback']:
                feedback.append(f"教师点评: {item['manual_feedback']}")
            ws.cell(row=row, column=5, value="\n".join(feedback))
        
        # 添加总体建议
        suggestion_row = len(analysis) + 7
        ws.cell(row=suggestion_row, column=1, value="总体建议")
        ws.merge_cells(f'A{suggestion_row}:D{suggestion_row}')
        
        # 计算平均得分率
        avg_score_rate = sum(a['final_score'] / a['total_score'] for a in analysis) / len(analysis)
        
        # 根据得分情况给出建议
        suggestions = []
        if avg_score_rate >= 0.8:
            suggestions.extend([
                "• 继续保持良好的学习状态",
                "• 可以尝试更具挑战性的题目",
                "• 帮助其他同学，提高自己的理解深度"
            ])
        elif avg_score_rate >= 0.6:
            suggestions.extend([
                "• 查漏补缺，巩固薄弱环节",
                "• 增加练习量，提高解题速度",
                "• 及时总结错题，避免重复错误"
            ])
        else:
            suggestions.extend([
                "• 重点复习基础知识点",
                "• 制定详细的学习计划",
                "• 建议参加课后辅导",
                "• 多与老师和同学交流学习方法"
            ])
        
        for i, suggestion in enumerate(suggestions):
            ws.cell(row=suggestion_row+1+i, column=1, value=suggestion)
            ws.merge_cells(f'A{suggestion_row+1+i}:D{suggestion_row+1+i}')
    else:
        ws['A5'] = "暂无答题分析数据"
        ws.merge_cells('A5:D5')
    
    # 设置列宽
    ws.column_dimensions['A'].width = 40  # 题目列宽
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 50  # 点评列宽
    
    # 设置样式
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=5):
        for cell in row:
            cell.font = Font(name='SimSun')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    # 设置表头格式
    for cell in ws[5]:
        cell.font = Font(name='SimSun', bold=True)
        cell.fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
    
    # 添加边框
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    for row in ws.iter_rows(min_row=5, max_row=5+len(analysis)):
        for cell in row:
            cell.border = thin_border
    
    # 保存文件
    wb.save(file_path)
    return file_path

def generate_word_improvement_report(student_info, analysis, filename):
    """生成 Word 格式的进步建议报告"""
    file_path = os.path.join('exports', f'{filename}.docx')
    doc = Document()
    
    # 添加标题
    doc.add_heading("学习进步建议报告", 0)
    
    # 添加学生信息
    doc.add_paragraph(f"姓名: {student_info['name']}")
    doc.add_paragraph(f"学号: {student_info['student_id']}")
    
    if analysis:
        # 添加答题分析表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '题目'
        header_cells[1].text = '总分'
        header_cells[2].text = '得分'
        header_cells[3].text = '得分率'
        header_cells[4].text = '点评'
        
        # 添加答题数据
        for item in analysis:
            score_rate = item['final_score'] / item['total_score']
            row_cells = table.add_row().cells
            row_cells[0].text = item['question_text']
            row_cells[1].text = str(item['total_score'])
            row_cells[2].text = str(item['final_score'])
            row_cells[3].text = f"{score_rate*100:.1f}%"
            
            # 合并AI点评和教师点评
            feedback = []
            if item['ai_feedback']:
                feedback.append(f"AI点评: {item['ai_feedback']}")
            if item['manual_feedback']:
                feedback.append(f"教师点评: {item['manual_feedback']}")
            row_cells[4].text = "\n".join(feedback)
        
        # 设置表格格式
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 1  # 居中对齐
                
        # 添加总体建议
        doc.add_heading("总体建议", level=1)
        
        # 计算平均得分率
        avg_score_rate = sum(a['final_score'] / a['total_score'] for a in analysis) / len(analysis)
        
        # 根据得分情况给出建议
        if avg_score_rate >= 0.8:
            suggestions = [
                "• 继续保持良好的学习状态",
                "• 可以尝试更具挑战性的题目",
                "• 帮助其他同学，提高自己的理解深度"
            ]
        elif avg_score_rate >= 0.6:
            suggestions = [
                "• 查漏补缺，巩固薄弱环节",
                "• 增加练习量，提高解题速度",
                "• 及时总结错题，避免重复错误"
            ]
        else:
            suggestions = [
                "• 重点复习基础知识点",
                "• 制定详细的学习计划",
                "• 建议参加课后辅导",
                "• 多与老师和同学交流学习方法"
            ]
        
        for suggestion in suggestions:
            doc.add_paragraph(suggestion)
    else:
        doc.add_paragraph("暂无答题分析数据")
    
    # 保存文件
    doc.save(file_path)
    return file_path

@app.route('/api/user/info')
@login_required
def get_user_info():
    """获取当前用户信息"""
    if 'email' not in session:
        return jsonify({'error': '未登录'}), 401
        
    query = """
        SELECT id, email, created_at
        FROM users
        WHERE email = ?
    """
    success, result = execute_query(query, (session['email'],))
    
    if not success or not result:
        return jsonify({'error': '获取用户信息失败'}), 500
        
    user = result[0]
    return jsonify({
        'id': user['id'],
        'email': user['email'],
        'created_at': user['created_at']
    })

@app.route('/api/user/change-password', methods=['POST'])
@login_required
def change_password():
    """修改用户密码"""
    if 'email' not in session:
        return jsonify({'error': '未登录'}), 401
        
    data = request.get_json()
    current_password = data.get('currentPassword')
    new_password = data.get('newPassword')
    
    if not all([current_password, new_password]):
        return jsonify({'error': '缺少必要参数'}), 400
        
    # 验证当前密码
    query = "SELECT id, password FROM users WHERE email = ?"
    success, result = execute_query(query, (session['email'],))
    
    if not success or not result:
        return jsonify({'error': '用户不存在'}), 404
        
    user = result[0]
    if current_password != user['password']:  # 实际应用中应使用密码哈希比较
        return jsonify({'error': '当前密码错误'}), 400
        
    # 更新密码
    update_query = "UPDATE users SET password = ? WHERE id = ?"
    success, _ = execute_query(update_query, (new_password, user['id']))


@app.route('/api/exam-sessions', methods=['GET'])
@login_required
def get_exam_sessions_index():
    # 获取当前用户ID
    user_id = get_current_user_id()
    """获取所有考试场次"""
    query = """
        SELECT id, name as session_name 
        FROM exam_sessions 
        WHERE status != 'completed' AND created_by = ?
        ORDER BY start_time ASC
    """
    
    success, sessions = execute_query(query, (user_id,))
    
    if not success:
        return jsonify({'error': '获取考试场次失败'}), 500
    
    # 确保返回的是列表
    sessions = [dict(session) for session in sessions] if sessions else []
    
    return jsonify(sessions)

@app.route('/api/review/sessions', methods=['GET'])
@login_required
def get_review_sessions():
    # 获取当前用户ID
    user_id = get_current_user_id()
    """获取需要复核的考试场次"""
    query = """
        SELECT 
            es.id,
            es.name,
            es.subject,
            es.start_time,
            es.end_time
        FROM exam_sessions es
        WHERE es.created_by = ?
    """
    
    print("Executing sessions query")  # 调试日志
    success, sessions = execute_query(query, (user_id,))
    print(f"Sessions result: {sessions}")  # 调试日志
    
    if not success:
        print("Failed to get sessions")  # 调试日志
        return jsonify({'error': '获取考试场次失败'}), 500
        
    return jsonify([dict(session) for session in (sessions or [])])

@app.route('/extracted_files/<path:filename>')
def uploaded_file(filename):
    """提供上传文件的访问"""
    try:
        return send_from_directory(EXTRACT_FOLDER, filename)
    except Exception as e:
        logger.error(f'访问文件失败: {filename} - {str(e)}')
        return jsonify({'error': '文件不存在'}), 404

@app.route('/api/start-grading', methods=['POST'])
@login_required
def start_grading():
    """开始自动阅卷"""
    try:
        # 获取考试场次ID
        session_id = request.form.get('session_id')
        if not session_id:
            return jsonify({'success': False, 'message': '缺少考试场次ID'}), 400

        # 处理上传的文件
        files = []
        for key in request.files:
            file = request.files[key]
            if file and allowed_file(file.filename):
                # 保存文件到 extracted_files 目录
                filename = secure_filename(file.filename)
                save_path = os.path.join(EXTRACT_FOLDER, filename)
                file.save(save_path)
                files.append(save_path)

        if not files:
            return jsonify({'success': False, 'message': '未上传有效的试卷文件'}), 400

        # 获取考试信息
        query = """
            SELECT id, name, subject, exam_score
            FROM exam_sessions
            WHERE id = ?
        """
        success, exam_info = execute_query(query, (session_id,))
        if not success or not exam_info:
            return jsonify({'success': False, 'message': '获取考试信息失败'}), 500
        
        exam_info = exam_info[0]
        total_score = float(exam_info['exam_score'])

        # 处理每个上传的文件
        for file_path in files:
            try:
                # 1. 使用 slip 模块进行 OCR 识别
                # 为每个文件创建单独的输出文件
                ocr_output = f"ocr_results_{os.path.basename(file_path)}.txt"
                split_columns_and_rows(file_path, ocr_output)

                # 2. 分割题目和答案
                from split_and_ocr.read.questionsplit import readexit
                readexit()

                # 3. 获取学生信息（从文件名或OCR结果中提取）
                # 这里假设文件名格式为：学号_姓名.pdf
                student_info = os.path.basename(file_path).split('.')[0].split('_')
                if len(student_info) != 2:
                    raise Exception("文件名格式错误，应为：学号_姓名.pdf")
                
                student_id, student_name = student_info

                # 检查学生是否存在，不存在则创建
                check_query = "SELECT id FROM students WHERE student_id = ?"
                success, student = execute_query(check_query, (student_id,))
                
                if not success:
                    raise Exception("查询学生信息失败")
                
                if not student:
                    # 创建新学生
                    insert_query = """
                        INSERT INTO students (name, student_id)
                        VALUES (?, ?)
                    """
                    success, student = execute_query(insert_query, (student_name, student_id))
                    if not success:
                        raise Exception("创建学生记录失败")
                    student_db_id = student.lastrowid
                else:
                    student_db_id = student[0]['id']

                    # 4. AI 评分
                    # 读取分割后的答案文件
                    current_dir = os.path.dirname(os.path.abspath(__file__))
                    read_dir = os.path.join(current_dir, "split_and_ocr", "read")
                    
                    # 获取read目录下所有txt文件
                    os.chdir(read_dir)  # 切换到read目录
                    files = [f for f in os.listdir() if f.endswith('.txt')]
                    
                    # 定义题型和对应的文件名匹配模式
                    question_patterns = {
                        '选择题': r'选择.+',
                        '填空题': r'填空.+',
                        '简答题': r'简答.+',
                        '判断题': r'判断.+',
                        '编程题': r'编程.+'
                    }
                    
                    # 处理每种题型
                    for q_type, pattern in question_patterns.items():
                        # 匹配对应题型的文件
                        matched_files = [f for f in files if re.search(pattern, f)]
                        
                        if matched_files:  # 如果找到匹配的文件
                            answer_file = os.path.join(read_dir, matched_files[0])
                            with open(answer_file, 'r', encoding='utf-8') as f:
                                answers = f.read()
                            f.close()
                                
                            # 从数据库获取该类型的题目信息
                            if q_type == '选择题':
                                query = """
                                    SELECT q.id, q.question_text, q.question_score,
                                            mcq.options, mcq.correct_option, mcq.explanation
                                    FROM questions q
                                    JOIN multiple_choice_questions mcq ON q.id = mcq.question_id
                                    WHERE q.session_id = ? AND q.question_type = '选择题'
                                    ORDER BY q.question_order
                                """
                            elif q_type == '编程题':
                                query = """
                                    SELECT q.id, q.question_text, q.question_score,
                                            pq.test_cases, pq.expected_output, pq.solution_template,
                                            pq.hints
                                    FROM questions q
                                    JOIN programming_questions pq ON q.id = pq.question_id
                                    WHERE q.session_id = ? AND q.question_type = '编程题'
                                    ORDER BY q.question_order
                                """
                            elif q_type == '简答题':
                                query = """
                                    SELECT q.id, q.question_text, q.question_score,
                                            saq.model_answer, saq.key_points, saq.grading_criteria
                                    FROM questions q
                                    JOIN short_answer_questions saq ON q.id = saq.question_id
                                    WHERE q.session_id = ? AND q.question_type = '简答题'
                                    ORDER BY q.question_order
                                """
                            elif q_type == '判断题':
                                query = """
                                    SELECT q.id, q.question_text, q.question_score,
                                            tfq.correct_answer, tfq.explanation
                                    FROM questions q
                                    JOIN true_false_questions tfq ON q.id = tfq.question_id
                                    WHERE q.session_id = ? AND q.question_type = '判断题'
                                    ORDER BY q.question_order
                                """
                            elif q_type == 'programming':
                                data = request.json
                                # 编程题处理
                                answer = data.get('answer', '')
                                test_cases = data.get('testCases', [])
                                sample_input = data.get('sampleInput', '')
                                sample_output = data.get('sampleOutput', '')
                                time_limit = data.get('timeLimit', 1000)
                                memory_limit = data.get('memoryLimit', 256)
                                
                                if mode == 'question-bank':
                                    # 题库模式，插入到题库表
                                    
                                    # 准备测试用例数据
                                    test_cases_json = json.dumps(test_cases)
                                    
                                    query = """
                                        INSERT INTO programming_questions 
                                        (subject, question_text, reference_solution, test_cases, 
                                            sample_input, sample_output, time_limit, memory_limit,
                                            hints, score, difficulty, created_by)
                                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                    """
                                    params = (
                                        subject, content, answer, test_cases_json, 
                                        sample_input, sample_output, time_limit, memory_limit,
                                        explanation, score, difficulty, user_id
                                    )
                                else:
                                    # 考试模式，类似选择题的处理...
                                    # 为简洁起见，这里省略具体实现
                                    pass
                            else:  # 填空题
                                query = """
                                    SELECT q.id, q.question_text, q.question_score,
                                            fbq.correct_answer, fbq.alternative_answers,
                                            fbq.explanation
                                    FROM questions q
                                    JOIN fill_blank_questions fbq ON q.id = fbq.question_id
                                    WHERE q.session_id = ? AND q.question_type = '填空题'
                                    ORDER BY q.question_order
                                """
                            
                            success, questions = execute_query(query, (session_id,))
                            
                            if not success:
                                raise Exception(f"获取{q_type}题目信息失败")

                            # 构建评分提示
                            if q_type == '选择题':
                                prompt = f"""
                                请对以下选择题答案进行评分。

                                考试信息：
                                - 考试名称：{exam_info['name']}
                                - 科目：{exam_info['subject']}
                                
                                题目信息：
                                {[f'''
                                第{idx+1}题（{q['question_score']}分）
                                题目：{q['question_text']}
                                选项：{q['options']}
                                正确答案：{q['correct_option']}
                                解析：{q['explanation'] or '无'}
                                ''' for idx, q in enumerate(questions)]}
                                
                                学生答案：
                                {answers}

                                请根据以下规则评分：
                                1. 答案完全正确得满分
                                2. 答案错误得0分
                                3. 如果有多个选择题，请分别给出每道题的得分
                                4. 给出评价时，说明答错的原因和正确的解析

                                请以JSON格式返回结果，包含以下字段：
                                score: 总分数（数字）
                                feedback: 详细的评价意见（文字）
                                question_scores: 每道题的得分列表（数组）
                                """
                            elif q_type == '填空题':
                                prompt = f"""
                                请对以下填空题答案进行评分。

                                考试信息：
                                - 考试名称：{exam_info['name']}
                                - 科目：{exam_info['subject']}
                                
                                题目信息：
                                {[f'''
                                第{idx+1}题（{q['question_score']}分）
                                题目：{q['question_text']}
                                标准答案：{q['correct_answer']}
                                其他可接受答案：{q['alternative_answers'] or '无'}
                                解析：{q['explanation'] or '无'}
                                ''' for idx, q in enumerate(questions)]}
                                
                                学生答案：
                                {answers}

                                请根据以下规则评分：
                                1. 答案完全匹配标准答案或可接受答案之一得满分
                                2. 答案部分正确酌情给分（根据关键词匹配程度）
                                3. 答案完全错误得0分
                                4. 如果有多个空，请分别给出每个空的得分
                                5. 给出评价时，指出错误之处，并给出正确答案

                                请以JSON格式返回结果，包含以下字段：
                                score: 总分数（数字）
                                feedback: 详细的评价意见（文字）
                                question_scores: 每道题的得分列表（数组）
                                """
                            else:  # 简答题
                                prompt = f"""
                                请对以下简答题答案进行评分。

                                考试信息：
                                - 考试名称：{exam_info['name']}
                                - 科目：{exam_info['subject']}
                                
                                题目信息：
                                {[f'''
                                第{idx+1}题（{q['question_score']}分）
                                题目：{q['question_text']}
                                参考答案：{q['model_answer']}
                                关键点：{q['key_points']}
                                评分标准：{q['grading_criteria'] or '无'}
                                ''' for idx, q in enumerate(questions)]}
                                
                                学生答案：
                                {answers}

                                请根据以下规则评分：
                                1. 答案要点完全正确且表述清晰得满分
                                2. 根据答案要点的覆盖程度和表述准确性酌情给分：
                                    - 核心概念准确：40%（关键点覆盖程度）
                                    - 论述完整性：30%（答案结构和逻辑）
                                    - 例子/证据支持：20%（实例说明）
                                    - 语言表达：10%（表述清晰度）
                                3. 给出评价时：
                                    - 指出答案中正确的关键点
                                    - 指出缺失或错误的关键点
                                    - 给出改进建议
                                    - 提供完整的参考答案

                                请以JSON格式返回结果，包含以下字段：
                                score: 总分数（数字）
                                feedback: 详细的评价意见（文字）
                                question_scores: 每道题的得分列表（数组）
                                scoring_details: {
                                    "核心概念": 分数,
                                    "论述完整性": 分数,
                                    "例子支持": 分数,
                                    "语言表达": 分数
                                }
                                """
                                
                            # 使用AI评分
                            from split_and_ocr.ai import aiapi
                            result = aiapi("", prompt)
                            try:
                                score_info = json.loads(result)
                            except:
                                score_info = {
                                    'score': 0,
                                    'feedback': '评分失败',
                                    'question_scores': [0] * len(questions)
                                }

                            # 保存每道题的评分结果
                            for idx, question in enumerate(questions):
                                question_score = score_info.get('question_scores', [])[idx] if idx < len(score_info.get('question_scores', [])) else 0
                                
                                insert_query = """
                                    INSERT INTO student_answers 
                                    (session_id, student_id, question_id, question_type, 
                                        answer_text, ai_score, ai_feedback, scoring_details)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                                """
                                success, _ = execute_query(
                                    insert_query, 
                                    (session_id, student_db_id, question['id'], q_type,
                                        answers, question_score, score_info['feedback'],
                                        json.dumps(score_info.get('scoring_details', {})))
                                )
                                if not success:
                                    raise Exception(f"{q_type}第{idx+1}题评分结果保存失败")
                
            except Exception as e:
                print(f"处理文件 {file_path} 时出错: {str(e)}")
                continue
            finally:
                # 清理临时文件
                if os.path.exists(file_path):
                    os.remove(file_path)

        # 更新考试场次状态
        update_query = """
            UPDATE exam_sessions 
            SET status = 'graded' 
            WHERE id = ?
        """
        success, _ = execute_query(update_query, (session_id,))
        
        if not success:
            return jsonify({'success': False, 'message': '更新考试状态失败'}), 500

        return jsonify({
            'success': True,
            'message': '阅卷完成'
        })

    except Exception as e:
        print(f"阅卷过程出错: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'阅卷过程出错: {str(e)}'
        }), 500
    
@app.route('/question/input')
@login_required
def question_input():
    """题目录入页面路由"""
    return render_template('question_input.html')

@app.route('/api/questions', methods=['POST'])
@login_required
def add_question():
    """添加题目的API端点"""
    try:
        print("============ 添加题目请求 ============")
        print(f"请求方法: {request.method}")
        print(f"URL参数: {request.args}")
        print(f"Content-Type: {request.headers.get('Content-Type')}")
        
        data = request.get_json()
        print(f"请求数据: {data}")
        
        # 获取当前用户ID
        user_id = get_current_user_id()
        
        # 从请求中获取题目数据
        mode = data.get('mode')
        print("当前模式:", mode)
        question_type = data.get('type')
        difficulty = data.get('difficulty')
        score = data.get('score')
        content = data.get('content')
        explanation = data.get('explanation', '')
        subject = data.get('subject', '未分类')  # 获取用户指定的科目，如果没有则默认为未分类
        print("接收到的请求数据:", data)
        
        # 获取考试ID，优先从URL参数获取，如果没有则从请求数据中获取
        url_exam_id = request.args.get('exam_id')
        if url_exam_id:
            data['examId'] = url_exam_id
            mode = 'exam'  # 强制设置为考试模式
            print(f"从URL参数获取考试ID: {url_exam_id}")
        
        # 根据题目类型处理不同的数据
        if question_type == 'single':
            options = data.get('options', [])
            answer = data.get('answer')
            
            # 确保选项数量足够
            if len(options) < 2:
                return jsonify({'status': 'error', 'message': '至少需要两个选项'}), 400
            
            # 确保答案有效
            if not answer:
                return jsonify({'status': 'error', 'message': '请选择正确答案'}), 400
            
            # 转换分值和难度为整数
            try:
                score = int(float(score))
                difficulty = int(difficulty)
            except (ValueError, TypeError):
                return jsonify({'status': 'error', 'message': '分值和难度等级必须是数字'}), 400
            
            # 将选项转换为A, B, C, D格式
            option_data = {}
            for i, option in enumerate(options):
                option_key = f"option_{chr(97 + i)}"  # a, b, c, d...
                # 确保选项有前缀 "A. ", "B. " 等
                if not option.startswith(f"{chr(65 + i)}. "):
                    option = f"{chr(65 + i)}. {option}"
                option_data[option_key] = option
            
            print(f"处理后的选项: {option_data}")  # 调试输出
            
            # 准备答案数据
            answer_text = answer if isinstance(answer, str) else ','.join(answer)
            
            # 将数字索引转换为对应的字母(A/B/C/D)
            try:
                answer_index = int(answer_text)
                answer_letter = chr(65 + answer_index)  # 将0转为A, 1转为B, 以此类推
                answer_text = answer_letter
                print(f"答案转换: 索引 {answer_index} -> 字母 {answer_letter}")  # 调试输出
            except (ValueError, TypeError):
                # 如果转换失败，保持原样
                print(f"答案未转换: {answer_text}")  # 调试输出
                pass
            
            # 创建选择题
            if mode == 'question-bank':
                # 题库模式，插入到题库表
                
                query = """
                    INSERT INTO multiple_choice_questions 
                    (subject, question_text, option_a, option_b, option_c, option_d, 
                     correct_answer, explanation, score, difficulty, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 确保选项字段都有值
                option_a = option_data.get('option_a', '')
                option_b = option_data.get('option_b', '')
                option_c = option_data.get('option_c', '')
                option_d = option_data.get('option_d', '')
                
                # 日志记录选项与答案
                print(f"选项A: {option_a}")
                print(f"选项B: {option_b}")
                print(f"选项C: {option_c}")
                print(f"选项D: {option_d}")
                print(f"正确答案(字母形式): {answer_text}")
                
                params = (
                    subject, content, 
                    option_a, option_b, option_c, option_d,
                    answer_text, explanation, score, difficulty, user_id
                )
                
                print(f"插入选择题: {params}")  # 调试输出
            else:
                # 考试模式，插入到考试题目表
                exam_id = data.get('examId') or request.args.get('exam_id')
                section_id = data.get('sectionId')
                
                if not exam_id:
                    return jsonify({'status': 'error', 'message': '请选择考试场次'}), 400

                # 验证考试ID是否有效
                query = "SELECT id FROM exam_sessions WHERE id = ?"
                success, result = execute_query(query, (exam_id,))
                if not success or not result:
                    return jsonify({
                        'status': 'error',
                        'message': '无效的考试ID'
                    }), 400
                
                # 获取该考试的最大题目序号
                query_max = """
                    SELECT MAX(question_order) as max_order FROM questions 
                    WHERE session_id = ?
                """
                success, result = execute_query(query_max, (exam_id,))
                if not success:
                    return jsonify({'status': 'error', 'message': '获取题目序号失败'}), 500
                
                order = (result[0]['max_order'] or 0) + 1 if result else 1
                
                # 首先插入到题库表
                query_mcq = """
                    INSERT INTO multiple_choice_questions 
                    (subject, question_text, option_a, option_b, option_c, option_d, 
                     correct_answer, explanation, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 确保选项字段都有值
                option_a = option_data.get('option_a', '')
                option_b = option_data.get('option_b', '')
                option_c = option_data.get('option_c', '')
                option_d = option_data.get('option_d', '')
                
                # 获取考试场次的科目信息
                query_subject = """
                    SELECT subject FROM exam_sessions WHERE id = ?
                """
                success, subject_result = execute_query(query_subject, (exam_id,))
                if not success or not subject_result:
                    return jsonify({'status': 'error', 'message': '获取考试科目信息失败'}), 500
                
                exam_subject = subject_result[0]['subject']
                
                params_mcq = (
                    exam_subject, content, 
                    option_a, option_b, option_c, option_d,
                    answer_text, explanation, score, difficulty, user_id, False
                )
                
                print(f"插入考试选择题: {params_mcq}")  # 调试输出
                
                success, mcq_result = execute_query(query_mcq, params_mcq)
                if not success:
                    return jsonify({'status': 'error', 'message': '保存选择题失败'}), 500
                
                mcq_id = mcq_result
                
                # 然后插入到考试题目表
                query = """
                    INSERT INTO questions 
                    (session_id, question_type, question_text, score, question_order, 
                     source_question_id, source_table)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                
                params = (
                    exam_id, 
                    'multiple_choice', 
                    content, score, order, mcq_id, 'multiple_choice_questions'
                )

                # 在插入考试题目表的代码后面添加：
                success, question_result = execute_query(query, params)
                if not success:
                    return jsonify({'status': 'error', 'message': '保存考试题目失败'}), 500

                return jsonify({
                    'status': 'success',
                    'message': '题目添加成功',
                    'data': {
                        'question_id': question_result
                    }
                }), 200
        elif question_type == 'truefalse':
            # 判断题处理
            answer = data.get('answer')
            
            # 调试信息
            print(f"处理判断题: data={data}")
            print(f"answer={answer}, 类型={type(answer)}")
            
            # 确保答案有效
            if answer is None or answer == '':
                return jsonify({'status': 'error', 'message': '请选择正确答案'}), 400
            
            try:
                # 转换答案为布尔值
                # 前端编码: 0="正确", 1="错误"
                # SQLite中布尔值: 0=False, 1=True
                # 前端->数据库: 0(正确)->1(True), 1(错误)->0(False)
                answer_bool = (answer == '0')
                
                # 转换分值和难度为整数
                try:
                    score = int(float(score))
                    difficulty = int(difficulty)
                except (ValueError, TypeError):
                    return jsonify({'status': 'error', 'message': '分值和难度等级必须是数字'}), 400
                
                # 创建判断题
                if mode == 'question-bank':
                    # 题库模式，插入到判断题表
                    query = """
                        INSERT INTO true_false_questions 
                        (subject, question_text, correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    
                    params = (
                        subject, content, 
                        answer_bool, explanation, score, difficulty, user_id, True
                    )
                    
                    print(f"插入判断题: {params}")  # 调试输出
                else:
                    # 考试模式，插入到考试题目表
                    exam_id = data.get('examId')
                    section_id = data.get('sectionId')
                    
                    if not exam_id:
                        return jsonify({'status': 'error', 'message': '请选择考试场次'}), 400
                    
                    # 获取该考试的最大题目序号
                    query_max = """
                        SELECT MAX(question_order) as max_order FROM questions 
                        WHERE session_id = ?
                    """
                    success, result = execute_query(query_max, (exam_id,))
                    if not success:
                        return jsonify({'status': 'error', 'message': '获取题目序号失败'}), 500
                    
                    order = (result[0]['max_order'] or 0) + 1 if result else 1
                    
                    # 首先插入到判断题表
                    query_tf = """
                        INSERT INTO true_false_questions 
                        (subject, question_text, correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    
                    # 获取考试场次的科目信息
                    query_subject = """
                        SELECT subject FROM exam_sessions WHERE id = ?
                    """
                    success, subject_result = execute_query(query_subject, (exam_id,))
                    if not success or not subject_result:
                        return jsonify({'status': 'error', 'message': '获取考试科目信息失败'}), 500
                    
                    exam_subject = subject_result[0]['subject']
                    
                    params_tf = (
                        exam_subject, content, 
                        answer_bool, explanation, score, difficulty, user_id, False
                    )
                    
                    print(f"插入考试判断题: {params_tf}")  # 调试输出
                    
                    success, tf_result = execute_query(query_tf, params_tf)
                    if not success:
                        return jsonify({'status': 'error', 'message': '保存判断题失败'}), 500
                    
                    tf_id = tf_result
                    
                    # 然后插入到考试题目表
                    query = """
                        INSERT INTO questions 
                        (session_id, question_type, question_text, score, question_order, 
                         source_question_id, source_table)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """
                    
                    params = (
                        exam_id, 
                        'true_false', 
                        content, score, order, tf_id, 'true_false_questions'
                    )

                    # 在插入考试题目表的代码后面添加：
                    success, question_result = execute_query(query, params)
                    if not success:
                        return jsonify({'status': 'error', 'message': '保存考试题目失败'}), 500

                    return jsonify({
                        'status': 'success',
                        'message': '题目添加成功',
                        'data': {
                            'question_id': question_result
                        }
                    }), 200
            except ValueError as ve:
                print(f"判断题答案格式错误: {ve}")
                return jsonify({'status': 'error', 'message': f'判断题答案格式错误: {ve}'}), 400
        elif question_type == 'programming':
            # 编程题处理
            app.logger.info(f"编程题完整数据: {json.dumps(data)}")
            app.logger.info(f"样例输入: {data.get('sample_input', '(无)')}")
            app.logger.info(f"样例输出: {data.get('sample_output', '(无)')}")
            # 编程题处理
            answer = data.get('answer', '')
            test_cases = data.get('test_cases', [])  # 修改为test_cases而不是testCases
            sample_input = data.get('sample_input', '')  # 修改为sample_input而不是sampleinput
            sample_output = data.get('sample_output', '')  # 修改为sample_output而不是sampleoutput
            time_limit = data.get('time_limit', 1000)
            memory_limit = data.get('memory_limit', 256)
            
            if mode == 'question-bank':
                # 题库模式，插入到题库表
                
                # 准备测试用例数据
                test_cases_json = json.dumps(test_cases)
                
                query = """
                    INSERT INTO programming_questions 
                    (subject, question_text, reference_solution, test_cases, 
                     sample_input, sample_output, time_limit, memory_limit,
                     hints, score, difficulty, created_by)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                params = (
                    subject, content, answer, test_cases_json, 
                    sample_input, sample_output, time_limit, memory_limit,
                    explanation, score, difficulty, user_id
                )
            else:
                # 考试模式，类似选择题的处理...
                # 为简洁起见，这里省略具体实现
                exam_id = data.get('examId')
                section_id = data.get('sectionId')
                
                if not exam_id:
                    return jsonify({'status': 'error', 'message': '请选择考试场次'}), 400
                
                # 获取该考试的最大题目序号
                query_max = """
                    SELECT MAX(question_order) as max_order FROM questions 
                    WHERE session_id = ?
                """
                success, result = execute_query(query_max, (exam_id,))
                if not success:
                    return jsonify({'status': 'error', 'message': '获取题目序号失败'}), 500
                
                order = (result[0]['max_order'] or 0) + 1 if result else 1
                
                # 准备测试用例数据
                test_cases_json = json.dumps(test_cases)
                
                # 首先插入到编程题表
                query_prog = """
                    INSERT INTO programming_questions 
                    (subject, question_text, reference_solution, test_cases, 
                     sample_input, sample_output, time_limit, memory_limit,
                     hints, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 获取考试场次的科目信息
                query_subject = """
                    SELECT subject FROM exam_sessions WHERE id = ?
                """
                success, subject_result = execute_query(query_subject, (exam_id,))
                if not success or not subject_result:
                    return jsonify({'status': 'error', 'message': '获取考试科目信息失败'}), 500
                
                exam_subject = subject_result[0]['subject']
                
                params_prog = (
                    exam_subject, content, answer, test_cases_json,
                    sample_input, sample_output, time_limit, memory_limit,
                    explanation, score, difficulty, user_id, False
                )
                
                print(f"插入考试编程题: {params_prog}")  # 调试输出
                
                success, prog_result = execute_query(query_prog, params_prog)
                if not success:
                    return jsonify({'status': 'error', 'message': '保存编程题失败'}), 500
                
                prog_id = prog_result
                
                # 然后插入到考试题目表
                query = """
                    INSERT INTO questions 
                    (session_id, question_type, question_text, score, question_order, 
                     source_question_id, source_table)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                
                params = (
                    exam_id, 
                    'programming', 
                    content, score, order, prog_id, 'programming_questions'
                )
        elif question_type == 'blank':
            # 填空题处理
            answer = data.get('answer')
            
            # 调试信息
            print(f"处理填空题: data={data}")
            print(f"answer={answer}, 类型={type(answer)}")
            
            # 确保答案有效
            if not answer:
                return jsonify({'status': 'error', 'message': '请输入正确答案'}), 400
            
            # 转换分值和难度为整数
            try:
                score = int(float(score))
                difficulty = int(difficulty)
            except (ValueError, TypeError):
                return jsonify({'status': 'error', 'message': '分值和难度等级必须是数字'}), 400
            
            # 处理可选的替代答案
            alternative_answers = data.get('alternativeAnswers', '')
            
            # 创建填空题
            if mode == 'question-bank':
                # 题库模式，插入到填空题表
                query = """
                    INSERT INTO fill_blank_questions 
                    (subject, question_text, correct_answer, alternative_answers, explanation, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                params = (
                    subject, content, 
                    answer, alternative_answers, explanation, score, difficulty, user_id, True
                )
                
                print(f"插入填空题: {params}")  # 调试输出
            else:
                # 考试模式，插入到考试题目表
                exam_id = data.get('examId')
                section_id = data.get('sectionId')
                
                if not exam_id:
                    return jsonify({'status': 'error', 'message': '请选择考试场次'}), 400
                
                # 获取该考试的最大题目序号
                query_max = """
                    SELECT MAX(question_order) as max_order FROM questions 
                    WHERE session_id = ?
                """
                success, result = execute_query(query_max, (exam_id,))
                if not success:
                    return jsonify({'status': 'error', 'message': '获取题目序号失败'}), 500
                
                order = (result[0]['max_order'] or 0) + 1 if result else 1
                
                # 首先插入到填空题表
                query_fb = """
                    INSERT INTO fill_blank_questions 
                    (subject, question_text, correct_answer, alternative_answers, explanation, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 获取考试场次的科目信息
                query_subject = """
                    SELECT subject FROM exam_sessions WHERE id = ?
                """
                success, subject_result = execute_query(query_subject, (exam_id,))
                if not success or not subject_result:
                    return jsonify({'status': 'error', 'message': '获取考试科目信息失败'}), 500
                
                exam_subject = subject_result[0]['subject']
                
                params_fb = (
                    exam_subject, content, 
                    answer, alternative_answers, explanation, score, difficulty, user_id, False
                )
                
                print(f"插入考试填空题: {params_fb}")  # 调试输出
                
                success, fb_result = execute_query(query_fb, params_fb)
                if not success:
                    return jsonify({'status': 'error', 'message': '保存填空题失败'}), 500
                
                fb_id = fb_result
                
                # 然后插入到考试题目表
                query = """
                    INSERT INTO questions 
                    (session_id, question_type, question_text, score, question_order, 
                     source_question_id, source_table)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                
                params = (
                    exam_id, 
                    'fill_blank', 
                    content, score, order, fb_id, 'fill_blank_questions'
                )
        elif question_type == 'essay':
            # 简答题处理
            answer = data.get('answer')
            
            # 调试信息
            print(f"处理简答题: data={data}")
            print(f"answer={answer}, 类型={type(answer)}")
            
            # 确保答案有效
            if not answer:
                return jsonify({'status': 'error', 'message': '请输入参考答案'}), 400
            
            # 转换分值和难度为整数
            try:
                score = int(float(score))
                difficulty = int(difficulty)
            except (ValueError, TypeError):
                return jsonify({'status': 'error', 'message': '分值和难度等级必须是数字'}), 400
            
            if mode == 'question-bank':
                # 题库模式，插入到简答题表
                query = """
                    INSERT INTO short_answer_questions 
                    (subject, question_text, reference_answer, key_points, grading_criteria, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 简单处理，将答案的前30个字符作为关键点
                key_points = answer[:30] + "..." if len(answer) > 30 else answer
                
                params = (
                    subject, content, 
                    answer, key_points, explanation, score, difficulty, user_id, True
                )
                
                print(f"插入简答题: {params}")  # 调试输出
            else:
                # 考试模式，插入到考试题目表
                exam_id = data.get('examId')
                section_id = data.get('sectionId')
                
                if not exam_id:
                    return jsonify({'status': 'error', 'message': '请选择考试场次'}), 400
                
                # 获取该考试的最大题目序号
                query_max = """
                    SELECT MAX(question_order) as max_order FROM questions 
                    WHERE session_id = ?
                """
                success, result = execute_query(query_max, (exam_id,))
                if not success:
                    return jsonify({'status': 'error', 'message': '获取题目序号失败'}), 500
                
                order = (result[0]['max_order'] or 0) + 1 if result else 1
                
                # 首先插入到简答题表
                query_sa = """
                    INSERT INTO short_answer_questions 
                    (subject, question_text, reference_answer, key_points, grading_criteria, score, difficulty, created_by, is_template)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                
                # 简单处理，将答案的前30个字符作为关键点
                key_points = answer[:30] + "..." if len(answer) > 30 else answer
                
                # 获取考试场次的科目信息
                query_subject = """
                    SELECT subject FROM exam_sessions WHERE id = ?
                """
                success, subject_result = execute_query(query_subject, (exam_id,))
                if not success or not subject_result:
                    return jsonify({'status': 'error', 'message': '获取考试科目信息失败'}), 500
                
                exam_subject = subject_result[0]['subject']
                
                params_sa = (
                    exam_subject, content, 
                    answer, key_points, explanation, score, difficulty, user_id, False
                )
                
                print(f"插入考试简答题: {params_sa}")  # 调试输出
                
                success, sa_result = execute_query(query_sa, params_sa)
                if not success:
                    return jsonify({'status': 'error', 'message': '保存简答题失败'}), 500
                
                sa_id = sa_result
                
                # 然后插入到考试题目表
                query = """
                    INSERT INTO questions 
                    (session_id, question_type, question_text, score, question_order, 
                     source_question_id, source_table)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                
                params = (
                    exam_id, 
                    'short_answer', 
                    content, score, order, sa_id, 'short_answer_questions'
                )
        else:
            # 其他题型
            pass
            
        # 执行数据库操作
        try:
            success, result = execute_query(query, params)
            if not success:
                print(f"保存题目失败: query={query}, params={params}")
                return jsonify({'status': 'error', 'message': '保存题目失败'}), 500
            
            return jsonify({'status': 'success', 'message': '题目保存成功'})
        except Exception as e:
            print(f"添加题目失败: {str(e)}")
            print(f"失败的SQL: query={query}, params={params}")
            return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500
    
    except Exception as e:
        print(f"添加题目失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/questions', methods=['GET'])
@login_required
def get_questions():
    """获取题目列表的API端点"""
    try:
        # 获取查询参数
        question_type = request.args.get('type', 'all')
        subject = request.args.get('subject')
        
        # 构建查询条件
        params = []
        where_clauses = []
        
        if question_type and question_type != 'all':
            if question_type == 'single' or question_type == 'multiple':
                table = 'multiple_choice_questions'
            elif question_type == 'truefalse':
                table = 'true_false_questions'  # 使用判断题表
            elif question_type == 'programming':
                table = 'programming_questions'
            elif question_type == 'blank':
                table = 'fill_blank_questions'
            elif question_type == 'essay':
                table = 'short_answer_questions'
            else:
                table = 'multiple_choice_questions'
        else:
            # 默认返回选择题
            table = 'multiple_choice_questions'
        
        if subject:
            where_clauses.append("subject LIKE ?")
            params.append(f'%{subject}%')
        
        # 只返回题库模板
        where_clauses.append("is_template = 1")
        
        # 构建查询语句
        query = f"SELECT * FROM {table}"
        if where_clauses:
            query += " WHERE " + " AND ".join(where_clauses)
        
        query += " ORDER BY created_at DESC LIMIT 100"
        
        # 执行查询
        success, questions = execute_query(query, params)
        if not success:
            return jsonify({'status': 'error', 'message': '获取题目失败'}), 500
        
        # 格式化结果
        result = []
        for q in questions or []:
            question_data = {
                'id': q['id'],
                'content': q['question_text'],
                'difficulty': q['difficulty'],
                'score': q['score'],
            }
            
            # 根据题目类型添加特定字段
            if table == 'multiple_choice_questions':
                question_data['options'] = [
                    q['option_a'], q['option_b'], 
                    q['option_c'], q['option_d']
                ]
                # 数据库中存储的是A/B/C/D格式的答案，直接传递给前端
                question_data['answer'] = q['correct_answer']
                
                # 注意：前端展示时使用字母，但提交时使用索引值(0,1,2,3)
                # 这里可以添加转换以匹配前端的索引，例如：
                # 如果需要将A转为0，B转为1等，可以使用如下代码：
                # if q['correct_answer'] and len(q['correct_answer']) == 1:
                #     letter = q['correct_answer'][0]
                #     if 'A' <= letter <= 'Z':
                #         # 将A转为0，B转为1，以此类推
                #         question_data['answer_index'] = ord(letter) - ord('A')
            elif table == 'true_false_questions':
                # 判断题只有两个固定选项：正确和错误
                question_data['options'] = ["A. 正确", "B. 错误"]
                # 0表示正确，1表示错误
                question_data['answer'] = str(q['correct_answer'])
            elif table == 'programming_questions':
                question_data['testCases'] = json.loads(q['test_cases']) if q['test_cases'] else []
                question_data['answer'] = q['reference_solution']
            elif table == 'fill_blank_questions':
                question_data['answer'] = q['correct_answer']
                question_data['alternativeAnswers'] = q['alternative_answers'] or ''
            
            # 添加通用字段
            question_data['explanation'] = q.get('explanation', '')
            question_data['tags'] = q['subject'].split(',') if q['subject'] else []
            
            result.append(question_data)
        
        return jsonify({'status': 'success', 'data': result})
    
    except Exception as e:
        print(f"获取题目失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/exams', methods=['GET'])
def get_exams():
    # 获取当前用户ID
    user_id = get_current_user_id()
    """获取考试场次列表的API端点"""
    try:
        print("============ 获取考试场次 ============")
        
        # 查询所有考试场次
        query = """
            SELECT id, name FROM exam_sessions 
            WHERE created_by = ?
            ORDER BY created_at DESC
        """
        print(f"执行查询: {query}")
        success, exams = execute_query(query, (user_id,))
        print(f"查询结果: success={success}, exams={exams}")
        
        if not success:
            print("获取考试场次失败")
            return jsonify({'error': '获取考试场次失败'}), 500
        
        # 确保返回的是列表
        result = [dict(exam) for exam in (exams or [])]
        print(f"返回数据: {result}")
        
        # 无论是否有数据，都返回JSON数组
        return jsonify(result)
    except Exception as e:
        print(f"获取考试场次失败: {str(e)}")
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

@app.route('/api/db/table_info/<table_name>', methods=['GET'])
@login_required
def get_table_info(table_name):
    """获取指定表的结构信息"""
    try:
        # 验证表名防止SQL注入
        allowed_tables = [
            'multiple_choice_questions', 
            'true_false_questions',
            'fill_blank_questions',
            'short_answer_questions',
            'programming_questions',
            'questions',
            'exam_sessions'
        ]
        
        if table_name not in allowed_tables:
            return jsonify({'error': '不允许的表名'}), 400
        
        # 查询表结构
        query = f"PRAGMA table_info({table_name})"
        success, columns = execute_query(query)
        
        if not success:
            return jsonify({'error': '获取表结构失败'}), 500
        
        # 查询表数据示例
        query_data = f"SELECT * FROM {table_name} LIMIT 5"
        success, data = execute_query(query_data)
        
        if not success:
            return jsonify({'error': '获取表数据失败'}), 500
        
        return jsonify({
            'table': table_name,
            'columns': columns,
            'sample_data': data
        })
    
    except Exception as e:
        print(f"获取表信息失败: {str(e)}")
        return jsonify({'error': f'服务器错误: {str(e)}'}), 500

@app.route('/api/db/query/true_false_questions', methods=['GET'])
@login_required
def get_truefalse_questions():
    """获取判断题列表"""
    try:
        query = "SELECT * FROM true_false_questions WHERE is_template = 1 ORDER BY created_at DESC LIMIT 100"
        success, questions = execute_query(query)
        
        if not success:
            return jsonify({'status': 'error', 'message': '获取判断题失败'}), 500
        
        return jsonify({'status': 'success', 'data': questions or []})
    except Exception as e:
        print(f"获取判断题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/true_false_questions/<int:question_id>', methods=['GET'])
@login_required
def get_truefalse_question(question_id):
    """获取单个判断题的详细信息"""
    try:
        query = "SELECT * FROM true_false_questions WHERE id = ?"
        success, result = execute_query(query, (question_id,))
        
        if not success or not result:
            return jsonify({'status': 'error', 'message': '获取判断题失败或判断题不存在'}), 404
        
        question = result[0]
        
        # 前端编码: 0="正确", 1="错误"
        # SQLite中布尔值: 0=False, 1=True
        # 数据库->前端: 1(True)->0(正确), 0(False)->1(错误)
        correct_answer_value = '0' if question['correct_answer'] else '1'
        
        # 格式化结果
        question_data = {
            'id': question['id'],
            'subject': question['subject'],
            'content': question['question_text'],
            'correct_answer': correct_answer_value,  # 转换为前端期望的格式
            'explanation': question['explanation'],
            'score': question['score'],
            'difficulty': question['difficulty'],
            'options': ["A. 正确", "B. 错误"]  # 判断题固定选项
        }
        
        return jsonify({'status': 'success', 'data': question_data})
    
    except Exception as e:
        print(f"获取判断题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/true_false_questions/<int:question_id>', methods=['PUT'])
@login_required
def update_truefalse_question(question_id):
    """更新判断题"""
    try:
        data = request.get_json()
        print(f"更新判断题数据: {data}")  # 调试输出
        
        # 获取要更新的数据
        subject = data.get('subject', '未分类')
        content = data.get('content')
        correct_answer = data.get('correct_answer')
        explanation = data.get('explanation', '')
        score = data.get('score')
        difficulty = data.get('difficulty')
        
        # 验证必填字段
        if not content:
            return jsonify({'status': 'error', 'message': '题目内容不能为空'}), 400
            
        if correct_answer is None:
            return jsonify({'status': 'error', 'message': '请选择正确答案'}), 400
        
        # 前端编码: 0="正确", 1="错误"
        # SQLite中布尔值: 0=False, 1=True
        # 前端->数据库: 0(正确)->1(True), 1(错误)->0(False)
        answer_bool = (correct_answer == '0')
        print(f"前端答案值: {correct_answer}, 转换为布尔值: {answer_bool}")  # 调试输出
        
        # 更新数据库
        query = """
            UPDATE true_false_questions 
            SET subject = ?, question_text = ?, correct_answer = ?, 
                explanation = ?, score = ?, difficulty = ?
            WHERE id = ?
        """
        
        params = (
            subject, content, answer_bool, 
            explanation, score, difficulty, question_id
        )
        
        print(f"执行更新: {query} 参数: {params}")  # 调试输出
        success, _ = execute_query(query, params)
        if not success:
            return jsonify({'status': 'error', 'message': '更新判断题失败'}), 500
        
        return jsonify({'status': 'success', 'message': '判断题更新成功'})
    
    except Exception as e:
        print(f"更新判断题失败: {str(e)}")
        import traceback
        traceback.print_exc()  # 打印完整的错误堆栈
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/true_false_questions/<int:question_id>', methods=['DELETE'])
@login_required
def delete_truefalse_question(question_id):
    """删除判断题"""
    try:
        # 检查题目是否存在
        query_check = "SELECT id FROM true_false_questions WHERE id = ?"
        success, result = execute_query(query_check, (question_id,))
        
        if not success or not result:
            return jsonify({'status': 'error', 'message': '判断题不存在'}), 404
        
        # 删除题目
        query = "DELETE FROM true_false_questions WHERE id = ?"
        success, _ = execute_query(query, (question_id,))
        
        if not success:
            return jsonify({'status': 'error', 'message': '删除判断题失败'}), 500
        
        return jsonify({'status': 'success', 'message': '判断题删除成功'})
    
    except Exception as e:
        print(f"删除判断题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

# 填空题API端点
@app.route('/api/db/query/fill_blank_questions', methods=['GET'])
@login_required
def get_fill_blank_questions():
    """获取填空题列表"""
    try:
        query = "SELECT * FROM fill_blank_questions WHERE is_template = 1 ORDER BY created_at DESC LIMIT 100"
        success, questions = execute_query(query)
        
        if not success:
            return jsonify({'status': 'error', 'message': '获取填空题失败'}), 500
        
        return jsonify({'status': 'success', 'data': questions or []})
    except Exception as e:
        print(f"获取填空题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/fill_blank_questions/<int:question_id>', methods=['GET'])
@login_required
def get_fill_blank_question(question_id):
    """获取单个填空题的详细信息"""
    try:
        query = "SELECT * FROM fill_blank_questions WHERE id = ?"
        success, result = execute_query(query, (question_id,))
        
        if not success or not result:
            return jsonify({'status': 'error', 'message': '获取填空题失败或填空题不存在'}), 404
        
        question = result[0]
        
        # 格式化结果
        question_data = {
            'id': question['id'],
            'subject': question['subject'],
            'content': question['question_text'],
            'correct_answer': question['correct_answer'],
            'alternative_answers': question['alternative_answers'] or '',
            'explanation': question['explanation'],
            'score': question['score'],
            'difficulty': question['difficulty']
        }
        
        return jsonify({'status': 'success', 'data': question_data})
    
    except Exception as e:
        print(f"获取填空题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/fill_blank_questions/<int:question_id>', methods=['PUT'])
@login_required
def update_fill_blank_question(question_id):
    """更新填空题"""
    try:
        data = request.get_json()
        print(f"更新填空题数据: {data}")  # 调试输出
        
        # 获取要更新的数据
        subject = data.get('subject', '未分类')
        content = data.get('content')
        correct_answer = data.get('correct_answer')
        alternative_answers = data.get('alternative_answers', '')
        explanation = data.get('explanation', '')
        score = data.get('score')
        difficulty = data.get('difficulty')
        
        # 验证必填字段
        if not content:
            return jsonify({'status': 'error', 'message': '题目内容不能为空'}), 400
            
        if not correct_answer:
            return jsonify({'status': 'error', 'message': '请填写正确答案'}), 400
        
        # 更新数据库
        query = """
            UPDATE fill_blank_questions 
            SET subject = ?, question_text = ?, correct_answer = ?, alternative_answers = ?,
                explanation = ?, score = ?, difficulty = ?
            WHERE id = ?
        """
        
        params = (
            subject, content, correct_answer, alternative_answers,
            explanation, score, difficulty, question_id
        )
        
        print(f"执行更新: {query} 参数: {params}")  # 调试输出
        success, _ = execute_query(query, params)
        if not success:
            return jsonify({'status': 'error', 'message': '更新填空题失败'}), 500
        
        return jsonify({'status': 'success', 'message': '填空题更新成功'})
    
    except Exception as e:
        print(f"更新填空题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/fill_blank_questions/<int:question_id>', methods=['DELETE'])
@login_required
def delete_fill_blank_question(question_id):
    """删除填空题"""
    try:
        # 检查题目是否存在
        query_check = "SELECT id FROM fill_blank_questions WHERE id = ?"
        success, result = execute_query(query_check, (question_id,))
        
        if not success or not result:
            return jsonify({'status': 'error', 'message': '填空题不存在'}), 404
        
        # 删除题目
        query = "DELETE FROM fill_blank_questions WHERE id = ?"
        success, _ = execute_query(query, (question_id,))
        
        if not success:
            return jsonify({'status': 'error', 'message': '删除填空题失败'}), 500
        
        return jsonify({'status': 'success', 'message': '填空题删除成功'})
    
    except Exception as e:
        print(f"删除填空题失败: {str(e)}")
        return jsonify({'status': 'error', 'message': f'服务器错误: {str(e)}'}), 500

@app.route('/api/upload-image', methods=['POST'])
def upload_image():
    """处理题目编辑器中的图片上传"""
    try:
        # 检查是否有文件
        if 'image' not in request.files:
            return jsonify({'status': 'error', 'message': '没有上传文件'}), 400
        
        file = request.files['image']
        
        # 检查文件名是否为空
        if file.filename == '':
            return jsonify({'status': 'error', 'message': '未选择文件'}), 400
        
        # 明确定义允许的图片扩展名
        IMAGE_ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
        
        # 检查文件类型
        if not allowed_file(file.filename, IMAGE_ALLOWED_EXTENSIONS):
            return jsonify({'status': 'error', 'message': '不支持的文件类型，请上传图片文件'}), 400
        
        # 检查文件大小（限制为2MB）
        content = file.read()
        if len(content) > 2 * 1024 * 1024:
            return jsonify({'status': 'error', 'message': '图片大小不能超过2MB'}), 400
        
        # 重置文件指针
        file.seek(0)
        
        # 生成安全的文件名
        original_filename = secure_filename(file.filename)
        # 使用UUID和时间戳组成唯一文件名
        filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex}_{original_filename}"
        
        # 确保上传目录存在
        os.makedirs(IMAGE_UPLOAD_FOLDER, exist_ok=True)
        
        # 保存文件
        file_path = os.path.join(IMAGE_UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # 构建URL（相对路径）
        image_url = f"/{IMAGE_UPLOAD_FOLDER}/{filename}"
        
        # 记录成功上传的日志
        print(f"图片上传成功: {image_url}")
        
        return jsonify({
            'status': 'success',
            'message': '图片上传成功',
            'imageUrl': image_url
        })
        
    except Exception as e:
        print(f"图片上传出错: {str(e)}")
        return jsonify({'status': 'error', 'message': f'图片上传失败: {str(e)}'}), 500

@app.route('/static/uploads/images/<path:filename>')
def uploaded_image(filename):
    """访问上传的图片"""
    return send_from_directory(IMAGE_UPLOAD_FOLDER, filename)

@app.route('/api/test-upload', methods=['GET'])
def test_upload_api():
    """测试图片上传API是否正常工作"""
    return jsonify({
        'status': 'success',
        'message': '图片上传API可用',
        'upload_folder': IMAGE_UPLOAD_FOLDER,
        'allowed_extensions': list(ALLOWED_EXTENSIONS),
        'server_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    })

@app.route('/api/subjects', methods=['GET'])
@login_required
def get_subjects():
    """获取所有科目列表"""
    query = """
        SELECT DISTINCT subject FROM exam_sessions
        ORDER BY subject
    """
    
    success, subjects = execute_query(query)
    
    if not success:
        return jsonify([]), 500
    
    # 提取科目列表
    subject_list = [subject['subject'] for subject in (subjects or [])]
    
    return jsonify(subject_list)

@app.route('/api/camera-mode-grading', methods=['POST'])
@login_required
def camera_mode_grading():
    """处理自动拍摄模式的阅卷请求"""
    try:
        # 确保存储目录存在
        EXTRACT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'extracted_files')
        if not os.path.exists(EXTRACT_FOLDER):
            os.makedirs(EXTRACT_FOLDER)
            print(f"创建目录: {EXTRACT_FOLDER}")
        # 获取考试场次ID
        session_id = request.form.get('session_id')
        if not session_id:
            return jsonify({'success': False, 'message': '缺少考试场次ID'}), 400

        # 获取PDF文件
        if 'pdf_file' not in request.files:
            return jsonify({'success': False, 'message': '未上传PDF文件'}), 400

        pdf_file = request.files['pdf_file']
        if not pdf_file or not pdf_file.filename:
            return jsonify({'success': False, 'message': 'PDF文件无效'}), 400
        print(f"接收到文件: {request.files['pdf_file'].filename}")

        # 保存PDF文件
        filename = secure_filename(pdf_file.filename)
        save_path = os.path.join(EXTRACT_FOLDER, filename)
        pdf_file.save(save_path)
        print(f"保存文件到: {save_path}")

        # 处理PDF文件
        try:
            # 1. 使用OCR处理PDF
            from split_and_ocr.pdf_ocr import process_pdf
            ocr_output = os.path.join(EXTRACT_FOLDER, f"ocr_results_{filename}.txt")
            process_pdf(save_path, ocr_output)
            with open(ocr_output, 'r', encoding='utf-8-sig') as f:
                res = f.read()
            f.close()
            print(res)
            # # 2. 分割题目和答案
            # from split_and_ocr.read.questionsplit import readexit
            # readexit()

            # 3. 获取考试信息
            query = """
                SELECT id, name, subject, exam_score
                FROM exam_sessions
                WHERE id = ?
            """
            success, exam_info = execute_query(query, (session_id,))
            if not success or not exam_info:
                raise Exception('获取考试信息失败')

            exam_info = exam_info[0]

            # # 4. AI评分处理
            # airead()

            return jsonify({
                'success': True,
                'message': '阅卷完成'
            })

        finally:
            # 清理临时文件
            if os.path.exists(save_path):
                os.remove(save_path)

    except Exception as e:
        print(f"阅卷过程出错: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'阅卷过程出错: {str(e)}'
        }), 500
    
@app.route('/api/exam/update-status/<int:exam_id>', methods=['POST'])
@login_required
def update_exam_status(exam_id):
    """更新考试状态"""
    try:
        # 检查权限，只有管理员可以更新考试状态
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'success': False, 'error': '未登录'}), 401
        
        # 检查用户权限
        role_query = "SELECT role FROM users WHERE id = ?"
        success, role_result = execute_query(role_query, (user_id,))
        if not success or not role_result or role_result[0]['role'] != 'admin':
            return jsonify({'success': False, 'error': '权限不足'}), 403
        
        # 获取要更新的状态
        data = request.get_json()
        new_status = data.get('status')
        
        # 验证状态是否有效
        if new_status not in ['pending', 'ongoing', 'completed', 'graded']:
            return jsonify({'success': False, 'error': '无效的状态值'}), 400
        
        # 更新考试状态
        update_query = """
            UPDATE exam_sessions
            SET status = ?
            WHERE id = ?
        """
        
        success, _ = execute_query(update_query, (new_status, exam_id))
        if not success:
            return jsonify({'success': False, 'error': '更新考试状态失败'}), 500
        
        # 如果更新状态为 graded，可能需要确保所有试卷都已评分
        if new_status == 'graded':
            # 检查是否有未评分的答案
            check_query = """
                SELECT COUNT(*) as ungraded_count
                FROM student_answers
                WHERE session_id = ? AND ai_score IS NULL
            """
            success, results = execute_query(check_query, (exam_id,))
            
            if success and results and results[0]['ungraded_count'] > 0:
                return jsonify({
                    'success': True, 
                    'warning': f'状态已更新为已评分，但仍有 {results[0]["ungraded_count"]} 个答案未评分'
                })
        
        return jsonify({'success': True, 'message': '考试状态更新成功'})
    except Exception as e:
        print(f"更新考试状态时出错: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# 获取考试场次列表
@app.route('/api/analysis/sessions')
@login_required
def get_analysis_sessions():
    """获取所有已完成的考试场次"""
    print("============== 开始获取分析考试场次 ==============")
    
    # 首先检查数据库中考试场次状态情况
    status_query = """
        SELECT status, COUNT(*) as count
        FROM exam_sessions 
        GROUP BY status
    """
    status_success, status_results = execute_query(status_query)
    
    if status_success and status_results:
        print("数据库中的考试场次状态统计:")
        for status_info in status_results:
            print(f"状态: {status_info['status']} - 数量: {status_info['count']}")
    else:
        print("获取考试状态统计失败或没有考试记录")
    
    # 修改查询，直接查询所有记录，不限制状态，方便调试
    all_query = """
        SELECT 
            id,
            name,
            subject,
            start_time,
            end_time,
            status
        FROM exam_sessions 
        ORDER BY start_time DESC
        LIMIT 10
    """
    
    print(f"执行所有考试场次查询: {all_query}")
    all_success, all_sessions = execute_query(all_query)
    
    if all_success and all_sessions:
        print("数据库中的所有考试记录:")
        for idx, session in enumerate(all_sessions):
            print(f"记录 {idx+1}: ID={session['id']}, 名称={session['name']}, 状态={session['status']}")
    
    # 查询所有的已评分考试记录
    query = """
        SELECT 
            id,
            name,
            subject,
            start_time,
            end_time,
            status
        FROM exam_sessions 
        ORDER BY start_time DESC
    """
    
    print(f"执行考试场次查询: {query}")
    success, sessions = execute_query(query)
    print(f"查询结果: success={success}, 找到记录数={len(sessions) if sessions else 0}")
    
    if sessions:
        print("找到的考试记录:")
        for idx, session in enumerate(sessions):
            print(f"记录 {idx+1}: ID={session['id']}, 名称={session['name']}, 状态={session['status']}")
    
    if not success:
        print("获取考试场次失败")
        return jsonify([])
    
    # 处理日期格式    
    result = []
    for session in (sessions or []):
        try:
            # 处理日期时间字段
            if session['start_time']:
                if '.' in session['start_time']:
                    start_time = datetime.strptime(session['start_time'], '%Y-%m-%d %H:%M:%S.%f')
                else:
                    start_time = datetime.strptime(session['start_time'], '%Y-%m-%d %H:%M:%S')
                
                formatted_start = start_time.strftime('%Y-%m-%d %H:%M')
            else:
                formatted_start = ''
                
            if session['end_time']:
                if '.' in session['end_time']:
                    end_time = datetime.strptime(session['end_time'], '%Y-%m-%d %H:%M:%S.%f')
                else:
                    end_time = datetime.strptime(session['end_time'], '%Y-%m-%d %H:%M:%S')
                
                formatted_end = end_time.strftime('%Y-%m-%d %H:%M')
            else:
                formatted_end = ''
            
            result.append({
                'id': session['id'],
                'name': session['name'],
                'subject': session['subject'],
                'start_time': formatted_start,
                'end_time': formatted_end,
                'status': session['status']
            })
        except Exception as e:
            print(f"处理考试场次数据时出错: {str(e)}, 数据: {session}")
    
    print(f"返回数据: {result}")
    print("============== 结束获取分析考试场次 ==============")
    return jsonify(result)


# 考试查询AI Assistant API接口
@app.route('/api/agent/exam-query', methods=['POST'])
@login_required
def exam_query():
    data = request.json
    query = data.get('query', '')
    
    if not query:
        return jsonify({'response': '请输入您的问题'})
    
    try:
        # 获取所有考试数据
        sql = """
        SELECT name, subject, start_time, end_time, status
        FROM exam_sessions
        ORDER BY start_time
        """
        success, exams = execute_query(sql)
        
        if not success or not exams:
            return jsonify({'response': '当前没有考试安排信息。'})
        
        # 构建AI提示
        exams_info = "\n".join([
            f"考试名称：{exam['name']}, 科目：{exam['subject']}, "
            f"开始时间：{exam['start_time']}, 结束时间：{exam['end_time']}, "
            f"状态：{exam['status']}"
            for exam in exams
        ])
        
        prompt = f"""
        作为考试查询助手，回答用户关于考试安排的问题。
        以下是当前系统中的所有考试安排：
        
        {exams_info}
        
        用户问题: {query}
        
        请提供简洁明了的回答，如果问题是关于特定日期的考试，请列出该日期的所有考试。
        如果没有考试，请明确告知。回答要简短友好。
        """
        
        # 调用AI接口获取回答
        response = new("你是一个考试查询助手", prompt)
        
        return jsonify({'response': response})
    except Exception as e:
        app.logger.error(f"查询考试出错: {str(e)}")
        return jsonify({'response': f'抱歉，查询时发生错误：{str(e)}'})
    
# 处理创建考试的自然语言请求
@app.route('/api/agent/process-exam-request', methods=['POST'])
@login_required
def process_exam_request():
    data = request.json
    query = data.get('query', '')
    
    if not query:
        return jsonify({'success': False, 'response': '请提供有效的请求'})
    
    try:
        # 获取当前用户ID
        user_id = get_current_user_id()
        
        # 使用AI解析请求中的考试信息
        prompt = f"""
        请从以下用户请求中提取考试信息，并按JSON格式返回：
        
        用户请求: {query}
        
        需要提取的信息包括：
        1. 考试名称（默认为"[科目]考试"）
        2. 科目（例如：Java、Python等）
        3. 考试日期（例如：2023-04-04）
        4. 开始时间（例如：16:00）
        5. 结束时间或持续时间
        
        请以下面的JSON格式返回结果：
        {{
          "exam_name": "考试名称",
          "subject": "科目名称",
          "date": "YYYY-MM-DD",
          "start_time": "HH:MM",
          "end_time": "HH:MM",
          "duration_minutes": 120
        }}
        
        如果无法确定某项信息，请将该字段设为null。请只返回解析结果的JSON，不要包含其他解释文本。
        """
        
        # 调用AI解析
        ai_response = new("你是一个考试信息提取助手", prompt)
        
        # 尝试解析AI返回的JSON
        try:
            exam_info = json.loads(ai_response)
        except json.JSONDecodeError:
            # 如果AI没有返回有效JSON，尝试从文本中提取JSON部分
            json_match = re.search(r'\{[\s\S]*\}', ai_response)
            if json_match:
                try:
                    exam_info = json.loads(json_match.group(0))
                except:
                    return jsonify({
                        'success': False, 
                        'response': '抱歉，我无法理解您的考试信息。请提供更明确的日期、时间和科目信息。'
                    })
            else:
                return jsonify({
                    'success': False, 
                    'response': '抱歉，我无法理解您的考试信息。请提供更明确的日期、时间和科目信息。'
                })
        
        # 验证必要的信息是否存在
        if not exam_info.get('subject'):
            return jsonify({
                'success': False, 
                'response': '请提供考试科目信息，例如："Java考试"或"Python测试"'
            })
            
        if not exam_info.get('date'):
            return jsonify({
                'success': False, 
                'response': '请提供考试日期，例如："4月4日"或"2023年4月4日"'
            })
            
        if not exam_info.get('start_time'):
            return jsonify({
                'success': False, 
                'response': '请提供考试开始时间，例如："下午3点"或"15:00"'
            })
        
        # 确保有考试名称，如果没有则使用科目+考试
        if not exam_info.get('exam_name'):
            exam_info['exam_name'] = f"{exam_info['subject']}考试"
            
        # 计算考试时长
        duration_minutes = exam_info.get('duration_minutes', 120)  # 默认2小时
        
        # 如果有结束时间但没有时长，计算时长
        if exam_info.get('end_time') and not exam_info.get('duration_minutes'):
            try:
                start_dt = datetime.strptime(f"{exam_info['date']} {exam_info['start_time']}", "%Y-%m-%d %H:%M")
                end_dt = datetime.strptime(f"{exam_info['date']} {exam_info['end_time']}", "%Y-%m-%d %H:%M")
                
                # 如果结束时间早于开始时间，可能是跨天
                if end_dt < start_dt:
                    end_dt = end_dt + timedelta(days=1)
                    
                duration_minutes = int((end_dt - start_dt).total_seconds() / 60)
            except:
                duration_minutes = 120  # 解析失败使用默认值
        
        # 生成开始和结束时间
        try:
            start_datetime = datetime.strptime(f"{exam_info['date']} {exam_info['start_time']}", "%Y-%m-%d %H:%M")
            end_datetime = start_datetime + timedelta(minutes=duration_minutes)
        except:
            return jsonify({
                'success': False, 
                'response': '抱歉，日期或时间格式有误。请使用标准格式，如"2023-04-04 16:00"。'
            })
        
        # 插入到数据库 - 不创建PDF文件，将文件路径设为空字符串
        query = """
            INSERT INTO exam_sessions 
            (name, subject, start_time, end_time, duration, exam_file_path, exam_score, status, created_by)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """
        
        success, result = execute_query(
            query, 
            (
                exam_info['exam_name'], 
                exam_info['subject'], 
                start_datetime.strftime('%Y-%m-%d %H:%M:%S'),
                end_datetime.strftime('%Y-%m-%d %H:%M:%S'),
                duration_minutes,
                "",  # 空文件路径
                100,  # 默认总分100
                'pending',
                user_id
            )
        )
        
        if not success:
            return jsonify({
                'success': False, 
                'response': '创建考试失败，请稍后再试。'
            })
        
        # 记录操作日志
        log_query = """
            INSERT INTO operation_logs (user_id, operation_type, operation_details)
            VALUES (?, ?, ?)
        """
        execute_query(
            log_query, 
            (user_id, 'add_exam', f'通过AI助手添加考试: {exam_info["exam_name"]}, 科目: {exam_info["subject"]}')
        )
        
        # 构建成功响应
        formatted_start = start_datetime.strftime('%Y年%m月%d日 %H:%M')
        formatted_end = end_datetime.strftime('%Y年%m月%d日 %H:%M')
        
        response_message = f"""
        已为您成功创建考试！

        考试名称：{exam_info['exam_name']}
        科目：{exam_info['subject']}
        开始时间：{formatted_start}
        结束时间：{formatted_end}
        考试时长：{duration_minutes}分钟

        考试状态已设置为"待开始"，您可以在考试场次中查看并管理这场考试。
        您可以稍后在考试管理页面上传试卷。
        """
        
        return jsonify({
            'success': True,
            'response': response_message,
            'exam_id': result,
            'exam_name': exam_info['exam_name']
        })
        
    except Exception as e:
        app.logger.error(f"处理考试创建请求失败: {str(e)}")
        return jsonify({
            'success': False, 
            'response': f'处理您的请求时出现错误: {str(e)}'
        })

# 处理考试PDF上传的API
@app.route('/api/agent/upload-exam-pdf', methods=['POST'])
@login_required
def upload_exam_pdf():
    try:
        # 获取考试ID
        exam_id = request.form.get('exam_id')
        if not exam_id:
            return jsonify({'success': False, 'message': '缺少考试ID'}), 400
        
        # 获取PDF文件
        pdf_file = request.files.get('pdf_file')
        if not pdf_file:
            return jsonify({'success': False, 'message': '未上传PDF文件'}), 400
        
        # 验证文件类型
        if not pdf_file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'message': '仅支持PDF格式的文件'}), 400
        
        # 获取正确的上传目录
        upload_folder = app.config.get('EXAM_UPLOAD_FOLDER', 'uploads')
        
        # 确保目录存在
        os.makedirs(upload_folder, exist_ok=True)
        
        # 生成安全的文件名并保存
        filename = secure_filename(f"exam_{exam_id}_{int(time.time())}_{pdf_file.filename}")
        file_path = os.path.join(upload_folder, filename)
        pdf_file.save(file_path)
        
        # 更新考试记录中的文件路径
        query = """
            UPDATE exam_sessions
            SET exam_file_path = ?
            WHERE id = ?
        """
        success, _ = execute_query(query, (file_path, exam_id))
        
        if not success:
            return jsonify({'success': False, 'message': '更新考试文件路径失败'}), 500
        
        # 记录操作日志
        user_id = get_current_user_id()
        log_query = """
            INSERT INTO operation_logs (user_id, operation_type, operation_details)
            VALUES (?, ?, ?)
        """
        execute_query(
            log_query, 
            (user_id, 'upload_exam_pdf', f'上传考试试卷: 考试ID {exam_id}, 文件 {filename}')
        )
        
        return jsonify({
            'success': True, 
            'message': 'PDF文件上传成功',
            'file_path': file_path
        })
        
    except Exception as e:
        app.logger.error(f"上传PDF文件失败: {str(e)}")
        return jsonify({'success': False, 'message': f'上传失败: {str(e)}'}), 500
    
# AI题目生成API
@app.route('/api/agent/generate-questions', methods=['POST'])
@login_required
def generate_questions():
    data = request.json
    query = data.get('query', '')
    
    if not query:
        return jsonify({'success': False, 'response': '请提供有效的请求'})
    
    try:
        # 获取当前用户ID
        user_id = get_current_user_id()
        
        # 使用AI生成题目
        prompt = f"""
        请根据以下请求生成题目：
        
        {query}
        
        请按照以下JSON格式返回结果，包含多个题目：
        
        {{
            "questions": [
                {{
                "type": "single", // 题目类型: single-单选题, truefalse-判断题, blank-填空题, essay-简答题, programming-编程题
                "content": "题目内容", // 题目描述
                "options": ["选项A", "选项B", "选项C", "选项D"], // 仅单选题需要
                "answer": "0", // 单选题为选项索引(0-3对应A-D), 判断题为true/false, 其他题型为答案文本
                "explanation": "答案解析", // 可选
                "difficulty": 2, // 难度1-3
                "score": 5, // 分值
                "subject": "科目名称" // 可选
                }},
                {{
                "type": "programming",
                "content": "编写一个Python函数...",
                "answer": "def solution(...)...",
                "sample_input": "输入示例",
                "sample_output": "输出示例",
                "test_cases": [...],
                "hints": "提示",
                "explanation": "解析",
                "difficulty": 2,
                "score": 10
                }},
                // 可包含多个题目...
            ]
        }}
        
        生成的题目应该符合以下要求：
        1. 题目内容要具体、清晰，符合专业标准
        2. 选择题必须有正确的选项和干扰项
        3. 答案必须正确且合理
        4. 解析应详细说明为什么答案是正确的
        
        仅返回JSON格式，不要包含其他说明文字。如果不能解析用户的请求，请返回空的questions数组并在response中解释原因。
        """
        
        # 调用AI接口
        ai_response = new("你是一个专业的考试题目生成助手", prompt)
        
        # 尝试解析JSON响应
        try:
            # 首先尝试直接解析整个响应
            result = json.loads(ai_response)
            questions = result.get('questions', [])
        except json.JSONDecodeError:
            # 如果直接解析失败，尝试提取JSON部分
            match = re.search(r'\{[\s\S]*\}', ai_response)
            if match:
                try:
                    result = json.loads(match.group(0))
                    questions = result.get('questions', [])
                except:
                    # 如果仍然解析失败，返回空列表
                    return jsonify({
                        'success': False,
                        'response': '抱歉，我无法生成符合要求的题目。请尝试更明确的描述，例如"10道关于Python异常处理的选择题"。',
                        'questions': []
                    })
            else:
                # 没有找到JSON格式的内容
                return jsonify({
                    'success': False,
                    'response': '抱歉，我无法生成符合要求的题目。请尝试更明确的描述，例如"10道关于Python异常处理的选择题"。',
                    'questions': []
                })
        
        # 验证生成的题目
        valid_questions = []
        for q in questions:
            # 基本验证：确保必要字段存在
            if 'type' in q and 'content' in q and 'answer' in q:
                # 根据题型验证
                if q['type'] == 'single' and 'options' in q and isinstance(q['options'], list):
                    valid_questions.append(q)
                elif q['type'] in ['truefalse', 'blank', 'essay', 'programming']:
                    valid_questions.append(q)
        
        # 构建响应消息
        if valid_questions:
            response_msg = f"已为您生成{len(valid_questions)}道题目。请选择您想要添加到试卷的题目。"
        else:
            response_msg = "抱歉，我无法生成符合要求的题目。请尝试更明确的描述，例如:10道关于Python异常处理的选择题。"
        
        # 记录生成题目的日志
        log_query = """
            INSERT INTO operation_logs (user_id, operation_type, operation_details)
            VALUES (?, ?, ?)
        """
        execute_query(
            log_query, 
            (user_id, 'generate_questions', f'AI生成题目: {len(valid_questions)}道 - 查询: {query}')
        )
        
        return jsonify({
            'success': True,
            'response': response_msg,
            'questions': valid_questions
        })
        
    except Exception as e:
        app.logger.error(f"生成题目失败: {str(e)}")
        return jsonify({
            'success': False,
            'response': f'生成题目时出现错误: {str(e)}',
            'questions': []
        })  
    
# 添加选中题目的API
@app.route('/api/agent/add-selected-questions', methods=['POST'])
@login_required
def add_selected_questions():
    data = request.json
    questions = data.get('questions', [])
    mode = data.get('mode', 'question-bank')  # 默认为题库模式
    exam_id = data.get('examId') if mode == 'exam' else None
    
    if not questions:
        return jsonify({'success': False, 'message': '没有提供题目'}), 400
    
    try:
        user_id = get_current_user_id()
        added_count = 0
        
        for question in questions:
            question_type = question.get('type')
            content = question.get('content')
            answer = question.get('answer')
            explanation = question.get('explanation', '')
            difficulty = int(question.get('difficulty', 2))
            score = float(question.get('score', 5))
            subject = question.get('subject', '未分类')
            
            # 根据题目类型和模式插入到不同的表
            if question_type == 'single':
                options = question.get('options', [])
                if len(options) < 2:
                    continue  # 跳过选项不足的题目
                
                # 将选项格式化为需要的格式
                option_a = options[0] if len(options) > 0 else ''
                option_b = options[1] if len(options) > 1 else ''
                option_c = options[2] if len(options) > 2 else ''
                option_d = options[3] if len(options) > 3 else ''
                
                # 转换答案格式（从索引到A/B/C/D）
                correct_option = chr(65 + int(answer)) if answer.isdigit() and 0 <= int(answer) < len(options) else 'A'
                
                if mode == 'question-bank':
                    # 插入到题库
                    query = """
                        INSERT INTO multiple_choice_questions 
                        (subject, question_text, option_a, option_b, option_c, option_d, 
                            correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, option_a, option_b, option_c, option_d, 
                            correct_option, explanation, score, difficulty, user_id, True)
                    )
                    
                    if success:
                        added_count += 1
                else:
                    # 考试模式 - 先插入选择题表获取ID
                    query = """
                        INSERT INTO multiple_choice_questions 
                        (subject, question_text, option_a, option_b, option_c, option_d, 
                            correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, option_a, option_b, option_c, option_d, 
                            correct_option, explanation, score, difficulty, user_id, False)
                    )
                    
                    # 再将题目关联到考试的questions表
                    if success and source_id:
                        question_order = get_next_question_order(exam_id)
                        query = """
                            INSERT INTO questions 
                            (session_id, question_type, question_text, score, question_order, 
                            source_question_id, source_table)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """
                        success, _ = execute_query(
                            query, 
                            (exam_id, 'multiple_choice', content, score, question_order, 
                            source_id, 'multiple_choice_questions')
                        )
                        
                        if success:
                            added_count += 1

            elif question_type == 'truefalse':
                # 判断题处理
                correct_answer = question.get('answer') in (True, 'true', 'True')
                
                if mode == 'question-bank':
                    query = """
                        INSERT INTO true_false_questions 
                        (subject, question_text, correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, correct_answer, explanation, score, difficulty, user_id, True)
                    )
                    
                    if success:
                        added_count += 1
                else:
                    # 考试模式 - 先插入判断题表获取ID
                    query = """
                        INSERT INTO true_false_questions 
                        (subject, question_text, correct_answer, explanation, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, correct_answer, explanation, score, difficulty, user_id, False)
                    )
                    
                    # 再将题目关联到考试的questions表
                    if success and source_id:
                        question_order = get_next_question_order(exam_id)
                        query = """
                            INSERT INTO questions 
                            (session_id, question_type, question_text, score, question_order, 
                            source_question_id, source_table)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """
                        success, _ = execute_query(
                            query, 
                            (exam_id, 'true_false', content, score, question_order, 
                            source_id, 'true_false_questions')
                        )
                        
                        if success:
                            added_count += 1

            elif question_type == 'blank':
                # 填空题处理
                correct_answer = question.get('answer', '')
                alternative_answers = question.get('alternative_answers', '')
                
                if mode == 'question-bank':
                    query = """
                        INSERT INTO fill_blank_questions 
                        (subject, question_text, correct_answer, alternative_answers, explanation, 
                            score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    try:
                        success, source_id = execute_query(
                            query, 
                            (subject, content, correct_answer, alternative_answers, explanation, 
                                int(score), min(max(1, difficulty), 5), user_id, True)
                        )
                        
                        if success:
                            added_count += 1
                        else:
                            app.logger.error(f"填空题插入失败: subject={subject}, content={content}")
                    except Exception as e:
                        app.logger.error(f"填空题插入异常: {str(e)}")
                else:
                    # 考试模式 - 先插入填空题表获取ID
                    query = """
                        INSERT INTO fill_blank_questions 
                        (subject, question_text, correct_answer, alternative_answers, explanation, 
                            score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    try:
                        success, source_id = execute_query(
                            query, 
                            (subject, content, correct_answer, alternative_answers, explanation, 
                                int(score), min(max(1, difficulty), 5), user_id, True)
                        )
                        
                        if success:
                            added_count += 1
                        else:
                            app.logger.error(f"填空题插入失败: subject={subject}, content={content}")
                    except Exception as e:
                        app.logger.error(f"填空题插入异常: {str(e)}")
                    
                    # 再将题目关联到考试的questions表
                    if success and source_id:
                        question_order = get_next_question_order(exam_id)
                        query = """
                            INSERT INTO questions 
                            (session_id, question_type, question_text, score, question_order, 
                            source_question_id, source_table)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """
                        success, _ = execute_query(
                            query, 
                            (exam_id, 'fill_blank', content, score, question_order, 
                            source_id, 'fill_blank_questions')
                        )
                        
                        if success:
                            added_count += 1

            elif question_type == 'essay':
                # 简答题处理
                reference_answer = question.get('answer', '')
                key_points = question.get('key_points', reference_answer)
                grading_criteria = question.get('grading_criteria', '')
                
                if mode == 'question-bank':
                    query = """
                        INSERT INTO short_answer_questions 
                        (subject, question_text, reference_answer, key_points, grading_criteria,
                            score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, reference_answer, key_points, grading_criteria,
                            score, difficulty, user_id, True)
                    )
                    
                    if success:
                        added_count += 1
                else:
                    # 考试模式 - 先插入简答题表获取ID
                    query = """
                        INSERT INTO short_answer_questions 
                        (subject, question_text, reference_answer, key_points, grading_criteria,
                            score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, reference_answer, key_points, grading_criteria,
                            score, difficulty, user_id, False)
                    )
                    
                    # 再将题目关联到考试的questions表
                    if success and source_id:
                        question_order = get_next_question_order(exam_id)
                        query = """
                            INSERT INTO questions 
                            (session_id, question_type, question_text, score, question_order, 
                            source_question_id, source_table)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """
                        success, _ = execute_query(
                            query, 
                            (exam_id, 'short_answer', content, score, question_order, 
                            source_id, 'short_answer_questions')
                        )
                        
                        if success:
                            added_count += 1

            elif question_type == 'programming':
                # 编程题处理
                app.logger.info(f"编程题数据: {json.dumps(question)}")
                reference_solution = question.get('answer', '')
                sample_input = question.get('sample_input', '')
                sample_output = question.get('sample_output', '')
                test_cases = question.get('test_cases', '[]')
                hints = question.get('hints', '')
                
                # 确保test_cases是JSON字符串
                if isinstance(test_cases, list):
                    test_cases = json.dumps(test_cases)
                
                if mode == 'question-bank':
                    query = """
                        INSERT INTO programming_questions 
                        (subject, question_text, sample_input, sample_output, test_cases,
                            reference_solution, hints, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, sample_input, sample_output, test_cases,
                            reference_solution, hints, score, difficulty, user_id, True)
                    )
                    
                    if success:
                        added_count += 1
                else:
                    # 考试模式 - 先插入编程题表获取ID
                    query = """
                        INSERT INTO programming_questions 
                        (subject, question_text, sample_input, sample_output, test_cases,
                            reference_solution, hints, score, difficulty, created_by, is_template)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    success, source_id = execute_query(
                        query, 
                        (subject, content, sample_input, sample_output, test_cases,
                            reference_solution, hints, score, difficulty, user_id, False)
                    )
                    
                    # 再将题目关联到考试的questions表
                    if success and source_id:
                        question_order = get_next_question_order(exam_id)
                        query = """
                            INSERT INTO questions 
                            (session_id, question_type, question_text, score, question_order, 
                            source_question_id, source_table)
                            VALUES (?, ?, ?, ?, ?, ?, ?)
                        """
                        success, _ = execute_query(
                            query, 
                            (exam_id, 'programming', content, score, question_order, 
                            source_id, 'programming_questions')
                        )
                        
                        if success:
                            added_count += 1
        
        # 记录操作日志
        log_query = """
            INSERT INTO operation_logs (user_id, operation_type, operation_details)
            VALUES (?, ?, ?)
        """
        execute_query(
            log_query, 
            (user_id, 'add_ai_questions', f'添加AI生成题目: {added_count}题到{mode}')
        )
        
        return jsonify({
            'success': True,
            'message': f'成功添加{added_count}个题目',
            'addedCount': added_count
        })
        
    except Exception as e:
        app.logger.error(f"添加题目失败: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'添加题目失败: {str(e)}'
        }), 500

# 获取下一个题目顺序号的辅助函数
def get_next_question_order(session_id):
    query = """
        SELECT MAX(question_order) as max_order
        FROM questions
        WHERE session_id = ?
    """
    success, result = execute_query(query, (session_id,))
    if success and result and result[0]['max_order'] is not None:
        return result[0]['max_order'] + 1
    return 1



@app.route('/api/agent/review-assistant', methods=['POST'])
def review_assistant():
    try:
        data = request.json
        user_message = data.get('user_message', '')
        session_id = data.get('session_id')
        student_id = data.get('student_id')
        question_data = data.get('question_data', [])
        
        if not user_message or not session_id or not student_id or not question_data:
            return jsonify({"error": "缺少必要参数"}), 400
        
        # 获取考试信息
        exam_info = get_exam_info(session_id)
        
        # 获取学生信息
        student_info = get_student_info(student_id)
        
        # 构建系统提示信息
        system_prompt = construct_system_prompt(exam_info, student_info, question_data)
        
        # 记录交互到日志
        app.logger.info(f"AI助手请求: 考试ID={session_id}, 学生ID={student_id}, 问题='{user_message}'")
        
        # 调用AI生成回答
        ai_response = new(system_prompt, user_message)
        
        # 记录交互历史
        try:
            log_interaction(session_id, student_id, user_message, ai_response)
        except Exception as log_error:
            app.logger.error(f"记录交互历史失败: {str(log_error)}")
        
        return jsonify({"response": ai_response})
    
    except Exception as e:
        app.logger.error(f"AI助手错误: {str(e)}")
        return jsonify({"error": f"处理请求时出错: {str(e)}"}), 500

# 修改get_exam_info函数
def get_exam_info(session_id):
    """获取考试信息"""
    query = """
        SELECT name, subject, total_score, duration_minutes
        FROM exam_sessions
        WHERE id = ?
    """
    success, result = execute_query(query, (session_id,))
    if success and result and len(result) > 0:
        return result[0]
    return {"name": "未知考试", "subject": "未知科目"}

# 修改get_student_info函数
def get_student_info(student_id):
    """获取学生信息"""
    query = """
        SELECT name, student_id as student_code, class
        FROM students
        WHERE id = ?
    """
    success, result = execute_query(query, (student_id,))
    if success and result and len(result) > 0:
        return result[0]
    return {"name": "未知学生", "student_code": "未知学号"}

# 修改log_interaction函数结尾
def log_interaction(session_id, student_id, user_message, ai_response):
    """记录交互历史"""
    try:
        query = """
            INSERT INTO ai_review_interactions
            (session_id, student_id, user_message, ai_response, created_at)
            VALUES (?, ?, ?, ?, datetime('now'))
        """
        success, _ = execute_query(query, (session_id, student_id, user_message, ai_response))
        return success
    except Exception as e:
        app.logger.error(f"记录交互历史失败: {str(e)}")
        return False

def construct_system_prompt(exam_info, student_info, question_data):
    """构建系统提示信息"""
    # 基础系统信息
    system_prompt = f"""你是一个专业的教育评估助手，帮助教师评审学生答案和提供指导。
当前考试: {exam_info.get('name', '未知考试')}
科目: {exam_info.get('subject', '未知科目')}
学生: {student_info.get('name', '未知学生')} (学号: {student_info.get('student_code', '未知学号')})

你将分析以下题目和学生答案。你需要保持客观公正，根据教育标准提供专业建议。
如果用户询问特定题目的解析，请详细解释题目的解题思路和关键知识点。
如果用户询问学生答案的评分，请分析学生答案的优缺点，并提出合理的评分建议。

题目信息:"""

    # 添加题目信息
    for q in question_data:
        system_prompt += f"""
题目 {q.get('order', '?')}: {q.get('text', '未知题目')}
满分: {q.get('totalScore', 0)}分
学生答案: {q.get('studentAnswer', '无答案')}
AI评分: {q.get('aiScore', 0)}分
AI反馈: {q.get('aiFeedback', '无反馈')}
"""
        if q.get('manualFeedback'):
            system_prompt += f"人工反馈: {q.get('manualFeedback')}\n"
    
    return system_prompt

@app.route('/api/check')
def api_check():
    return jsonify({"status": "ok"})

@app.route('/api/review/submit-all', methods=['POST'])
def submit_all_reviews():
    try:
        data = request.json
        session_id = data.get('session_id')
        student_id = data.get('student_id')
        scores = data.get('scores', [])
        
        if not session_id or not student_id or not scores:
            return jsonify({"error": "缺少必要参数"}), 400
        
        # 开始事务
        conn = get_db_connection()
        conn.execute('BEGIN TRANSACTION')
        
        updated_count = 0
        try:
            for score_item in scores:
                question_id = score_item.get('question_id')
                final_score = score_item.get('final_score')
                
                if not question_id or final_score is None:
                    continue
                
                # 更新答案评分和复核状态
                query = """
                    UPDATE student_answers
                    SET final_score = ?, review_status = 'reviewed'
                    WHERE session_id = ? AND student_id = ? AND question_id = ?
                """
                cursor = conn.execute(query, (final_score, session_id, student_id, question_id))
                if cursor.rowcount > 0:
                    updated_count += 1        

            # 标记学生复核状态
            if updated_count > 0:
                query = """
                    UPDATE student_exams
                    SET status = 'graded'
                    WHERE session_id = ? AND student_id = ?
                """
                conn.execute(query, (session_id, student_id))
            
            # 提交事务
            conn.execute('COMMIT')

            stats_query = """
                SELECT sum(final_score) FROM student_answers WHERE session_id = ? AND student_id = ?;
            """
            success, stats_data = execute_query(stats_query, (session_id, student_id))
            if success:
                student_score = stats_data[0]['sum(final_score)']

            if updated_count > 0:
                query = """
                    UPDATE student_exams
                    SET score = ?
                    WHERE session_id = ? AND student_id = ?
                """
                execute_query(query, (student_score, session_id, student_id))
            
            return jsonify({
                "success": True,
                "message": f"成功更新{updated_count}个题目的评分",
                "updated_count": updated_count
            })
            
        except Exception as e:
            # 回滚事务
            conn.execute('ROLLBACK')
            raise e
        finally:
            conn.close()
            
    except Exception as e:
        app.logger.error(f"一键复核失败: {str(e)}")
        return jsonify({"error": f"处理请求时出错: {str(e)}"}), 500

# 导入AI分析模块
from ai_analysis import ai_analysis as ai_analysis_blueprint

# 注册蓝图
app.register_blueprint(ai_analysis_blueprint)

@app.route('/api/review/mark-as-graded', methods=['POST'])
@login_required
def mark_exam_as_graded():
    """将考试标记为已评分并生成统计数据"""
    try:
        data = request.get_json()
        session_id = data.get('session_id')
        
        if not session_id:
            return jsonify({'success': False, 'error': '缺少考试ID'}), 400
        
        # 1. 更新考试状态为已评分
        update_query = """
            UPDATE exam_sessions
            SET status = 'graded'
            WHERE id = ?
        """
        success, _ = execute_query(update_query, (session_id,))
        
        if not success:
            return jsonify({'success': False, 'error': '更新考试状态失败'}), 500
        
        # 2. 先获取每个学生的总分
        student_scores_query = """
            SELECT 
                student_id,
                SUM(COALESCE(final_score, ai_score, 0)) as total_score
            FROM student_answers
            WHERE session_id = ?
            GROUP BY student_id
        """
        success, student_scores = execute_query(student_scores_query, (session_id,))
        
        if not success:
            return jsonify({'success': False, 'error': '计算学生分数失败'}), 500
        
        # 3. 逐个更新学生分数
        for student in student_scores:
            update_score_query = """
                UPDATE student_exams
                SET score = ?,
                    status = 'graded'
                WHERE student_id = ? AND session_id = ?
            """
            success, _ = execute_query(update_score_query, (
                float(student['total_score']),  # 确保转换为浮点数
                student['student_id'],
                session_id
            ))
            
            if not success:
                print(f"更新学生 {student['student_id']} 的分数失败")
        
        # 4. 计算统计数据
        stats_query = """
            SELECT 
                AVG(score) as average_score,
                SUM(CASE WHEN score >= 60 THEN 1 ELSE 0 END) * 1.0 / COUNT(*) as pass_rate,
                MAX(score) as highest_score,
                (SELECT student_id FROM student_exams WHERE session_id = ? ORDER BY score DESC LIMIT 1) as highest_score_student_id,
                COUNT(*) as total_students
            FROM student_exams
            WHERE session_id = ? AND status = 'graded'
        """
        success, stats = execute_query(stats_query, (session_id, session_id))
        
        if not success or not stats:
            return jsonify({'success': False, 'error': '计算统计数据失败'}), 500
        
        # 5. 存储统计数据
        check_query = "SELECT id FROM score_statistics WHERE session_id = ?"
        success, existing = execute_query(check_query, (session_id,))
        
        if success and existing:
            update_stats_query = """
                UPDATE score_statistics
                SET average_score = ?, pass_rate = ?, highest_score = ?, 
                    highest_score_student_id = ?, updated_at = CURRENT_TIMESTAMP
                WHERE session_id = ?
            """
            execute_query(update_stats_query, (
                stats[0]['average_score'],
                stats[0]['pass_rate'],
                stats[0]['highest_score'],
                stats[0]['highest_score_student_id'],
                session_id
            ))
        else:
            insert_stats_query = """
                INSERT INTO score_statistics 
                (session_id, average_score, pass_rate, highest_score, highest_score_student_id)
                VALUES (?, ?, ?, ?, ?)
            """
            execute_query(insert_stats_query, (
                session_id,
                stats[0]['average_score'],
                stats[0]['pass_rate'],
                stats[0]['highest_score'],
                stats[0]['highest_score_student_id']
            ))
        
        # 在mark_exam_as_graded函数中更新分数分布数据部分
        # 删除旧记录
        delete_dist_query = "DELETE FROM score_distribution WHERE session_id = ?"
        execute_query(delete_dist_query, (session_id,))

        # 按分数段统计学生并记录
        ranges = ['90-100', '80-89', '70-79', '60-69', '0-59']
        for range_str in ranges:
            # 获取分数段的范围
            if range_str == '90-100':
                min_score, max_score = 90, 100
            elif range_str == '80-89':
                min_score, max_score = 80, 89.99
            elif range_str == '70-79':
                min_score, max_score = 70, 79.99
            elif range_str == '60-69':
                min_score, max_score = 60, 69.99
            else:  # 0-59
                min_score, max_score = 0, 59.99
            
            # 查询该分数段的学生数和未评分数
            count_query = """
                SELECT COUNT(*) as total,
                    SUM(CASE WHEN status != 'graded' THEN 1 ELSE 0 END) as ungraded
                FROM student_exams
                WHERE session_id = ? AND score >= ? AND score <= ?
            """
            success, counts = execute_query(count_query, (session_id, min_score, max_score))
            
            if success and counts:
                total_count = counts[0]['total'] or 0
                ungraded_count = counts[0]['ungraded'] or 0
                
                # 插入分数分布记录
                insert_query = """
                    INSERT INTO score_distribution 
                    (session_id, score_range, student_count, ungraded_count)
                    VALUES (?, ?, ?, ?)
                """
                execute_query(insert_query, (session_id, range_str, total_count, ungraded_count))
        
        # 9. 计算题目分析数据并存入question_analysis表
        # 先获取所有题目
        questions_query = """
            SELECT id, score
            FROM questions
            WHERE session_id = ?
        """
        success, questions = execute_query(questions_query, (session_id,))
        
        if not success:
            return jsonify({'success': False, 'error': '获取题目数据失败'}), 500
        
        # 计算每道题的分析数据并存入数据库
        for question in questions:
            # 删除现有分析数据
            delete_query = """
                DELETE FROM question_analysis 
                WHERE session_id = ? AND question_id = ?
            """
            execute_query(delete_query, (session_id, question['id']))
            
            # 计算平均得分率
            avg_score_query = """
                SELECT AVG(COALESCE(final_score, ai_score, 0) / ?) as avg_score_rate
                FROM student_answers
                WHERE session_id = ? AND question_id = ?
            """
            success, avg_score = execute_query(avg_score_query, (
                question['score'], 
                session_id, 
                question['id']
            ))

            if not success or not avg_score:
                continue
                
            # 确保平均得分率在0-1范围内
            avg_score_rate = avg_score[0]['avg_score_rate'] or 0
            if avg_score_rate > 1:
                avg_score_rate = 1.0
            elif avg_score_rate < 0:
                avg_score_rate = 0.0

            # 计算难度系数（等于1减去平均得分率）
            difficulty_coefficient = 1 - avg_score_rate
            
            # 计算区分度（高分组与低分组的得分率差异）
            # 先获取所有学生成绩
            students_query = """
                SELECT student_id, SUM(COALESCE(final_score, ai_score, 0)) as total_score
                FROM student_answers
                WHERE session_id = ?
                GROUP BY student_id
                ORDER BY total_score DESC
            """
            success, students = execute_query(students_query, (session_id,))
            
            if not success or not students:
                continue
                
            # 分高分组和低分组（各取27%）
            total_students = len(students)
            group_size = max(1, int(total_students * 0.27))
            
            high_group = [s['student_id'] for s in students[:group_size]]
            low_group = [s['student_id'] for s in students[-group_size:]]
            
            # 计算高分组在此题的得分率
            high_score_query = """
                SELECT AVG(COALESCE(final_score, ai_score, 0) / ?) as score_rate
                FROM student_answers
                WHERE session_id = ? AND question_id = ? AND student_id IN ({})
            """.format(','.join(['?'] * len(high_group)))
            
            params = [question['score'], session_id, question['id']] + high_group
            success, high_result = execute_query(high_score_query, params)
            
            high_score_rate = high_result[0]['score_rate'] if success and high_result else 0
            
            # 计算低分组在此题的得分率
            low_score_query = """
                SELECT AVG(COALESCE(final_score, ai_score, 0) / ?) as score_rate
                FROM student_answers
                WHERE session_id = ? AND question_id = ? AND student_id IN ({})
            """.format(','.join(['?'] * len(low_group)))
            
            params = [question['score'], session_id, question['id']] + low_group
            success, low_result = execute_query(low_score_query, params)
            
            low_score_rate = low_result[0]['score_rate'] if success and low_result else 0
            
            # 区分度 = 高分组得分率 - 低分组得分率
            discrimination_degree = high_score_rate - low_score_rate
            
            # 插入分析数据
            insert_query = """
                INSERT INTO question_analysis 
                (session_id, question_id, average_score_rate, difficulty_coefficient, discrimination_degree)
                VALUES (?, ?, ?, ?, ?)
            """
            execute_query(insert_query, (
                session_id,
                question['id'],
                avg_score_rate,
                difficulty_coefficient,
                discrimination_degree
            ))
        
        return jsonify({'success': True, 'message': '考试已标记为已评分'}), 200
    except Exception as e:
        print(f"标记考试为已评分时出错: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# 在应用启动时初始化数据库表
if __name__ == '__main__':
    # 确保数据库目录存在
    os.makedirs('database', exist_ok=True)
    # 启动应用
    app.run(debug=True, host='0.0.0.0', port=5000)


