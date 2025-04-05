# ai_analysis.py
from flask import Blueprint, request, jsonify, session, current_app, send_from_directory
from functools import wraps
import sqlite3
import json
import os
import traceback
from datetime import datetime
import matplotlib.pyplot as plt
import io
import base64
import pandas as pd
from werkzeug.utils import secure_filename
# 导入AI模块
from split_and_ocr.ai import new as ai_new
import time

# 创建蓝图
ai_analysis = Blueprint('ai_analysis', __name__)

# 存储最近分析结果
latest_ai_analysis = None

# 数据库连接函数
def get_db_connection():
    conn = sqlite3.connect('database/It_g.db')
    conn.row_factory = sqlite3.Row
    return conn

# 解析用户查询
def parse_query(query):
    """使用AI模型解析用户的自然语言查询"""
    system_prompt = """你是一个教育数据分析助手，负责解析用户的查询意图。
请根据用户输入的自然语言查询，判断用户想要什么类型的分析，并提取出相关实体。

查询类型包括:
1. student_comparison: 比较两个学生的成绩
2. student_analysis: 分析单个学生的成绩
3. exam_analysis: 分析一次考试的情况
4. general_analysis: 通用查询，无法明确归类

请返回一个JSON格式的输出，包含以下字段:
{
  "query_type": "查询类型",
  "entities": {
    // 根据查询类型包含不同的字段
    // 对于student_comparison，包含student1和student2
    // 对于student_analysis，包含student
    // 对于exam_analysis，可能包含exam，例如"latest"表示最近一次考试
  }
}

例如:
用户查询: "请分析张三和李四的成绩对比"
返回: {"query_type": "student_comparison", "entities": {"student1": "张三", "student2": "李四"}}

用户查询: "请分析张三的成绩"
返回: {"query_type": "student_analysis", "entities": {"student": "张三"}}

用户查询: "分析最近一次考试的情况"
返回: {"query_type": "exam_analysis", "entities": {"exam": "latest"}}

请注意:
- 学生姓名应该是1-3个汉字
- 如果查询中没有明确的学生姓名，但应该是student_analysis类型，请尝试从查询中提取可能的学生姓名
- 如果是与python相关的考试查询，格式可能是"给我沈亦豪这个学生python这次考试的分析"，请正确提取学生姓名"沈亦豪"
"""
    
    user_prompt = query
    
    try:
        # 调用AI解析查询
        response = ai_new(system_prompt, user_prompt)
        
        # 解析JSON响应
        try:
            result = json.loads(response)
            # 验证基本结构
            if "query_type" not in result or "entities" not in result:
                print(f"AI返回结构不完整: {response}")
                return "general_analysis", {}
                
            return result["query_type"], result["entities"]
        except json.JSONDecodeError:
            print(f"无法解析AI返回的JSON: {response}")
            # 尝试提取关键信息，兜底处理
            if "student_comparison" in response and "student1" in response and "student2" in response:
                # 尝试提取学生名称
                import re
                students = re.findall(r'"student\d": "([^"]+)"', response)
                if len(students) >= 2:
                    return "student_comparison", {"student1": students[0], "student2": students[1]}
            
            if "student_analysis" in response and "student" in response:
                # 尝试提取学生名称
                import re
                student = re.search(r'"student": "([^"]+)"', response)
                if student:
                    return "student_analysis", {"student": student.group(1)}
            
            return "general_analysis", {}
    except Exception as e:
        print(f"AI解析查询时出错: {str(e)}")
        return "general_analysis", {}

@ai_analysis.route('/api/ai/analyze', methods=['POST'])
def analyze_data():
    """处理AI分析请求"""
    global latest_ai_analysis
    
    try:
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({'message': '请提供查询内容', 'analysis': None})
        
        query = data['query']
        session_id = data.get('session_id')
        
        # 使用AI解析用户查询
        query_type, entities = parse_query(query)
        print(f"解析查询结果: 类型={query_type}, 实体={entities}")
        
        # 根据查询类型进行不同的处理
        if query_type == "student_comparison":
            if len(entities) < 2 or "student1" not in entities or "student2" not in entities:
                # 使用AI生成友好提示
                error_prompt = f"""用户想要比较学生成绩，但未提供足够信息。
原始查询: "{query}"
请用友好的语气提醒用户需要提供两个学生的姓名。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or "请提供两个学生的姓名进行比较，例如'请分析张三和李四的成绩对比'",
                    'analysis': None
                })
            
            # 获取两名学生的信息
            conn = get_db_connection()
            
            # 查询第一个学生
            student1 = conn.execute(
                'SELECT id, name, student_id FROM students WHERE name = ?', 
                (entities["student1"],)
            ).fetchone()
            
            # 查询第二个学生
            student2 = conn.execute(
                'SELECT id, name, student_id FROM students WHERE name = ?', 
                (entities["student2"],)
            ).fetchone()
            
            if not student1 or not student2:
                missing_students = []
                if not student1:
                    missing_students.append(entities["student1"])
                if not student2:
                    missing_students.append(entities["student2"])
                
                conn.close()
                
                # 使用AI生成友好提示
                error_prompt = f"""用户想要比较学生成绩，但以下学生不在数据库中: {', '.join(missing_students)}
请用友好的语气提醒用户检查这些学生的姓名是否正确。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or f"找不到学生: {', '.join(missing_students)}，请检查姓名是否正确",
                    'analysis': None
                })
            
            # 获取学生考试记录 - 使用正确的表结构
            student1_exams = conn.execute(
                '''SELECT se.id, se.score, se.session_id, es.name as exam_name, 
                   es.exam_score as total_score, es.start_time as exam_date
                   FROM student_exams se
                   JOIN exam_sessions es ON se.session_id = es.id
                   WHERE se.student_id = ? AND se.status = 'graded' 
                   ORDER BY es.start_time DESC''', 
                (student1['id'],)
            ).fetchall()
            
            student2_exams = conn.execute(
                '''SELECT se.id, se.score, se.session_id, es.name as exam_name, 
                   es.exam_score as total_score, es.start_time as exam_date
                   FROM student_exams se
                   JOIN exam_sessions es ON se.session_id = es.id
                   WHERE se.student_id = ? AND se.status = 'graded'
                   ORDER BY es.start_time DESC''', 
                (student2['id'],)
            ).fetchall()
            
            # 将查询结果转为字典列表
            student1_exams_list = [dict(exam) for exam in student1_exams]
            student2_exams_list = [dict(exam) for exam in student2_exams]
            
            conn.close()
            
            if not student1_exams_list or not student2_exams_list:
                # 检查是否有学生没有考试记录
                no_exams_students = []
                if not student1_exams_list:
                    no_exams_students.append(student1['name'])
                if not student2_exams_list:
                    no_exams_students.append(student2['name'])
                
                error_prompt = f"""以下学生没有考试记录: {', '.join(no_exams_students)}
请用友好的语气告知用户这些学生尚无成绩数据可供分析。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or f"学生 {', '.join(no_exams_students)} 没有考试记录，无法进行分析",
                    'analysis': None
                })
            
            # 生成比较分析
            comparison_result = compare_students(
                {"name": student1['name'], "student_id": student1['student_id'], "exams": student1_exams_list},
                {"name": student2['name'], "student_id": student2['student_id'], "exams": student2_exams_list}
            )
            
            # 保存分析结果用于导出
            latest_ai_analysis = {
                'type': 'student_comparison',
                'data': comparison_result,
                'query': query,
                'timestamp': datetime.now().isoformat(),
                'student1': {"id": student1['id'], "name": student1['name'], "student_id": student1['student_id']},
                'student2': {"id": student2['id'], "name": student2['name'], "student_id": student2['student_id']}
            }
            
            # 使用AI生成更友好的消息
            success_prompt = f"""已成功生成{student1['name']}和{student2['name']}的成绩对比分析。
请用友好的语气告知用户分析已生成，并简短提示用户可以查看分析结果或导出报告。"""
            message = ai_new("你是一个友好的教育助手", success_prompt)
            
            return jsonify({
                'message': message or f"已生成{student1['name']}和{student2['name']}的成绩对比分析",
                'analysis': comparison_result
            })
            
        elif query_type == "student_analysis":
            if "student" not in entities:
                # 使用AI生成友好提示
                error_prompt = f"""用户想要分析学生成绩，但未提供学生姓名。
原始查询: "{query}"
请用友好的语气提醒用户需要提供学生姓名。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or "请提供学生姓名进行分析，例如'请分析张三的成绩'",
                    'analysis': None
                })
            
            # 查询学生信息
            conn = get_db_connection()
            
            student = conn.execute(
                'SELECT id, name, student_id FROM students WHERE name = ?', 
                (entities["student"],)
            ).fetchone()
            
            if not student:
                conn.close()
                
                # 使用AI生成友好提示
                error_prompt = f"""用户想要分析学生成绩，但学生"{entities['student']}"不在数据库中。
请用友好的语气提醒用户检查学生姓名是否正确。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or f"找不到学生: {entities['student']}，请检查姓名是否正确",
                    'analysis': None
                })
            
            # 获取学生考试记录
            exams = conn.execute(
                '''SELECT se.id, se.score, se.session_id, es.name as exam_name, 
                   es.exam_score as total_score, es.start_time as exam_date,
                   (SELECT COUNT(*) + 1 FROM student_exams se2 
                    WHERE se2.session_id = se.session_id AND se2.score > se.score) as class_rank
                   FROM student_exams se
                   JOIN exam_sessions es ON se.session_id = es.id
                   WHERE se.student_id = ? AND se.status = 'graded'
                   ORDER BY es.start_time DESC''', 
                (student['id'],)
            ).fetchall()
            
            conn.close()
            
            if not exams:
                # 使用AI生成友好提示
                error_prompt = f"""学生"{student['name']}"没有考试记录。
请用友好的语气告知用户该学生尚无成绩数据可供分析。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or f"学生 {student['name']} 没有考试记录，无法进行分析",
                    'analysis': None
                })
            
            # 将查询结果转为字典列表
            exams_list = []
            for exam in exams:
                exam_dict = dict(exam)
                # 计算得分率
                exam_dict['score_percentage'] = round(exam_dict['score'] / exam_dict['total_score'] * 100, 2)
                exams_list.append(exam_dict)
            
            # 检查Python考试的特殊处理
            if "python" in query.lower():
                python_exams = [e for e in exams_list if "python" in e['exam_name'].lower()]
                if python_exams:
                    # 使用最近的Python考试
                    latest_python_exam = python_exams[0]
                    # 生成详细的Python考试分析
                    python_analysis = generate_python_exam_analysis(student, latest_python_exam)
                    
                    # 保存分析结果用于导出
                    latest_ai_analysis = {
                        'type': 'python_exam_analysis',
                        'data': python_analysis,
                        'query': query,
                        'timestamp': datetime.now().isoformat(),
                        'student': {"id": student['id'], "name": student['name'], "student_id": student['student_id']},
                        'exam': latest_python_exam
                    }
                    
                    # 使用AI生成友好消息
                    success_prompt = f"""已生成{student['name']}的Python考试分析报告。
请用友好的语气告知用户分析已生成，并简短提示用户可以查看分析结果。"""
                    message = ai_new("你是一个友好的教育助手", success_prompt)
                    
                    return jsonify({
                        'message': message or f"已生成{student['name']}的Python考试分析报告",
                        'analysis': python_analysis
                    })
                else:
                    # 找不到Python考试
                    error_prompt = f"""学生"{student['name']}"没有Python考试记录。
请用友好的语气告知用户该学生尚无Python考试数据可供分析。"""
                    message = ai_new("你是一个友好的教育助手", error_prompt)
                    return jsonify({
                        'message': message or f"学生 {student['name']} 没有Python考试记录，无法进行分析",
                        'analysis': None
                    })
            
            # 生成一般分析
            analysis_result = analyze_student({
                "student": {"id": student['id'], "name": student['name'], "student_id": student['student_id']},
                "history": exams_list
            })
            
            # 保存分析结果用于导出
            latest_ai_analysis = {
                'type': 'student_analysis',
                'data': analysis_result,
                'query': query,
                'timestamp': datetime.now().isoformat(),
                'student': {"id": student['id'], "name": student['name'], "student_id": student['student_id']},
                'history': exams_list
            }
            
            # 使用AI生成友好消息
            success_prompt = f"""已生成{student['name']}的成绩分析报告。
请用友好的语气告知用户分析已生成，并简短提示用户可以查看分析结果或导出报告。"""
            message = ai_new("你是一个友好的教育助手", success_prompt)
            
            return jsonify({
                'message': message or f"已生成{student['name']}的成绩分析报告",
                'analysis': analysis_result
            })
            
        elif query_type == "exam_analysis":
            # 处理考试ID
            if not session_id and "exam" in entities and entities["exam"] == "latest":
                # 获取最新考试
                conn = get_db_connection()
                latest_exam = conn.execute(
                    'SELECT id FROM exam_sessions ORDER BY start_time DESC LIMIT 1'
                ).fetchone()
                
                if latest_exam:
                    session_id = latest_exam['id']
                conn.close()
            
            if not session_id:
                # 使用AI生成友好提示
                error_prompt = f"""用户想要分析考试情况，但未提供具体考试。
原始查询: "{query}"
请用友好的语气提醒用户需要在分析页面选择一个考试或明确指定最近一次考试。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or "请选择一个考试或者指定最近一次考试",
                    'analysis': None
                })
            
            # 获取考试信息和统计数据
            conn = get_db_connection()
            
            # 考试基本信息
            exam_info = conn.execute(
                'SELECT id, name, subject, start_time, exam_score as total_score FROM exam_sessions WHERE id = ?',
                (session_id,)
            ).fetchone()
            
            if not exam_info:
                conn.close()
                
                # 使用AI生成友好提示
                error_prompt = f"""用户想要分析考试ID为{session_id}的考试，但该考试不存在。
请用友好的语气告知用户找不到指定的考试。"""
                message = ai_new("你是一个友好的教育助手", error_prompt)
                return jsonify({
                    'message': message or f"找不到考试，ID: {session_id}",
                    'analysis': None
                })
            
            # 获取统计数据
            stats = conn.execute(
                'SELECT average_score, pass_rate, highest_score FROM score_statistics WHERE session_id = ?',
                (session_id,)
            ).fetchone()
            
            # 获取分数分布
            distribution = conn.execute(
                'SELECT score_range, student_count FROM score_distribution WHERE session_id = ?',
                (session_id,)
            ).fetchall()
            
            # 获取题目分析
            questions = conn.execute(
                '''SELECT q.id, q.question_text, q.score, qa.average_score_rate, 
                   qa.difficulty_coefficient as difficulty, qa.discrimination_degree as discrimination 
                   FROM questions q
                   JOIN question_analysis qa ON q.id = qa.question_id
                   WHERE q.session_id = ? AND qa.session_id = ?
                   ORDER BY q.question_order''',
                (session_id, session_id)
            ).fetchall()
            
            # 获取学生人数
            student_count = conn.execute(
                'SELECT COUNT(*) as count FROM student_exams WHERE session_id = ?',
                (session_id,)
            ).fetchone()
            
            conn.close()
            
            # 组织数据
            exam_data = {
                'basic_stats': {
                    'exam_name': exam_info['name'],
                    'subject': exam_info['subject'],
                    'exam_date': exam_info['start_time'],
                    'total_score': exam_info['total_score'],
                    'average_score': stats['average_score'] if stats else 0,
                    'pass_rate': stats['pass_rate'] * 100 if stats else 0,
                    'highest_score': stats['highest_score'] if stats else 0,
                    'total_students': student_count['count'] if student_count else 0
                },
                'distribution': [dict(d) for d in distribution],
                'questions': [dict(q) for q in questions]
            }
            
            # 生成分析
            analysis_result = generate_exam_ai_analysis(exam_data)
            
            # 保存分析结果用于导出
            latest_ai_analysis = {
                'type': 'exam_analysis',
                'data': analysis_result,
                'query': query,
                'timestamp': datetime.now().isoformat(),
                'session_id': session_id,
                'raw_data': exam_data
            }
            
            # 使用AI生成友好消息
            success_prompt = f"""已生成{exam_info['name']}的考试分析报告。
请用友好的语气告知用户分析已生成，并简短提示用户可以查看分析结果或导出报告。"""
            message = ai_new("你是一个友好的教育助手", success_prompt)
            
            return jsonify({
                'message': message or f"已生成{exam_info['name']}的分析报告",
                'analysis': analysis_result
            })
            
        else:
            # 通用查询或未识别的查询类型
            # 使用AI解释助手能做什么
            system_prompt = """你是一位教育分析助手，帮助用户分析学生成绩和考试数据。
请用友好的语气解释你能够做什么，并举几个例子，如：
1. 分析某个学生的历史成绩
2. 比较两个学生的成绩
3. 分析某次考试的整体情况
以简短、友好的语气回复。"""
            
            explanation = ai_new(system_prompt, query)
            return jsonify({
                'message': explanation or "我能够帮您分析学生成绩或考试数据，例如：\n- 请分析张三和李四的成绩对比\n- 请分析张三的历史成绩\n- 分析最近一次考试的情况",
                'analysis': None
            })
            
    except Exception as e:
        traceback.print_exc()
        
        # 使用AI生成错误消息
        error_prompt = f"""分析过程中出现错误: {str(e)}
请用友好的语气告知用户系统遇到了问题，并给出简单的建议。"""
        message = ai_new("你是一个友好的教育助手", error_prompt)
        
        return jsonify({
            'message': message or f"分析过程中出错: {str(e)}",
            'analysis': None
        }), 500

@ai_analysis.route('/api/ai/export-analysis', methods=['POST'])
def export_ai_analysis():
    """导出AI分析报告"""
    try:
        if not latest_ai_analysis:
            print("没有可导出的分析结果")
            return jsonify({'error': '没有可导出的分析结果'}), 400
            
        data = request.get_json()
        format_type = data.get('format', 'pdf')  # 默认为PDF格式
        
        # 确定报告类型和文件名
        report_type = latest_ai_analysis.get('type')
        timestamp = int(time.time())
        print(f"准备导出报告，类型: {report_type}, 格式: {format_type}")
        
        if report_type == 'student_comparison':
            filename = f'student_comparison_{timestamp}'
            report_data = latest_ai_analysis.get('data', {})
        elif report_type == 'student_analysis':
            student_name = latest_ai_analysis.get('data', {}).get('student_name', 'unknown')
            filename = f'student_analysis_{student_name}_{timestamp}'
            report_data = latest_ai_analysis.get('data', {})
        elif report_type == 'exam_analysis':
            exam_name = latest_ai_analysis.get('data', {}).get('exam_name', 'unknown')
            filename = f'exam_analysis_{exam_name}_{timestamp}'
            report_data = latest_ai_analysis.get('data', {})
        else:
            # 对于其他类型，统一使用ai_analysis前缀
            filename = f'ai_analysis_{timestamp}'
            report_data = latest_ai_analysis.get('data', {})
        
        # 确保目录存在
        report_dir = os.path.join('exports')
        os.makedirs(report_dir, exist_ok=True)
        file_path = os.path.join(report_dir, f'{filename}.{format_type}')
        
        print(f"正在生成报告文件: {file_path}")
        
        # 根据报告类型生成相应格式的报告
        success = False
        if report_type == 'student_comparison':
            success = generate_comparison_report(report_data, format_type, file_path)
        elif report_type == 'student_analysis':
            success = generate_student_report(report_data, format_type, file_path)
        elif report_type == 'exam_analysis' or report_type == 'python_exam_analysis':
            # 处理常规考试分析和Python考试分析
            success = generate_exam_report(report_data, format_type, file_path)
        else:
            # 对于未知类型，使用学生报告格式（因为结构最简单）
            print(f"未知报告类型: {report_type}，使用通用报告格式")
            success = generate_student_report(report_data, format_type, file_path)
        
        if not success:
            print(f"报告生成失败: {file_path}")
            return jsonify({'error': '报告生成失败，请查看服务器日志'}), 500
            
        print(f"报告生成成功: {file_path}")
        
        # 尝试记录导出行为
        try:
            user_id = session.get('user_id')
            if user_id is not None:
                # 如果用户已登录，记录导出行为
                conn = sqlite3.connect('examination.db')
                c = conn.cursor()
                c.execute('''
                    INSERT INTO export_logs (user_id, export_type, export_format, export_time)
                    VALUES (?, ?, ?, ?)
                ''', (user_id, report_type, format_type, int(time.time())))
                conn.commit()
                conn.close()
        except Exception as e:
            # 记录错误，但不影响文件下载
            print(f"记录导出行为失败: {str(e)}")
        
        # 返回文件下载链接
        download_filename = f'{filename}.{format_type}'
        file_url = f"/api/ai/download-report/{download_filename}"
        
        # 验证文件是否确实存在
        if not os.path.exists(file_path):
            print(f"警告：文件路径存在但文件不存在: {file_path}")
            return jsonify({'error': '报告文件生成后未找到，请联系管理员'}), 500
            
        print(f"文件下载链接: {file_url}")
        
        return jsonify({
            'success': True,
            'message': '报告生成成功',
            'file_url': file_url,
            'filename': download_filename
        })
            
    except Exception as e:
        print(f"导出分析报告失败: {str(e)}")
        traceback.print_exc()  # 打印完整的错误堆栈
        return jsonify({'error': f'导出失败: {str(e)}'}), 500

# 辅助函数 - 学生对比分析
def compare_students(student1_data, student2_data):
    """比较两个学生的成绩数据"""
    student1 = student1_data
    student2 = student2_data
    
    # 找出共同的考试
    exam_ids1 = {exam['session_id'] for exam in student1['exams']}
    exam_ids2 = {exam['session_id'] for exam in student2['exams']}
    common_exams = exam_ids1.intersection(exam_ids2)
    
    if not common_exams:
        return {
            'text': f"<p>{student1['name']}和{student2['name']}没有共同参加的考试，无法进行比较。</p>"
        }
    
    # 构建共同考试的数据
    common_data = []
    for exam_id in common_exams:
        s1_exam = next((e for e in student1['exams'] if e['session_id'] == exam_id), None)
        s2_exam = next((e for e in student2['exams'] if e['session_id'] == exam_id), None)
        
        if s1_exam and s2_exam:
            s1_percentage = s1_exam['score'] / s1_exam['total_score'] * 100
            s2_percentage = s2_exam['score'] / s2_exam['total_score'] * 100
            
            common_data.append({
                'exam_name': s1_exam['exam_name'],
                'student1_score': s1_exam['score'],
                'student2_score': s2_exam['score'],
                'student1_percentage': round(s1_percentage, 2),
                'student2_percentage': round(s2_percentage, 2),
                'total_score': s1_exam['total_score']
            })
    
    # 对考试按时间排序
    common_data.sort(key=lambda x: x['exam_name'])
    
    # 构建图表数据
    chart_data = {
        'type': 'line',
        'data': {
            'labels': [d['exam_name'] for d in common_data],
            'datasets': [
                {
                    'label': student1['name'],
                    'data': [d['student1_percentage'] for d in common_data],
                    'borderColor': 'rgb(75, 192, 192)',
                    'backgroundColor': 'rgba(75, 192, 192, 0.1)',
                    'tension': 0.1
                },
                {
                    'label': student2['name'],
                    'data': [d['student2_percentage'] for d in common_data],
                    'borderColor': 'rgb(255, 99, 132)',
                    'backgroundColor': 'rgba(255, 99, 132, 0.1)',
                    'tension': 0.1
                }
            ]
        }
    }
    
    # 计算统计数据
    s1_avg = sum(d['student1_percentage'] for d in common_data) / len(common_data) if common_data else 0
    s2_avg = sum(d['student2_percentage'] for d in common_data) / len(common_data) if common_data else 0
    
    s1_wins = sum(1 for d in common_data if d['student1_percentage'] > d['student2_percentage'])
    s2_wins = sum(1 for d in common_data if d['student2_percentage'] > d['student1_percentage'])
    
    # 使用AI生成分析结果
    system_prompt = f"""你是一位教育分析专家，擅长分析学生成绩数据。
请分析以下两名学生的成绩对比数据，给出专业、客观的分析:

学生1: {student1['name']} (学号: {student1['student_id']})
学生2: {student2['name']} (学号: {student2['student_id']})

共同参加的考试次数: {len(common_data)}
{student1['name']}的平均得分率: {s1_avg:.2f}%
{student2['name']}的平均得分率: {s2_avg:.2f}%
{student1['name']}表现更好的考试次数: {s1_wins}
{student2['name']}表现更好的考试次数: {s2_wins}

请提供详细的分析，包括两位学生的整体表现对比、优势劣势以及学习建议。
以HTML格式输出，使用<h3>、<p>等标签。
"""
    
    user_prompt = f"请分析{student1['name']}和{student2['name']}的考试成绩对比"
    
    try:
        # 调用AI生成分析文本
        ai_response = ai_new(system_prompt, user_prompt)
        if ai_response and len(ai_response) > 50:  # 确保有有效的返回
            analysis_text = ai_response
        else:
            # 备用生成方式
            analysis_text = f"""
            <h3>{student1['name']}和{student2['name']}的成绩对比分析</h3>
            <p>共同参加的考试: {len(common_data)}次</p>
            <p>{student1['name']}的平均得分率: {s1_avg:.2f}%</p>
            <p>{student2['name']}的平均得分率: {s2_avg:.2f}%</p>
            <p>整体表现: {student1['name'] if s1_avg > s2_avg else student2['name']}的总体表现更好</p>
            <p>各次考试对比: {student1['name']}在{s1_wins}次考试中表现更好，{student2['name']}在{s2_wins}次考试中表现更好</p>
            """
    except Exception as e:
        print(f"调用AI分析出错: {str(e)}")
        # 备用生成方式
        analysis_text = f"""
        <h3>{student1['name']}和{student2['name']}的成绩对比分析</h3>
        <p>共同参加的考试: {len(common_data)}次</p>
        <p>{student1['name']}的平均得分率: {s1_avg:.2f}%</p>
        <p>{student2['name']}的平均得分率: {s2_avg:.2f}%</p>
        <p>整体表现: {student1['name'] if s1_avg > s2_avg else student2['name']}的总体表现更好</p>
        <p>各次考试对比: {student1['name']}在{s1_wins}次考试中表现更好，{student2['name']}在{s2_wins}次考试中表现更好</p>
        """
    
    return {
        'text': analysis_text,
        'chart': chart_data
    }

# 辅助函数 - 单个学生分析
def analyze_student(student_data):
    """分析单个学生的历史成绩"""
    student = student_data.get('student', {})
    history = student_data.get('history', [])
    
    if not history:
        return {
            'text': f"<p>{student.get('name')}没有历史成绩记录，无法进行分析。</p>"
        }
    
    # 按时间排序
    sorted_history = sorted(history, key=lambda x: x.get('exam_date', ''))
    
    # 计算统计数据
    avg_score = sum(h.get('score_percentage', 0) for h in history) / len(history) if history else 0
    
    trend_direction = "上升" if sorted_history[-1].get('score_percentage', 0) > sorted_history[0].get('score_percentage', 0) else "下降"
    
    if len(sorted_history) >= 2:
        recent_trend = "上升" if sorted_history[-1].get('score_percentage', 0) > sorted_history[-2].get('score_percentage', 0) else "下降"
    else:
        recent_trend = "稳定"
    
    # 构建图表数据
    chart_data = {
        'type': 'line',
        'data': {
            'labels': [h.get('exam_name', '未知考试') for h in sorted_history],
            'datasets': [{
                'label': '得分率(%)',
                'data': [h.get('score_percentage', 0) for h in sorted_history],
                'borderColor': 'rgb(75, 192, 192)',
                'backgroundColor': 'rgba(75, 192, 192, 0.1)',
                'tension': 0.1,
                'fill': True
            }]
        }
    }
    
    # 使用AI生成分析结果
    scores_info = "\n".join([
        f"考试: {h.get('exam_name', '未知考试')}, 得分: {h.get('score', 0)}/{h.get('total_score', 100)}, " + 
        f"得分率: {h.get('score_percentage', 0):.2f}%, 班级排名: {h.get('class_rank', 'N/A')}"
        for h in sorted_history
    ])
    
    system_prompt = f"""你是一位教育分析专家，擅长分析学生成绩数据。
请分析以下学生的历史成绩数据，给出专业、客观的分析:

学生: {student.get('name')} (学号: {student.get('student_id')})
考试次数: {len(history)}
平均得分率: {avg_score:.2f}%
成绩整体趋势: {trend_direction}
最近考试趋势: {recent_trend}

详细成绩数据:
{scores_info}

请提供详细的分析，包括学生的整体表现、成绩变化趋势、优势劣势以及学习建议。
以HTML格式输出，使用<h3>、<p>等标签。
"""
    
    user_prompt = f"请分析{student.get('name')}的历史成绩数据"
    
    try:
        # 调用AI生成分析文本
        ai_response = ai_new(system_prompt, user_prompt)
        if ai_response and len(ai_response) > 50:  # 确保有有效的返回
            analysis_text = ai_response
        else:
            # 备用生成方式
            analysis_text = f"""
            <h3>{student.get('name')}的成绩分析</h3>
            <p>共参加考试: {len(history)}次</p>
            <p>平均得分率: {avg_score:.2f}%</p>
            <p>成绩走势: 整体呈{trend_direction}趋势，最近一次考试相比前一次{recent_trend}</p>
            <p>班级排名: 最好排名 {min([h.get('class_rank', float('inf')) for h in history])}</p>
            """
    except Exception as e:
        print(f"调用AI分析出错: {str(e)}")
        # 备用生成方式
        analysis_text = f"""
        <h3>{student.get('name')}的成绩分析</h3>
        <p>共参加考试: {len(history)}次</p>
        <p>平均得分率: {avg_score:.2f}%</p>
        <p>成绩走势: 整体呈{trend_direction}趋势，最近一次考试相比前一次{recent_trend}</p>
        <p>班级排名: 最好排名 {min([h.get('class_rank', float('inf')) for h in history])}</p>
        """
    
    return {
        'text': analysis_text,
        'chart': chart_data
    }

# 辅助函数 - 考试分析
def generate_exam_ai_analysis(exam_data):
    """生成考试的AI分析报告"""
    basic_stats = exam_data.get('basic_stats', {})
    distribution = exam_data.get('distribution', [])
    questions = exam_data.get('questions', [])
    
    # 获取关键指标
    avg_score = basic_stats.get('average_score', 0)
    pass_rate = basic_stats.get('pass_rate', 0)
    highest_score = basic_stats.get('highest_score', 0)
    total_students = basic_stats.get('total_students', 0)
    
    # 构建图表数据 - 使用分布数据
    chart_data = {
        'type': 'bar',
        'data': {
            'labels': [d.get('score_range', '') for d in distribution],
            'datasets': [{
                'label': '学生人数',
                'data': [d.get('student_count', 0) for d in distribution],
                'backgroundColor': 'rgba(75, 192, 192, 0.2)',
                'borderColor': 'rgba(75, 192, 192, 1)',
                'borderWidth': 1
            }]
        }
    }
    
    # 分析题目数据，找出最难和最简单的题目
    hardest_idx = "未知"
    easiest_idx = "未知"
    
    if questions:
        # 按难度系数排序
        sorted_by_difficulty = sorted(questions, key=lambda q: q.get('average_score_rate', 0))
        hardest = sorted_by_difficulty[0] if sorted_by_difficulty else None
        easiest = sorted_by_difficulty[-1] if sorted_by_difficulty else None
        
        if hardest and easiest:
            # 找出题目的顺序编号
            for i, q in enumerate(questions):
                if q['id'] == hardest['id']:
                    hardest_idx = i + 1
                if q['id'] == easiest['id']:
                    easiest_idx = i + 1
    
    # 找出分布最集中的区间
    max_count_range = max(distribution, key=lambda d: d.get('student_count', 0)) if distribution else {'score_range': '未知', 'student_count': 0}
    
    # 准备题目难度分析
    question_analysis = ""
    if questions:
        for i, q in enumerate(questions):
            question_analysis += (f"题目{i+1}: 平均得分率 {q.get('average_score_rate', 0)*100:.2f}%, " +
                                f"难度系数 {q.get('difficulty', 0):.2f}, " +
                                f"区分度 {q.get('discrimination', 0):.2f}\n")
    
    # 使用AI生成分析结果
    system_prompt = f"""你是一位教育分析专家，擅长分析考试数据。
请分析以下考试数据，给出专业、客观的分析:

考试名称: {basic_stats.get('exam_name', '未知考试')}
科目: {basic_stats.get('subject', '未知科目')}
考试日期: {basic_stats.get('exam_date', '未知日期')}
参考人数: {total_students}人
平均分: {avg_score}
及格率: {pass_rate}%
最高分: {highest_score}
分数分布: 主要集中在{max_count_range.get('score_range', '未知')}区间，共{max_count_range.get('student_count', 0)}人

题目分析:
{question_analysis}

最难题目序号: {hardest_idx}
最简单题目序号: {easiest_idx}

请提供详细的分析，包括考试总体情况、学生分数分布特点、题目难度分析以及教学建议。
以HTML格式输出，使用<h3>、<p>等标签。
"""
    
    user_prompt = f"请分析{basic_stats.get('exam_name', '考试')}的考试数据"
    
    try:
        # 调用AI生成分析文本
        ai_response = ai_new(system_prompt, user_prompt)
        if ai_response and len(ai_response) > 50:  # 确保有有效的返回
            analysis_text = ai_response
        else:
            # 备用生成方式
            analysis_text = f"""
            <h3>{basic_stats.get('exam_name', '考试')}智能分析</h3>
            <p>参考人数: {total_students}人</p>
            <p>平均分: {avg_score}</p>
            <p>及格率: {pass_rate}%</p>
            <p>最高分: {highest_score}</p>
            <p>分数分布: 主要集中在{max_count_range.get('score_range', '未知')}区间，共{max_count_range.get('student_count', 0)}人</p>
            <p>题目分析: 第{hardest_idx}题最难，第{easiest_idx}题最简单</p>
            """
    except Exception as e:
        print(f"调用AI分析出错: {str(e)}")
        # 备用生成方式
        analysis_text = f"""
        <h3>{basic_stats.get('exam_name', '考试')}智能分析</h3>
        <p>参考人数: {total_students}人</p>
        <p>平均分: {avg_score}</p>
        <p>及格率: {pass_rate}%</p>
        <p>最高分: {highest_score}</p>
        <p>分数分布: 主要集中在{max_count_range.get('score_range', '未知')}区间，共{max_count_range.get('student_count', 0)}人</p>
        <p>题目分析: 第{hardest_idx}题最难，第{easiest_idx}题最简单</p>
        """
    
    return {
        'text': analysis_text,
        'chart': chart_data
    }

# 报告生成函数
def generate_comparison_report(analysis_data, format_type, file_path):
    """生成学生比较分析报告"""
    try:
        # 根据报告类型生成相应格式的报告
        if format_type == 'pdf':
            generate_pdf_comparison_report(analysis_data, file_path)
        elif format_type == 'excel':
            generate_excel_comparison_report(analysis_data, file_path)
        elif format_type == 'word':
            generate_word_comparison_report(analysis_data, file_path)
        return True
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        traceback.print_exc()  # 打印完整错误堆栈
        return False

def generate_student_report(analysis_data, format_type, file_path):
    """生成学生个人分析报告"""
    try:
        # 根据报告类型生成相应格式的报告
        if format_type == 'pdf':
            generate_pdf_student_report(analysis_data, file_path)
        elif format_type == 'excel':
            generate_excel_student_report(analysis_data, file_path)
        elif format_type == 'word':
            generate_word_student_report(analysis_data, file_path)
        return True
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        traceback.print_exc()  # 打印完整错误堆栈
        return False

def generate_exam_report(analysis_data, format_type, file_path):
    """生成考试分析报告"""
    try:
        # 根据报告类型生成相应格式的报告
        if format_type == 'pdf':
            generate_pdf_exam_report(analysis_data, file_path)
        elif format_type == 'excel': 
            generate_excel_exam_report(analysis_data, file_path)
        elif format_type == 'word':
            generate_word_exam_report(analysis_data, file_path)
        return True
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        traceback.print_exc()  # 打印完整错误堆栈
        return False

# 各种格式的报告生成函数
def generate_pdf_comparison_report(data, file_path):
    """生成PDF格式的学生比较报告"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # 标题
        title = Paragraph("学生比较分析报告", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # 添加比较内容
        student1 = data.get('student1', {})
        student2 = data.get('student2', {})
        
        # 学生基本信息
        elements.append(Paragraph(f"学生1: {student1.get('name', 'N/A')}", styles['Heading2']))
        elements.append(Paragraph(f"学生2: {student2.get('name', 'N/A')}", styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        # 比较分析
        if 'comparison' in data:
            elements.append(Paragraph("比较分析", styles['Heading2']))
            elements.append(Paragraph(data['comparison'], styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # 成绩对比
        if 'scores' in data:
            elements.append(Paragraph("成绩对比", styles['Heading2']))
            scores_data = [['考试', '学生1', '学生2']]
            
            for exam in data.get('scores', []):
                scores_data.append([
                    exam.get('exam_name', 'N/A'),
                    str(exam.get('student1_score', 'N/A')),
                    str(exam.get('student2_score', 'N/A'))
                ])
            
            scores_table = Table(scores_data)
            scores_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(scores_table)
            elements.append(Spacer(1, 12))
        
        # 构建文档
        doc.build(elements)
        print(f"PDF比较报告已生成: {file_path}")
        return True
    except Exception as e:
        print(f"生成PDF比较报告失败: {str(e)}")
        return False

def generate_excel_comparison_report(data, file_path):
    """生成Excel格式的学生比较报告"""
    try:
        import pandas as pd
        
        # 创建Excel写入器
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            # 基本信息
            student1 = data.get('student1', {})
            student2 = data.get('student2', {})
            
            info_df = pd.DataFrame({
                '比较项': ['学生姓名'],
                '学生1': [student1.get('name', 'N/A')],
                '学生2': [student2.get('name', 'N/A')]
            })
            
            info_df.to_excel(writer, sheet_name='基本信息', index=False)
            
            # 成绩对比
            if 'scores' in data:
                scores_data = []
                for exam in data.get('scores', []):
                    scores_data.append({
                        '考试': exam.get('exam_name', 'N/A'),
                        f"{student1.get('name', '学生1')}分数": exam.get('student1_score', 'N/A'),
                        f"{student2.get('name', '学生2')}分数": exam.get('student2_score', 'N/A')
                    })
                
                if scores_data:
                    scores_df = pd.DataFrame(scores_data)
                    scores_df.to_excel(writer, sheet_name='成绩对比', index=False)
            
            # 比较分析
            if 'comparison' in data:
                analysis_df = pd.DataFrame({
                    '比较分析': [data['comparison']]
                })
                analysis_df.to_excel(writer, sheet_name='分析结果', index=False)
        
        print(f"Excel比较报告已生成: {file_path}")
        return True
    except Exception as e:
        print(f"生成Excel比较报告失败: {str(e)}")
        return False

def generate_word_comparison_report(data, file_path):
    """生成Word格式的学生比较报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('学生比较分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 学生基本信息
        student1 = data.get('student1', {})
        student2 = data.get('student2', {})
        
        doc.add_heading('基本信息', level=1)
        p = doc.add_paragraph()
        p.add_run(f"学生1: {student1.get('name', 'N/A')}").bold = True
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"学生2: {student2.get('name', 'N/A')}").bold = True
        
        # 比较分析
        if 'comparison' in data:
            doc.add_heading('比较分析', level=1)
            doc.add_paragraph(data['comparison'])
        
        # 成绩对比
        if 'scores' in data:
            doc.add_heading('成绩对比', level=1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # 表头
            headers = table.rows[0].cells
            headers[0].text = '考试'
            headers[1].text = '学生1'
            headers[2].text = '学生2'
            
            # 数据行
            for exam in data.get('scores', []):
                row = table.add_row().cells
                row[0].text = exam.get('exam_name', 'N/A')
                row[1].text = str(exam.get('student1_score', 'N/A'))
                row[2].text = str(exam.get('student2_score', 'N/A'))
        
        # 保存文档
        doc.save(file_path)
        print(f"Word比较报告已生成: {file_path}")
        return True
    except Exception as e:
        print(f"生成Word比较报告失败: {str(e)}")
        return False

def generate_pdf_student_report(data, file_path):
    """生成PDF格式的学生个人分析报告"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # 标题
        student_name = data.get('student_name', 'N/A')
        title = Paragraph(f"{student_name}的学习分析报告", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # 学生基本信息
        if 'student_info' in data:
            elements.append(Paragraph("学生信息", styles['Heading2']))
            info = data.get('student_info', {})
            elements.append(Paragraph(f"姓名: {student_name}", styles['Normal']))
            elements.append(Paragraph(f"学号: {info.get('student_id', 'N/A')}", styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # 分析内容
        if 'analysis' in data:
            elements.append(Paragraph("分析结果", styles['Heading2']))
            elements.append(Paragraph(data['analysis'], styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # 成绩情况
        if 'exam_scores' in data:
            elements.append(Paragraph("考试成绩", styles['Heading2']))
            scores_data = [['考试名称', '成绩', '考试时间']]
            
            for exam in data.get('exam_scores', []):
                scores_data.append([
                    exam.get('exam_name', 'N/A'),
                    str(exam.get('score', 'N/A')),
                    exam.get('exam_date', 'N/A')
                ])
            
            scores_table = Table(scores_data)
            scores_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(scores_table)
            elements.append(Spacer(1, 12))
        
        # 构建文档
        doc.build(elements)
        print(f"PDF学生报告已生成: {file_path}")
        return True
    except Exception as e:
        print(f"生成PDF学生报告失败: {str(e)}")
        return False

def generate_excel_student_report(data, file_path):
    """生成Excel格式的学生个人分析报告"""
    import pandas as pd
    
    # 安全地获取数据
    student = data.get('student', {'name': '学生'})
    history = data.get('history', [])
    
    # 创建Excel工作簿
    writer = pd.ExcelWriter(file_path, engine='xlsxwriter')
    
    # 创建基本信息表
    info_df = pd.DataFrame({
        '项目': ['分析类型', '生成时间', '学生姓名', '学号'],
        '内容': [
            '学生成绩分析',
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            student.get('name', ''),
            student.get('student_id', '')
        ]
    })
    
    info_df.to_excel(writer, sheet_name='基本信息', index=False)
    
    # 如果有历史成绩数据，创建成绩表
    if history:
        history_df = pd.DataFrame(history)
        history_df.to_excel(writer, sheet_name='历史成绩', index=False)
    
    # 保存并关闭文件
    writer.close()

def generate_word_student_report(data, file_path):
    """生成Word格式的学生个人分析报告"""
    try:
        import docx
        
        # 安全地获取数据
        student = data.get('student', {'name': '学生'})
        data = data.get('data', {})
        
        # 创建Word文档
        doc = docx.Document()
        
        # 添加标题
        doc.add_heading(f"{student.get('name', '学生')}的成绩分析报告", 0)
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 添加分析文本
        if data.get('text'):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(data['text'], 'html.parser')
            text_content = soup.get_text()
            
            for paragraph in text_content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        
        # 保存文档
        doc.save(file_path)
    except Exception as e:
        print(f"生成Word报告时发生错误: {str(e)}")
        # 创建一个简单的报告，即使docx库不可用
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"{student.get('name', '学生')}的成绩分析报告\n")
            f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            if data.get('text'):
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(data['text'], 'html.parser')
                f.write(soup.get_text())

def generate_pdf_exam_report(data, file_path):
    """生成PDF格式的考试分析报告"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 注册中文字体
        try:
            # 尝试注册微软雅黑字体
            pdfmetrics.registerFont(TTFont('MSYaHei', 'C:/Windows/Fonts/msyh.ttc'))
            default_font = 'MSYaHei'
        except:
            try:
                # 备选：尝试注册宋体
                pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                default_font = 'SimSun'
            except:
                # 如果都失败，使用默认字体，但中文可能显示为乱码
                default_font = 'Helvetica'
                print("警告：无法注册中文字体，PDF中的中文可能显示为乱码")
        
        # 创建自定义样式
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='ChineseTitle',
                                 fontName=default_font,
                                 fontSize=18,
                                 leading=22,
                                 alignment=1))  # 居中对齐
        
        styles.add(ParagraphStyle(name='ChineseHeading',
                                 fontName=default_font,
                                 fontSize=14,
                                 leading=18))
        
        styles.add(ParagraphStyle(name='ChineseNormal',
                                 fontName=default_font,
                                 fontSize=10,
                                 leading=14))
        
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        elements = []
        
        # 标题
        title = Paragraph("考试分析报告", styles['ChineseTitle'])
        elements.append(title)
        elements.append(Spacer(1, 20))
        
        # 考试基本信息
        elements.append(Paragraph("考试信息", styles['ChineseHeading']))
        elements.append(Spacer(1, 6))
        
        exam_name = data.get('exam_name', 'N/A')
        elements.append(Paragraph(f"考试名称: {exam_name}", styles['ChineseNormal']))
        elements.append(Spacer(1, 6))
        
        # 统计数据
        elements.append(Paragraph("统计数据", styles['ChineseHeading']))
        elements.append(Spacer(1, 6))
        
        stats = data.get('stats', {})
        elements.append(Paragraph(f"参考人数: {stats.get('total_students', 0)}", styles['ChineseNormal']))
        elements.append(Paragraph(f"平均分: {stats.get('average_score', 0.00)}", styles['ChineseNormal']))
        elements.append(Paragraph(f"及格率: {stats.get('pass_rate', '0%')}", styles['ChineseNormal']))
        elements.append(Paragraph(f"最高分: {stats.get('highest_score', 0)}", styles['ChineseNormal']))
        elements.append(Spacer(1, 10))
        
        # 分析文本
        elements.append(Paragraph("分析", styles['ChineseHeading']))
        elements.append(Spacer(1, 6))
        
        analysis_text = data.get('text', '无分析数据')
        # 如果HTML标签存在，转换为纯文本
        if '<' in analysis_text and '>' in analysis_text:
            import re
            analysis_text = re.sub(r'<[^>]+>', '', analysis_text)
        elements.append(Paragraph(analysis_text, styles['ChineseNormal']))
        elements.append(Spacer(1, 10))
        
        # 题目分析
        elements.append(Paragraph("题目分析", styles['ChineseHeading']))
        elements.append(Spacer(1, 6))
        
        # 处理questions数据，确保它是字符串
        questions_data = data.get('questions', [])
        if isinstance(questions_data, list):
            # 如果是列表，将其转换为字符串
            questions_str = "\n".join([str(q) for q in questions_data]) if questions_data else "无题目数据"
        else:
            # 如果已经是字符串，直接使用
            questions_str = questions_data if questions_data else "无题目数据"
        elements.append(Paragraph(questions_str, styles['ChineseNormal']))
        
        # 构建文档
        doc.build(elements)
        print(f"PDF考试报告已生成: {file_path}")
        return True
    except Exception as e:
        print(f"生成PDF考试报告失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def generate_excel_exam_report(data, file_path):
    """生成Excel格式的考试分析报告"""
    import pandas as pd
    
    # 安全地获取数据
    raw_data = data.get('raw_data', {})
    basic_stats = raw_data.get('basic_stats', {})
    distribution = raw_data.get('distribution', [])
    
    # 创建Excel工作簿
    writer = pd.ExcelWriter(file_path, engine='xlsxwriter')
    
    # 创建基本信息表
    info_df = pd.DataFrame({
        '项目': ['考试名称', '学科', '考试日期', '总分', '平均分', '及格率', '最高分', '参考人数'],
        '内容': [
            basic_stats.get('exam_name', ''),
            basic_stats.get('subject', ''),
            basic_stats.get('exam_date', ''),
            basic_stats.get('total_score', 0),
            basic_stats.get('average_score', 0),
            f"{basic_stats.get('pass_rate', 0)}%",
            basic_stats.get('highest_score', 0),
            basic_stats.get('total_students', 0)
        ]
    })
    
    info_df.to_excel(writer, sheet_name='考试信息', index=False)
    
    # 如果有分布数据，创建分布表
    if distribution:
        dist_df = pd.DataFrame(distribution)
        dist_df.to_excel(writer, sheet_name='分数分布', index=False)
    
    # 保存并关闭文件
    writer.close()

def generate_word_exam_report(data, file_path):
    """生成Word格式的考试分析报告"""
    try:
        import docx
        
        # 安全地获取数据
        data = data.get('data', {})
        raw_data = data.get('raw_data', {})
        basic_stats = raw_data.get('basic_stats', {})
        
        # 创建Word文档
        doc = docx.Document()
        
        # 添加标题
        exam_name = basic_stats.get('exam_name', '考试')
        doc.add_heading(f"{exam_name}的分析报告", 0)
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # 添加分析文本
        if data.get('text'):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(data['text'], 'html.parser')
            text_content = soup.get_text()
            
            for paragraph in text_content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        
        # 保存文档
        doc.save(file_path)
    except Exception as e:
        print(f"生成Word报告时发生错误: {str(e)}")
        # 创建一个简单的报告，即使docx库不可用
        with open(file_path, 'w', encoding='utf-8') as f:
            exam_name = basic_stats.get('exam_name', '考试')
            f.write(f"{exam_name}的分析报告\n")
            f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            if data.get('text'):
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(data['text'], 'html.parser')
                f.write(soup.get_text())

# 新增函数：生成Python考试分析
def generate_python_exam_analysis(student, exam_data):
    """生成特定学生的Python考试分析"""
    # 构建图表数据
    chart_data = {
        'type': 'bar',
        'data': {
            'labels': ['得分率'],
            'datasets': [{
                'label': exam_data['exam_name'],
                'data': [exam_data['score_percentage']],
                'backgroundColor': 'rgba(75, 192, 192, 0.2)',
                'borderColor': 'rgba(75, 192, 192, 1)',
                'borderWidth': 1
            }]
        }
    }
    
    # 使用AI生成Python考试分析
    system_prompt = f"""你是一位Python教育专家，擅长分析学生的Python考试成绩。
请针对以下学生的Python考试数据，给出专业、详细的分析:

学生: {student['name']} (学号: {student['student_id']})
考试: {exam_data['exam_name']}
得分: {exam_data['score']}/{exam_data['total_score']} (得分率: {exam_data['score_percentage']}%)
班级排名: {exam_data.get('class_rank', 'N/A')}
考试日期: {exam_data.get('exam_date', 'N/A')}

请分析以下内容:
1. 整体表现评价
2. Python学习情况分析
3. 学习建议和改进方向
4. 针对Python编程的具体提高策略

以HTML格式输出，使用<h3>、<p>、<ul>、<li>等标签，使分析内容清晰易读。
"""
    
    user_prompt = f"请分析{student['name']}的Python考试情况"
    
    try:
        # 调用AI生成分析文本
        ai_response = ai_new(system_prompt, user_prompt)
        if ai_response and len(ai_response) > 50:  # 确保有有效的返回
            analysis_text = ai_response
        else:
            # 备用生成方式
            analysis_text = f"""
            <h3>{student['name']}的Python考试分析</h3>
            <p>考试名称: {exam_data['exam_name']}</p>
            <p>得分: {exam_data['score']}/{exam_data['total_score']} (得分率: {exam_data['score_percentage']}%)</p>
            <p>班级排名: {exam_data.get('class_rank', 'N/A')}</p>
            <h3>学习建议</h3>
            <p>建议针对Python基础概念进行复习，特别是语法和数据结构部分。</p>
            <p>建议多做练习题，提高编程实践能力。</p>
            """
    except Exception as e:
        print(f"调用AI Python分析出错: {str(e)}")
        # 备用生成方式
        analysis_text = f"""
        <h3>{student['name']}的Python考试分析</h3>
        <p>考试名称: {exam_data['exam_name']}</p>
        <p>得分: {exam_data['score']}/{exam_data['total_score']} (得分率: {exam_data['score_percentage']}%)</p>
        <p>班级排名: {exam_data.get('class_rank', 'N/A')}</p>
        <h3>学习建议</h3>
        <p>建议针对Python基础概念进行复习，特别是语法和数据结构部分。</p>
        <p>建议多做练习题，提高编程实践能力。</p>
        """
    
    return {
        'text': analysis_text,
        'chart': chart_data
    }

# 增加路由用于下载报告
@ai_analysis.route('/api/ai/download-report/<filename>')
def download_ai_report(filename):
    """下载AI分析报告"""
    try:
        reports_dir = os.path.join(os.getcwd(), 'exports')
        file_path = os.path.join(reports_dir, filename)
        
        print(f"尝试下载文件: {file_path}")
        
        # 验证文件是否存在
        if not os.path.exists(file_path):
            print(f"文件不存在: {file_path}")
            # 列出目录中的所有文件，以便调试
            print(f"目录 {reports_dir} 中的文件: {os.listdir(reports_dir)}")
            return jsonify({'error': '请求的文件不存在'}), 404
            
        print(f"文件存在，准备发送: {file_path}")
        return send_from_directory(reports_dir, filename, as_attachment=True)
    except Exception as e:
        print(f"文件下载失败: {str(e)}")
        traceback.print_exc()  # 打印完整的错误堆栈
        return jsonify({'error': '文件下载失败: ' + str(e)}), 404